"""
障碍者福祉事業所 運営ガイド - テンプレート一括生成スクリプト
Excel・Word テンプレートを自動生成します
"""

import os
from openpyxl import Workbook
from openpyxl.styles import (
    Font, PatternFill, Alignment, Border, Side, numbers
)
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))

# ===== ヘルパー =====
TEAL  = "0E7490"
WHITE = "FFFFFF"
LIGHT = "CFFAFE"
GRAY  = "F8FAFC"
WARN  = "FEF3C7"
RED   = "FEE2E2"

def hdr_fill(): return PatternFill("solid", fgColor=TEAL)
def sub_fill(): return PatternFill("solid", fgColor=LIGHT)
def warn_fill(): return PatternFill("solid", fgColor=WARN)
def red_fill():  return PatternFill("solid", fgColor=RED)

def hdr_font(sz=10, bold=True, color=WHITE):
    return Font(name="Meiryo UI", size=sz, bold=bold, color=color)

def body_font(sz=9, bold=False):
    return Font(name="Meiryo UI", size=sz, bold=bold)

def thin_border():
    s = Side(style="thin")
    return Border(left=s, right=s, top=s, bottom=s)

def set_col(ws, col_letter, width):
    ws.column_dimensions[col_letter].width = width

def set_row(ws, row_num, height):
    ws.row_dimensions[row_num].height = height

def header_row(ws, row, cols, fills=None):
    """cols = list of (letter, text, width)"""
    for letter, text, width in cols:
        c = ws[f"{letter}{row}"]
        c.value = text
        c.font = hdr_font()
        c.fill = fills(letter) if fills else hdr_fill()
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border = thin_border()
        ws.column_dimensions[letter].width = width
    set_row(ws, row, 30)

def body_cell(ws, ref, value="", bold=False, wrap=True, halign="left"):
    c = ws[ref]
    c.value = value
    c.font = body_font(bold=bold)
    c.alignment = Alignment(horizontal=halign, vertical="center", wrap_text=wrap)
    c.border = thin_border()
    return c

def example_cell(ws, ref, value=""):
    c = ws[ref]
    c.value = value
    c.font = Font(name="Meiryo UI", size=9, color="6B7280", italic=True)
    c.fill = PatternFill("solid", fgColor=GRAY)
    c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    c.border = thin_border()
    return c

def save(wb, folder, filename):
    path = os.path.join(BASE, folder, filename)
    wb.save(path)
    print(f"✅ {folder}/{filename}")

# ======================================================================
# 02_訪問看護記録様式
# ======================================================================

def make_houmon_kiroku2():
    """訪問看護記録書Ⅱ（経過記録）"""
    wb = Workbook()
    ws = wb.active
    ws.title = "訪問看護記録書Ⅱ"
    ws.page_setup.orientation = "landscape"

    # タイトル
    ws.merge_cells("A1:L1")
    t = ws["A1"]
    t.value = "訪問看護記録書Ⅱ（経過記録）"
    t.font = Font(name="Meiryo UI", size=14, bold=True, color=TEAL)
    t.alignment = Alignment(horizontal="center", vertical="center")
    set_row(ws, 1, 28)

    # 基本情報行
    ws.merge_cells("A2:C2"); ws["A2"].value = "利用者氏名："
    ws.merge_cells("D2:F2")
    ws.merge_cells("G2:H2"); ws["G2"].value = "生年月日："
    ws.merge_cells("I2:J2")
    ws.merge_cells("K2:L2"); ws["K2"].value = f"（　　　）年度"
    for r in ["A2","G2","K2"]:
        ws[r].font = body_font(bold=True)
        ws[r].alignment = Alignment(horizontal="right", vertical="center")
    set_row(ws, 2, 20)

    # ヘッダー行
    cols = [
        ("A","訪問日",8), ("B","開始時刻",8), ("C","終了時刻",8),
        ("D","体温\n(℃)",6), ("E","血圧\n(mmHg)",8), ("F","脈拍\n(/分)",6),
        ("G","SpO2\n(%)",6), ("H","呼吸\n(/分)",6),
        ("I","訪問看護の内容・利用者の状態・反応（SOAPまたは自由記述）",40),
        ("J","特記事項・連絡事項",20),
        ("K","次回訪問予定",10), ("L","記録者",8)
    ]
    header_row(ws, 3, cols)

    # データ行（20行）+ 記入例1行
    example_data = [
        "5/1(木)", "10:00", "11:00", "36.5", "120/80", "72", "98", "18",
        "【S】特に問題なし。昨日の夜は眠れたとのこと。【O】顔色良好。痰の喀出あり。吸引実施、少量の白色痰。【A】バイタル安定。誤嚥リスク継続。【P】引き続き観察継続。",
        "主治医への報告不要。次回も同様のケアを継続。",
        "5/8(木) 10:00", "山田 看護師"
    ]
    letters = [get_column_letter(i) for i in range(1, 13)]
    example_cell(ws, f"A4", "（記入例）")
    for i, val in enumerate(example_data):
        example_cell(ws, f"{letters[i]}4", val)

    for row in range(5, 25):
        for letter in letters:
            body_cell(ws, f"{letter}{row}")
        set_row(ws, row, 45)

    # 備考
    ws.merge_cells(f"A25:L25")
    ws["A25"].value = "※ 修正は二重線で消し、訂正者の署名・日付を記入してください。修正液（白塗り）の使用は禁止です。"
    ws["A25"].font = Font(name="Meiryo UI", size=8, color="DC2626", bold=True)
    ws["A25"].alignment = Alignment(horizontal="left", vertical="center")
    set_row(ws, 25, 18)

    save(wb, "02_訪問看護記録様式", "03_訪問看護記録書Ⅱ経過記録.xlsx")


def make_houmon_houkoku():
    """訪問看護報告書"""
    wb = Workbook()
    ws = wb.active
    ws.title = "訪問看護報告書"

    ws.merge_cells("A1:H1")
    ws["A1"].value = "訪問看護報告書"
    ws["A1"].font = Font(name="Meiryo UI", size=14, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    set_row(ws, 1, 28)

    # 基本情報
    info_rows = [
        ("利用者氏名", "様", "報告対象月", "　　年　　月分"),
        ("主治医氏名", "先生", "訪問看護ステーション名", ""),
        ("疾患・傷病名", "", "管理者名", ""),
        ("訪問看護指示書 有効期限", "", "担当看護師", ""),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info_rows, start=2):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value = l2
        ws.merge_cells(f"G{i}:H{i}"); ws[f"G{i}"].value = v2
        for cell in [f"A{i}", f"E{i}"]:
            ws[cell].font = body_font(bold=True)
            ws[cell].fill = sub_fill()
            ws[cell].alignment = Alignment(horizontal="right", vertical="center")
            ws[cell].border = thin_border()
        for cell in [f"C{i}", f"G{i}"]:
            ws[cell].font = body_font()
            ws[cell].alignment = Alignment(horizontal="left", vertical="center")
            ws[cell].border = thin_border()
        set_row(ws, i, 22)

    sections = [
        (6, "訪問回数・状況", 25,
         "当月の訪問回数：　　回　 / うち緊急訪問：　　回 / 特別訪問看護指示期間：　なし　あり（　　日〜　　日）"),
        (8, "病状・バイタルサインの経過", 60, "・体温：〜℃台　・血圧：/〜/mmHg台　・脈拍：〜回/分台\n・意識レベル：清明 / 傾眠 / その他（　　　　）\n・ADL：（　　）\n・特記すべき病状変化："),
        (11, "看護・リハビリの内容", 60, "・清拭・入浴介助：　有　無\n・創傷処置：　有（部位：　　）　無\n・服薬管理・指導：　有　無\n・吸引：　有　無\n・褥瘡処置：　有（部位：　　）　無\n・その他："),
        (14, "家族への指導・支援内容", 40, ""),
        (16, "今後の看護・支援の方向性（主治医へ）", 50, ""),
        (18, "特記事項・連絡事項", 40, ""),
    ]

    for row, title, height, example in sections:
        ws.merge_cells(f"A{row}:H{row}")
        ws[f"A{row}"].value = title
        ws[f"A{row}"].font = hdr_font(sz=9)
        ws[f"A{row}"].fill = hdr_fill()
        ws[f"A{row}"].alignment = Alignment(horizontal="left", vertical="center")
        ws[f"A{row}"].border = thin_border()
        set_row(ws, row, 20)

        ws.merge_cells(f"A{row+1}:H{row+1}")
        c = ws[f"A{row+1}"]
        c.value = example
        c.font = Font(name="Meiryo UI", size=9, color="6B7280", italic=True)
        c.fill = PatternFill("solid", fgColor=GRAY)
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        c.border = thin_border()
        set_row(ws, row+1, height)

    # 署名欄
    sign_row = 20
    ws.merge_cells(f"A{sign_row}:H{sign_row}")
    ws[f"A{sign_row}"].value = "報告日：　　年　　月　　日　　担当看護師署名：　　　　　　　　　　管理者確認：　　　　　　　"
    ws[f"A{sign_row}"].font = body_font()
    ws[f"A{sign_row}"].alignment = Alignment(horizontal="left", vertical="center")
    ws[f"A{sign_row}"].border = thin_border()
    set_row(ws, sign_row, 30)

    for col_letter, width in [("A",12),("B",12),("C",12),("D",12),("E",12),("F",12),("G",12),("H",12)]:
        ws.column_dimensions[col_letter].width = width

    save(wb, "02_訪問看護記録様式", "04_訪問看護報告書.xlsx")


def make_jisseki():
    """サービス提供実績記録（月次）"""
    wb = Workbook()
    ws = wb.active
    ws.title = "サービス提供実績"
    ws.page_setup.orientation = "landscape"

    ws.merge_cells("A1:AH1")
    ws["A1"].value = "サービス提供実績記録（月次）　　　　年　　月分"
    ws["A1"].font = Font(name="Meiryo UI", size=12, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    set_row(ws, 1, 24)

    ws["A2"].value = "利用者氏名"
    ws["A2"].font = body_font(bold=True)
    ws["B2"].value = "サービス種別"
    ws["B2"].font = body_font(bold=True)
    for i, day in enumerate(range(1, 32), start=3):
        c = ws.cell(row=2, column=i)
        c.value = str(day)
        c.font = hdr_font(sz=8)
        c.fill = hdr_fill()
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = thin_border()
        ws.column_dimensions[get_column_letter(i)].width = 3.5
    c = ws.cell(row=2, column=34); c.value = "利用日数"
    c.font = hdr_font(sz=8); c.fill = hdr_fill()
    c.alignment = Alignment(horizontal="center", vertical="center")
    c.border = thin_border()
    ws.column_dimensions[get_column_letter(34)].width = 6
    for cell in [ws["A2"], ws["B2"]]:
        cell.fill = hdr_fill()
        cell.font = hdr_font(sz=8)
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin_border()
    ws.column_dimensions["A"].width = 14
    ws.column_dimensions["B"].width = 12
    set_row(ws, 2, 22)

    for row in range(3, 23):
        for col in range(1, 35):
            c = ws.cell(row=row, column=col)
            c.border = thin_border()
            c.font = body_font(sz=8)
            c.alignment = Alignment(horizontal="center", vertical="center")
        if row == 3:
            ws.cell(row=3, column=1).value = "（記入例）山田 太郎"
            ws.cell(row=3, column=1).font = Font(name="Meiryo UI", size=8, color="6B7280")
            ws.cell(row=3, column=2).value = "訪問看護"
            ws.cell(row=3, column=2).font = Font(name="Meiryo UI", size=8, color="6B7280")
            ws.cell(row=3, column=3).value = "○"
            ws.cell(row=3, column=5).value = "○"
            ws.cell(row=3, column=34).value = 2
        set_row(ws, row, 18)

    save(wb, "02_訪問看護記録様式", "05_サービス提供実績記録（月次）.xlsx")


# ======================================================================
# 06_事故苦情ヒヤリ様式
# ======================================================================

def make_jiko():
    """事故報告書"""
    wb = Workbook()
    ws = wb.active
    ws.title = "事故報告書"

    ws.merge_cells("A1:H1")
    ws["A1"].value = "事故報告書"
    ws["A1"].font = Font(name="Meiryo UI", size=14, bold=True, color=WHITE)
    ws["A1"].fill = PatternFill("solid", fgColor="DC2626")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    set_row(ws, 1, 28)

    rows_def = [
        (2, "発生日時", "　　年　　月　　日（　曜日）　　時　　分頃", "事業所名", ""),
        (3, "発生場所", "", "担当者氏名", ""),
        (4, "当事者氏名（利用者）", "", "当事者の障害種別", "身体 ・ 精神 ・ 知的"),
        (5, "事故の種類", "転倒 ・ 転落 ・ 誤薬 ・ 行方不明 ・ 医療事故 ・ その他（　　　　）", "重症度", "軽傷 ・ 中等症 ・ 重症 ・ 死亡"),
    ]
    for rn, l1, v1, l2, v2 in rows_def:
        ws.merge_cells(f"A{rn}:B{rn}"); ws[f"A{rn}"].value = l1
        ws.merge_cells(f"C{rn}:D{rn}"); ws[f"C{rn}"].value = v1
        ws.merge_cells(f"E{rn}:F{rn}"); ws[f"E{rn}"].value = l2
        ws.merge_cells(f"G{rn}:H{rn}"); ws[f"G{rn}"].value = v2
        for cell in [f"A{rn}", f"E{rn}"]:
            ws[cell].font = body_font(bold=True); ws[cell].fill = sub_fill()
            ws[cell].alignment = Alignment(horizontal="right", vertical="center")
            ws[cell].border = thin_border()
        for cell in [f"C{rn}", f"G{rn}"]:
            ws[cell].font = body_font(); ws[cell].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[cell].border = thin_border()
        set_row(ws, rn, 22)

    text_sections = [
        (6, "【事故の経緯・状況】（発生時の状況を時系列で具体的に記述してください）", 120,
         "例）10:30頃、訪問看護実施中に利用者が浴室内で転倒。右膝に擦り傷あり。本人は「足が滑った」と話していた。"),
        (8, "【緊急対応の内容】（発見後の対応・連絡先・対応時刻）", 80,
         "例）10:31 傷の確認・止血処置実施。10:35 管理者へ報告。10:40 主治医（〇〇クリニック）へ連絡・指示を受ける。10:45 家族へ報告・受診不要の判断。"),
        (10, "【原因分析】（なぜ起きたか。人・設備・環境・手順等の観点で）", 60,
         "例）浴室マットの滑り止めが劣化していた。本人の下肢筋力低下があり、見守りが必要な状況だった。"),
        (12, "【再発防止策】（具体的な対策と実施期限）", 60,
         "例）①滑り止めマットを新品に交換（5/10まで）②訪問中の浴室使用時は必ず見守りを実施③スタッフへの周知徹底（5/9朝礼）"),
        (14, "【行政への報告】", 30,
         "要　・　不要　／　報告先：　　　　　　　　報告日：　　年　　月　　日"),
    ]

    for rn, title, height, example in text_sections:
        ws.merge_cells(f"A{rn}:H{rn}"); ws[f"A{rn}"].value = title
        ws[f"A{rn}"].font = hdr_font(sz=9); ws[f"A{rn}"].fill = hdr_fill()
        ws[f"A{rn}"].alignment = Alignment(horizontal="left", vertical="center")
        ws[f"A{rn}"].border = thin_border(); set_row(ws, rn, 20)

        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c = ws[f"A{rn+1}"]
        c.value = example; c.font = Font(name="Meiryo UI", size=9, color="6B7280", italic=True)
        c.fill = PatternFill("solid", fgColor=GRAY)
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        c.border = thin_border(); set_row(ws, rn+1, height)

    sign = 16
    ws.merge_cells(f"A{sign}:D{sign}"); ws[f"A{sign}"].value = "報告者氏名・日付：　　　　　　　　　　　　年　　月　　日"
    ws.merge_cells(f"E{sign}:H{sign}"); ws[f"E{sign}"].value = "管理者確認・日付：　　　　　　　　　　　　年　　月　　日"
    for cell in [f"A{sign}", f"E{sign}"]:
        ws[cell].font = body_font(); ws[cell].border = thin_border()
        ws[cell].alignment = Alignment(horizontal="left", vertical="center")
    set_row(ws, sign, 28)

    for col, width in zip("ABCDEFGH", [14,14,14,14,14,14,14,14]):
        ws.column_dimensions[col].width = width

    save(wb, "06_事故苦情ヒヤリ様式", "01_事故報告書.xlsx")


def make_hiyari():
    """ヒヤリハット報告書"""
    wb = Workbook()
    ws = wb.active
    ws.title = "ヒヤリハット報告書"

    ws.merge_cells("A1:F1")
    ws["A1"].value = "ヒヤリハット報告書"
    ws["A1"].font = Font(name="Meiryo UI", size=14, bold=True, color=WHITE)
    ws["A1"].fill = PatternFill("solid", fgColor="D97706")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    set_row(ws, 1, 28)

    top_rows = [
        (2, "発生日時", "　　年　　月　　日（　）　　時　　分", "報告者氏名", ""),
        (3, "発生場所", "", "当事者氏名（利用者）", ""),
        (4, "ヒヤリハットの種類", "転倒ヒヤリ ・ 誤薬ヒヤリ ・ 急変ヒヤリ ・ その他（　　　）", "深刻度",
         "①繰り返す可能性あり  ②重大事故に繋がる可能性あり"),
    ]
    for rn, l1, v1, l2, v2 in top_rows:
        ws.merge_cells(f"A{rn}:B{rn}"); ws[f"A{rn}"].value = l1
        ws.merge_cells(f"C{rn}:D{rn}"); ws[f"C{rn}"].value = v1
        ws.merge_cells(f"E{rn}:F{rn}"); ws[f"E{rn}"].value = l2
        ws.merge_cells(f"G{rn}:H{rn}"); ws[f"G{rn}"].value = v2
        for cell in [f"A{rn}", f"E{rn}"]:
            ws[cell].font = body_font(bold=True); ws[cell].fill = sub_fill()
            ws[cell].alignment = Alignment(horizontal="right", vertical="center")
            ws[cell].border = thin_border()
        for cell in [f"C{rn}", f"G{rn}"]:
            ws[cell].font = body_font(); ws[cell].border = thin_border()
            ws[cell].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        set_row(ws, rn, 22)

    sections = [
        (5, "【発生した状況】（どのような状況でヒヤリハットが起きたか）", 70,
         "例）服薬介助の際に、同日処方された2種類の薬のうち1種類の投与を失念しそうになった。利用者本人からの指摘で気づいた。"),
        (7, "【なぜヒヤリとしたか・原因分析】", 50,
         "例）新しい薬が追加されたことを担当者間で共有できていなかった。薬の一覧表が最新版に更新されていなかった。"),
        (9, "【対応策・再発防止策】（具体的な対策と期限）", 50,
         "例）①薬の一覧表を即日更新し、全担当者にメール共有。②今後の薬の変更は必ず記録簿に即日記入する運用に変更。"),
    ]

    for rn, title, height, example in sections:
        ws.merge_cells(f"A{rn}:H{rn}"); ws[f"A{rn}"].value = title
        ws[f"A{rn}"].font = hdr_font(sz=9)
        ws[f"A{rn}"].fill = PatternFill("solid", fgColor="D97706")
        ws[f"A{rn}"].alignment = Alignment(horizontal="left", vertical="center")
        ws[f"A{rn}"].border = thin_border(); set_row(ws, rn, 20)
        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c = ws[f"A{rn+1}"]
        c.value = example; c.font = Font(name="Meiryo UI", size=9, color="6B7280", italic=True)
        c.fill = PatternFill("solid", fgColor=GRAY)
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        c.border = thin_border(); set_row(ws, rn+1, height)

    ws.merge_cells("A11:H11")
    ws["A11"].value = "報告日：　　年　　月　　日　　報告者署名：　　　　　　　　管理者確認：　　　　　　（　　年　　月　　日）"
    ws["A11"].font = body_font(); ws["A11"].border = thin_border()
    ws["A11"].alignment = Alignment(horizontal="left", vertical="center")
    set_row(ws, 11, 28)

    for col, width in zip("ABCDEFGH", [14,14,14,14,14,14,14,14]):
        ws.column_dimensions[col].width = width

    save(wb, "06_事故苦情ヒヤリ様式", "02_ヒヤリハット報告書.xlsx")


def make_kujo():
    """苦情受付記録"""
    wb = Workbook()
    ws = wb.active
    ws.title = "苦情受付記録"

    ws.merge_cells("A1:G1")
    ws["A1"].value = "苦情受付記録"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    set_row(ws, 1, 26)

    ws.merge_cells("A2:G2")
    ws["A2"].value = "※ 苦情が「0件」の月も「0件」と記録してください。全件を記録・保管することが運営指導の確認対象です。"
    ws["A2"].font = Font(name="Meiryo UI", size=9, color="DC2626", bold=True)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center")
    set_row(ws, 2, 20)

    cols = [
        ("A","受付日時",14), ("B","申立者\n（利用者/家族/その他）",16), ("C","申立者との関係",10),
        ("D","苦情の内容",40), ("E","対応内容・対応日",30), ("F","解決状況",12), ("G","担当者",10)
    ]
    header_row(ws, 3, cols)

    # 記入例
    ex = ["2026/4/15 14:30", "利用者家族（長男）", "家族", "「職員の言葉が乱暴だ」との申し出。特定の職員の敬語使用が不十分とのこと。",
          "4/15：管理者が電話で謝罪・状況確認。4/16：当該職員と面談・指導実施。4/18：長男へ改善報告の電話を実施。",
          "解決（4/18）", "山田 管理者"]
    letters = "ABCDEFG"
    for i, val in enumerate(ex):
        example_cell(ws, f"{letters[i]}4", val)
    set_row(ws, 4, 50)

    for row in range(5, 25):
        for letter in letters:
            body_cell(ws, f"{letter}{row}")
        set_row(ws, row, 45)

    save(wb, "06_事故苦情ヒヤリ様式", "03_苦情受付記録.xlsx")


def make_kinkyu():
    """緊急連絡先一覧"""
    wb = Workbook()
    ws = wb.active
    ws.title = "緊急連絡先一覧"

    ws.merge_cells("A1:E1")
    ws["A1"].value = "緊急連絡先一覧"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    set_row(ws, 1, 26)

    ws.merge_cells("A2:C2"); ws["A2"].value = "利用者氏名："
    ws.merge_cells("D2:E2"); ws["D2"].value = "作成日：　　年　　月　　日　 最終更新：　　年　　月　　日"
    ws["A2"].font = body_font(bold=True); ws["D2"].font = body_font()
    set_row(ws, 2, 20)

    cols = [("A","区分",12),("B","氏名・機関名",20),("C","続柄・役割",12),
            ("D","電話番号（日中）",16),("E","電話番号（夜間・緊急）",16)]
    header_row(ws, 3, cols)

    rows = [
        ("家族①（緊急時最優先）","","","",""),
        ("家族②","","","",""),
        ("家族③","","","",""),
        ("主治医","","","",""),
        ("緊急対応病院","","","",""),
        ("担当相談支援専門員","","","",""),
        ("市区町村（障害福祉担当）","","","",""),
        ("事業所管理者（夜間）","","","",""),
    ]
    for i, (cat, *rest) in enumerate(rows, start=4):
        c = ws[f"A{i}"]
        c.value = cat
        c.font = body_font(bold=True); c.fill = sub_fill()
        c.border = thin_border(); c.alignment = Alignment(horizontal="left", vertical="center")
        for j, val in enumerate(rest):
            body_cell(ws, f"{'BCDE'[j]}{i}", val)
        set_row(ws, i, 22)

    save(wb, "06_事故苦情ヒヤリ様式", "04_緊急連絡先一覧.xlsx")


# ======================================================================
# スタッフ管理系
# ======================================================================

def make_staff_meibo():
    """職員名簿・資格証台帳"""
    wb = Workbook()
    ws = wb.active
    ws.title = "職員名簿・資格台帳"

    ws.merge_cells("A1:J1")
    ws["A1"].value = "職員名簿・資格証台帳"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    set_row(ws, 1, 26)

    cols = [
        ("A","氏名",12), ("B","職種",12), ("C","雇用形態",10),
        ("D","入社日",10), ("E","資格・免許名",18),
        ("F","免許証番号",14), ("G","取得年月日",10), ("H","有効期限",10),
        ("I","更新研修\n期限",10), ("J","備考",16)
    ]
    header_row(ws, 2, cols)

    ex = ["山田 花子", "看護師", "常勤", "2022/4/1", "看護師免許", "第123456号", "2015/3/20", "なし", "—", "管理者兼務"]
    for i, val in enumerate(ex):
        example_cell(ws, f"{get_column_letter(i+1)}3", val)
    set_row(ws, 3, 22)

    for row in range(4, 24):
        for col in range(1, 11):
            body_cell(ws, f"{get_column_letter(col)}{row}")
        set_row(ws, row, 22)

    ws.merge_cells("A25:J25")
    ws["A25"].value = "【更新が必要な資格の例】サービス管理責任者現任研修：5年毎 / 相談支援専門員現任研修：5年毎 / 認定看護師：5年毎"
    ws["A25"].font = Font(name="Meiryo UI", size=8, color="6B7280")
    ws["A25"].alignment = Alignment(horizontal="left", vertical="center")
    set_row(ws, 25, 18)

    save(wb, "06_事故苦情ヒヤリ様式", "05_職員名簿・資格証台帳.xlsx")


def make_kenshu():
    """研修受講記録"""
    wb = Workbook()
    ws = wb.active
    ws.title = "研修受講記録"

    ws.merge_cells("A1:H1")
    ws["A1"].value = "研修受講記録"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    set_row(ws, 1, 26)

    ws.merge_cells("A2:D2"); ws["A2"].value = "事業所名："
    ws.merge_cells("E2:H2"); ws["E2"].value = "　　年度"
    ws["A2"].font = body_font(bold=True)
    set_row(ws, 2, 20)

    cols = [
        ("A","実施日",12), ("B","研修名・テーマ",28), ("C","主催・講師",14),
        ("D","形式",10), ("E","参加者氏名",20), ("F","参加者数",8),
        ("G","内容（要旨）",30), ("H","確認者",10)
    ]
    header_row(ws, 3, cols)

    ex_rows = [
        ["2026/4/10","虐待防止・権利擁護研修","管理者（社内）","集合","全職員",6,
         "虐待の定義・通報義務・日常的な権利擁護の重要性を確認。事例をもとにグループ討議実施。","山田 管理者"],
        ["2026/5/20","個人情報保護研修","管理者（社内）","集合","全職員",6,
         "個人情報保護法の改正点・要配慮個人情報の取扱い・漏洩時の対応を学習。","山田 管理者"],
        ["2026/6/15","感染症対策研修","管理者（社内）","集合","全職員",6,
         "標準予防策・手指衛生・PPEの正しい使用方法を確認。","山田 管理者"],
    ]
    letters = "ABCDEFGH"
    for i, row_data in enumerate(ex_rows, start=4):
        for j, val in enumerate(row_data):
            example_cell(ws, f"{letters[j]}{i}", str(val))
        set_row(ws, i, 45)

    for row in range(7, 27):
        for letter in letters:
            body_cell(ws, f"{letter}{row}")
        set_row(ws, row, 40)

    ws.merge_cells("A27:H27")
    ws["A27"].value = "【最低限年1回実施が必要な研修】虐待防止研修 / 個人情報保護研修 / 感染症対策研修 / 身体拘束防止研修"
    ws["A27"].font = Font(name="Meiryo UI", size=8, color="DC2626", bold=True)
    ws["A27"].alignment = Alignment(horizontal="left", vertical="center")
    set_row(ws, 27, 18)

    save(wb, "06_事故苦情ヒヤリ様式", "06_研修受講記録.xlsx")


# ======================================================================
# 就労支援記録
# ======================================================================

def make_kojin_shien():
    """個別支援計画（就労継続支援A型・B型用）追加様式"""
    wb = Workbook()
    ws = wb.active
    ws.title = "個別支援計画（就労支援版）"

    ws.merge_cells("A1:H1")
    ws["A1"].value = "個別支援計画書（就労継続支援A型・B型用）"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    set_row(ws, 1, 28)

    ws.merge_cells("A2:H2")
    ws["A2"].value = "※ 本様式は京都府版の補助様式です。公式様式（05_個別支援計画書様式（京都府）.docx）と併用してください。"
    ws["A2"].font = Font(name="Meiryo UI", size=8, color="DC2626")
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center")
    set_row(ws, 2, 18)

    info = [
        (3, "利用者氏名", "（様）", "生年月日", "　　年　　月　　日（　　歳）"),
        (4, "障害の種別・等級", "身体(　級) ・ 精神(　級) ・ 知的(　) ・ 難病", "障害支援区分", "区分　　　/認定日：　　年　　月"),
        (5, "計画作成日", "　　年　　月　　日", "計画期間", "　　年　　月　〜　　年　　月"),
        (6, "担当サービス管理責任者", "", "前回計画作成日", "　　年　　月　　日"),
    ]
    for rn, l1, v1, l2, v2 in info:
        ws.merge_cells(f"A{rn}:B{rn}"); ws[f"A{rn}"].value = l1
        ws.merge_cells(f"C{rn}:D{rn}"); ws[f"C{rn}"].value = v1
        ws.merge_cells(f"E{rn}:F{rn}"); ws[f"E{rn}"].value = l2
        ws.merge_cells(f"G{rn}:H{rn}"); ws[f"G{rn}"].value = v2
        for cell in [f"A{rn}", f"E{rn}"]:
            ws[cell].font = body_font(bold=True); ws[cell].fill = sub_fill()
            ws[cell].alignment = Alignment(horizontal="right", vertical="center"); ws[cell].border = thin_border()
        for cell in [f"C{rn}", f"G{rn}"]:
            ws[cell].font = body_font(); ws[cell].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[cell].border = thin_border()
        set_row(ws, rn, 22)

    text_secs = [
        (7, "【本人の希望・夢・目標】", 50, "（例）週4日、午後2時まで安定して作業できるようになりたい。将来的にパン屋でアルバイトしたい。"),
        (9, "【長期目標（6ヶ月〜1年）】", 40, "（例）毎日4時間の安定した就労ができる"),
        (11,"【短期目標（〜3ヶ月）】", 40, "（例）体調管理をして月間欠席2日以内を目指す"),
        (13,"【具体的な支援内容】", 80,
         "①日常的な体調確認と声かけ\n②休憩の取り方・体調悪化時の自己申告ができるよう支援\n③作業工程の分割・簡略化（本人の能力に合わせた調整）\n④一般就労への移行に向けた職場体験の機会を探る"),
        (15,"【医療・保健との連携】", 40, "主治医：〇〇クリニック　田中医師　/ 訪問頻度：月1回以上の情報共有"),
        (17,"【家族・居住環境との連携】", 40, ""),
        (19,"【本人の同意】", 30, "上記の個別支援計画について説明を受け、内容に同意します。\n\n署名：　　　　　　　　　　　年　　月　　日"),
    ]
    for rn, title, height, example in text_secs:
        ws.merge_cells(f"A{rn}:H{rn}"); ws[f"A{rn}"].value = title
        ws[f"A{rn}"].font = hdr_font(sz=9); ws[f"A{rn}"].fill = hdr_fill()
        ws[f"A{rn}"].alignment = Alignment(horizontal="left", vertical="center")
        ws[f"A{rn}"].border = thin_border(); set_row(ws, rn, 20)
        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c = ws[f"A{rn+1}"]
        c.value = example; c.font = Font(name="Meiryo UI", size=9, color="6B7280", italic=True)
        c.fill = PatternFill("solid", fgColor=GRAY)
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        c.border = thin_border(); set_row(ws, rn+1, height)

    for col, width in zip("ABCDEFGH", [14,14,14,14,14,14,14,14]):
        ws.column_dimensions[col].width = width

    save(wb, "03_就労支援記録様式", "06_個別支援計画書（就労継続支援版）.xlsx")


def make_kochin():
    """工賃台帳（B型）"""
    wb = Workbook()
    ws = wb.active
    ws.title = "工賃台帳"

    ws.merge_cells("A1:H1")
    ws["A1"].value = "工賃台帳（就労継続支援B型用）　　　　年　　月分"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    set_row(ws, 1, 26)

    ws.merge_cells("A2:H2")
    ws["A2"].value = "【法的根拠】障害者総合支援法 第43条 / 目標工賃：前年度平均工賃月額が目標値以上で「目標工賃達成加算（10単位/日）」を算定可能"
    ws["A2"].font = Font(name="Meiryo UI", size=8, color=TEAL, bold=True)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center")
    set_row(ws, 2, 20)

    cols = [("A","利用者氏名",14),("B","利用日数",8),("C","作業時間\n合計（h）",10),
            ("D","作業種別（主な作業）",16),("E","工賃単価\n（円/h）",10),
            ("F","当月工賃（円）",12),("G","累計工賃\n（年度内）",12),("H","支払日・確認",12)]
    header_row(ws, 3, cols)

    ex = [["山田 太郎",18,54,"封入作業・箱折り",100,5400,5400,"2026/5/25"],
          ["佐藤 花子",20,60,"清掃・洗濯補助",110,6600,6600,"2026/5/25"]]
    for i, row_data in enumerate(ex, start=4):
        for j, val in enumerate(row_data):
            example_cell(ws, f"{get_column_letter(j+1)}{i}", str(val))
        set_row(ws, i, 22)

    for row in range(6, 26):
        for col in range(1, 9):
            body_cell(ws, f"{get_column_letter(col)}{row}")
        set_row(ws, row, 22)

    total_row = 26
    ws.merge_cells(f"A{total_row}:E{total_row}")
    ws[f"A{total_row}"].value = "月間合計"
    ws[f"A{total_row}"].font = body_font(bold=True)
    ws[f"A{total_row}"].fill = sub_fill()
    ws[f"A{total_row}"].alignment = Alignment(horizontal="right", vertical="center")
    ws[f"A{total_row}"].border = thin_border()
    for col in ["F","G","H"]:
        body_cell(ws, f"{col}{total_row}")
    set_row(ws, total_row, 22)

    ws.merge_cells(f"A{total_row+1}:H{total_row+1}")
    ws[f"A{total_row+1}"].value = "平均工賃月額 = 月間工賃支払総額 ÷ 当月の利用者数　　【目標：月額3,000円以上（努力義務）】"
    ws[f"A{total_row+1}"].font = Font(name="Meiryo UI", size=9, color="DC2626", bold=True)
    ws[f"A{total_row+1}"].alignment = Alignment(horizontal="left", vertical="center")
    set_row(ws, total_row+1, 20)

    save(wb, "03_就労支援記録様式", "07_工賃台帳（B型）.xlsx")


def make_seisan():
    """生産活動収支記録（A型・B型共通）"""
    wb = Workbook()
    ws = wb.active
    ws.title = "生産活動収支記録"

    ws.merge_cells("A1:G1")
    ws["A1"].value = "生産活動収支記録（就労継続支援A型・B型用）　　　　年度"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    set_row(ws, 1, 26)

    ws.merge_cells("A2:G2")
    ws["A2"].value = "【A型のポイント】生産活動収入 ≧ 利用者賃金総額 → スコア「生産活動」分野での高評価につながります"
    ws["A2"].font = Font(name="Meiryo UI", size=8, color="DC2626", bold=True)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center")
    set_row(ws, 2, 20)

    cols = [("A","月",8),("B","生産活動\n収入（円）",14),("C","材料費・\n経費（円）",14),
            ("D","収支差額\n（B-C）",14),("E","利用者賃金\n総額（円）※A型",16),
            ("F","収入≧賃金\n（A型判定）",12),("G","備考（作業内容・特記）",28)]
    header_row(ws, 3, cols)

    months = ["4月","5月","6月","7月","8月","9月","10月","11月","12月","1月","2月","3月"]
    for i, m in enumerate(months, start=4):
        ws[f"A{i}"].value = m
        ws[f"A{i}"].font = body_font(bold=True); ws[f"A{i}"].fill = sub_fill()
        ws[f"A{i}"].alignment = Alignment(horizontal="center", vertical="center"); ws[f"A{i}"].border = thin_border()
        for col in "BCDEFG":
            body_cell(ws, f"{col}{i}", halign="right" if col in "BCDEF" else "left")
        set_row(ws, i, 22)

    total = 16
    ws[f"A{total}"].value = "年度合計"
    ws[f"A{total}"].font = body_font(bold=True); ws[f"A{total}"].fill = sub_fill()
    ws[f"A{total}"].alignment = Alignment(horizontal="center", vertical="center"); ws[f"A{total}"].border = thin_border()
    for col in "BCDEFG":
        body_cell(ws, f"{col}{total}")
    set_row(ws, total, 22)

    save(wb, "03_就労支援記録様式", "08_生産活動収支記録.xlsx")


# ======================================================================
# 07_運営規程ひな形（Word）
# ======================================================================

def make_uneikitei_houmon():
    """運営規程ひな形（訪問看護）"""
    doc = Document()
    doc.styles["Normal"].font.name = "Meiryo UI"
    doc.styles["Normal"].font.size = Pt(10)

    style = doc.styles["Normal"]
    section = doc.sections[0]
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3); section.right_margin = Cm(2.5)

    h = doc.add_heading("訪問看護ステーション 運営規程（ひな形）", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x0E, 0x74, 0x90)
    h.runs[0].font.size = Pt(14)

    doc.add_paragraph("※ このひな形は参考例です。実際の申請では事業所の実態に合わせて修正し、京都府・京都市の最新様式を使用してください。")

    chapters = [
        ("第1条（目的）", "この運営規程は、○○訪問看護ステーション（以下「当ステーション」という。）が実施する訪問看護事業の適正な運営を確保するために必要な人員及び管理運営に関する事項を定め、利用者が心身の機能の維持回復を図り、その居宅において自立した日常生活を営むことができるよう、適正な訪問看護の提供を目的とする。"),
        ("第2条（運営の方針）", "当ステーションが実施する訪問看護事業の運営は、以下の方針により行う。\n①利用者の意思及び人格を尊重し、常にその者の立場に立ったサービス提供に努める。\n②地域の医療機関・保健・福祉機関との連携を密にして、総合的なサービス提供に努める。\n③利用者の人権の擁護、虐待の防止のため、研修の実施等、必要な措置を講じる。\n④個人情報の保護に努め、業務上知り得た情報を適切に管理する。"),
        ("第3条（事業所の名称及び所在地）", "事業所の名称：○○訪問看護ステーション\n事業所の所在地：京都府○○市○○町○○番地"),
        ("第4条（従業者の職種、員数及び職務内容）", "当ステーションに勤務する職員の職種、員数及び職務内容は次のとおりとする。\n（１）管理者：1名　訪問看護ステーションの管理、業務の実施状況の把握、業務の改善\n（２）看護師・准看護師：常勤換算2.5名以上　訪問看護の実施、記録の作成\n（３）理学療法士・作業療法士・言語聴覚士：必要に応じて配置\n（４）事務員：必要に応じて配置"),
        ("第5条（営業日及び営業時間）", "営業日：月曜日から金曜日（祝日・年末年始〔12月29日から1月3日〕を除く）\n営業時間：午前8時30分から午後5時30分まで\n訪問看護の提供時間：上記営業時間中。ただし、緊急時は24時間対応する（24時間対応体制加算取得事業所の場合）。"),
        ("第6条（訪問看護の内容及び利用料等）", "（１）提供するサービスの内容\n・病状の観察及び判断\n・医療処置の補助\n・リハビリテーション（理学療法士・作業療法士等による場合）\n・医療機器の管理・操作指導\n・認知症のケア・精神科訪問看護\n・ターミナルケア\n\n（２）利用料\n主治医の指示に基づく訪問看護は、医療保険（健康保険・国民健康保険）の給付対象となる。利用者負担割合は保険の種類によって異なる（原則：3割または1割）。\n自立支援医療（精神通院）の対象者は、原則1割（月額上限あり）となる。"),
        ("第7条（苦情処理）", "苦情受付窓口：管理者\n受付方法：面接、電話、書面\n苦情処理の手順：受付→記録→調査→対応→報告→記録の保管\n第三者委員：○○様（連絡先：○○）"),
        ("第8条（緊急時等における対応方法）", "訪問中に利用者の病状が急変した場合は、速やかに主治医に連絡をとり、必要に応じて救急搬送の手配を行う。管理者に報告し、利用者家族への連絡も速やかに行う。"),
        ("第9条（虐待の防止のための措置）", "当ステーションは利用者への虐待防止のため、以下の措置を講ずる。\n①虐待防止責任者の設置（管理者）\n②虐待防止に関する研修の年1回以上の実施\n③苦情・通報窓口の整備・周知\n④虐待発見時の速やかな通報（市区町村への通報義務）"),
        ("第10条（秘密保持等）", "従業者は業務上知り得た利用者及びその家族の個人情報の保持に努め、正当な理由なくして第三者に漏洩しない。また、業務を離れた後もこの義務は継続する。"),
        ("附則", "この運営規程は、令和○年○月○日から施行する。"),
    ]

    for title, content in chapters:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.paragraph_format.space_before = Pt(12)
        doc.add_paragraph(content)

    path = os.path.join(BASE, "07_運営規程ひな形", "01_運営規程ひな形_訪問看護.docx")
    doc.save(path)
    print(f"✅ 07_運営規程ひな形/01_運営規程ひな形_訪問看護.docx")


def make_uneikitei_shuro():
    """運営規程ひな形（就労継続支援B型）"""
    doc = Document()
    doc.styles["Normal"].font.name = "Meiryo UI"
    doc.styles["Normal"].font.size = Pt(10)
    section = doc.sections[0]
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3); section.right_margin = Cm(2.5)

    h = doc.add_heading("就労継続支援B型 運営規程（ひな形）", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x0E, 0x74, 0x90)
    h.runs[0].font.size = Pt(14)

    doc.add_paragraph("※ このひな形は参考例です。A型の場合は「雇用契約を締結する」旨・就業規則の参照・最低賃金の記載が必要です。")

    chapters = [
        ("第1条（目的）", "この運営規程は、○○就労継続支援事業所（以下「当事業所」という。）が実施する就労継続支援B型事業の運営の適正を確保するために必要な事項を定め、障碍者が自立した日常生活又は社会生活を営めるよう支援することを目的とする。"),
        ("第2条（運営の方針）", "当事業所は、利用者の意思及び人格を尊重し、自立した日常生活・社会生活の支援を目的として、以下の方針により運営する。\n①利用者の意向を尊重した個別支援計画を作成し、適切なサービスを提供する。\n②関係機関・医療機関と連携し、総合的な支援に努める。\n③利用者の人権擁護・虐待防止のために必要な措置を講じる。\n④個人情報を適切に管理する。"),
        ("第3条（事業所の名称及び所在地）", "事業所の名称：○○就労継続支援事業所\n事業所の所在地：京都府○○市○○町○○番地"),
        ("第4条（従業者の職種、員数及び職務内容）", "（１）管理者：1名（常勤・専従）事業所全体の管理\n（２）サービス管理責任者：1名以上（利用定員60名以下）個別支援計画の作成・管理・モニタリング\n（３）職業指導員：常勤換算○名　作業・就労に関する指導支援\n（４）生活支援員：常勤換算○名　日常生活に関する相談支援\n※職業指導員と生活支援員の合計：利用者数÷10以上（常勤換算）"),
        ("第5条（利用定員）", "利用定員：○名"),
        ("第6条（営業日及び営業時間）", "営業日：月曜日から金曜日（祝日・年末年始を除く）\n営業時間：午前9時00分から午後4時00分まで"),
        ("第7条（サービスの内容及び利用料）", "（１）提供するサービス：生産活動（作業）・職業指導・生活支援・各種相談支援\n（２）生産活動の種類：○○（具体的な作業内容を記載）\n（３）利用者負担：障害者総合支援法に基づき算定。原則1割（所得に応じた月額上限あり）\n（４）工賃の支払い：生産活動の実績に応じて月1回支払う（目標：月平均3,000円以上）"),
        ("第8条（個別支援計画の作成）", "サービス管理責任者は、利用者ごとに個別支援計画を作成する。計画の内容を利用者及び家族に説明し、同意を得た上でサービスを提供する。計画は6ヶ月ごとに見直しを行う。"),
        ("第9条（苦情処理）", "苦情受付窓口：管理者（TEL：○○）\n苦情処理の手順：受付→記録→調査・対応→報告→記録保管\n第三者委員：○○様"),
        ("第10条（虐待の防止）", "虐待防止責任者を管理者と定め、年1回以上の虐待防止研修を実施する。虐待を発見した場合は速やかに市区町村に通報する。"),
        ("第11条（秘密保持）", "従業者は利用者及び家族の個人情報を正当な理由なく第三者に漏洩しない。退職後も同様とする。"),
        ("附則", "この運営規程は、令和○年○月○日から施行する。"),
    ]

    for title, content in chapters:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.paragraph_format.space_before = Pt(12)
        doc.add_paragraph(content)

    path = os.path.join(BASE, "07_運営規程ひな形", "02_運営規程ひな形_就労継続支援B型.docx")
    doc.save(path)
    print(f"✅ 07_運営規程ひな形/02_運営規程ひな形_就労継続支援B型.docx")


def make_juyo_jiko():
    """重要事項説明書ひな形"""
    doc = Document()
    doc.styles["Normal"].font.name = "Meiryo UI"
    doc.styles["Normal"].font.size = Pt(10)
    section = doc.sections[0]
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3); section.right_margin = Cm(2.5)

    h = doc.add_heading("重要事項説明書（訪問看護・就労継続支援 共通ひな形）", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x0E, 0x74, 0x90)
    h.runs[0].font.size = Pt(13)

    doc.add_paragraph("この書面は、障害者の日常生活及び社会生活を総合的に支援するための法律（障害者総合支援法）及び関係法令に基づき、サービスの利用に関する重要な事項についてご説明するものです。\n※ A型・B型・訪問看護など、サービス種別に応じて内容を修正してください。")

    sections_word = [
        ("１．事業者の概要", "名称：\n指定番号：\n所在地：\n電話：　　　　　　FAX：\nメールアドレス："),
        ("２．提供するサービスの種類・内容", "サービス種別：\n主な内容："),
        ("３．利用料金", "①障害福祉サービス費に係る利用者負担額\n原則として費用の1割。ただし、所得状況により月額上限があります。\n\n②食費等の実費負担\n食費：1食　　　円（月額　　　円）\n\n③その他加算\n加算の種類："),
        ("４．営業日・営業時間・サービス提供地域", "営業日：\n営業時間：\nサービス提供地域："),
        ("５．緊急時の対応方法", "サービス提供中に体調急変等が生じた場合は、速やかに管理者・主治医・家族へ連絡します。救急搬送が必要と判断した場合は119番に通報します。"),
        ("６．非常災害対策", "地震・火災等の非常時は、「防災マニュアル」に従い利用者の安全確保を最優先とします。避難場所：　　　　（市区町村指定の避難場所）"),
        ("７．苦情の受付・解決", "苦情受付窓口：\n受付方法：面接・電話・書面\n\n京都府・京都市に設置された苦情解決機関への相談も可能です。"),
        ("８．事故発生時の対応", "サービスの提供中に事故が発生した場合は、利用者の安全確保を最優先とし、速やかに管理者・主治医・家族へ連絡します。行政への報告が必要な場合は速やかに報告します。"),
        ("９．虐待防止・身体拘束の禁止", "当事業所は利用者への虐待・身体拘束を絶対に行いません。虐待を発見した場合は速やかに通報します。"),
        ("１０．個人情報の取り扱い", "利用者の個人情報は、サービス提供・請求・関係機関との連携の目的のみに使用します。第三者への提供は同意書に基づきます。"),
        ("１１．サービスの中止・解約", "30日前の書面による申し出で解約できます。事業者側からの解約は正当な理由がある場合に限り、30日前に書面で通知します。"),
    ]

    for title, content in sections_word:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True; r.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(10)
        doc.add_paragraph(content)

    doc.add_paragraph("\n\n【説明者（事業所）署名欄】\n\n説明者氏名（署名）：　　　　　　　　　　　　　日付：　　年　　月　　日")
    doc.add_paragraph("\n【説明を受けた方（利用者・家族）署名欄】\n\n上記の重要事項について説明を受け、内容を理解しました。\n\n氏名（署名）：　　　　　　　　　　　（本人 / 家族）　　日付：　　年　　月　　日\n\n続柄：")

    path = os.path.join(BASE, "07_運営規程ひな形", "03_重要事項説明書ひな形.docx")
    doc.save(path)
    print(f"✅ 07_運営規程ひな形/03_重要事項説明書ひな形.docx")


# ======================================================================
# 実行
# ======================================================================
if __name__ == "__main__":
    print("=== 訪問看護・障碍者福祉 テンプレート生成 ===")
    print("\n--- 02_訪問看護記録様式 ---")
    make_houmon_kiroku2()
    make_houmon_houkoku()
    make_jisseki()

    print("\n--- 06_事故苦情ヒヤリ様式 ---")
    make_jiko()
    make_hiyari()
    make_kujo()
    make_kinkyu()
    make_staff_meibo()
    make_kenshu()

    print("\n--- 03_就労支援記録様式 ---")
    make_kojin_shien()
    make_kochin()
    make_seisan()

    print("\n--- 07_運営規程ひな形 ---")
    make_uneikitei_houmon()
    make_uneikitei_shuro()
    make_juyo_jiko()

    print("\n=== 完了 ===")
