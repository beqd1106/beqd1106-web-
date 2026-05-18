"""追加テンプレートバッチ4 — 指定申請・処遇改善・不足様式"""
import os, sys
sys.stdout.reconfigure(encoding='utf-8')
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
TEAL="FF0E7490"; WHITE="FFFFFFFF"; LIGHT="FFCFFAFE"; GRAY="FFF8FAFC"; RED="FFDC2626"

def hf(): return PatternFill("solid", fgColor=TEAL)
def sf(): return PatternFill("solid", fgColor=LIGHT)
def gf(): return PatternFill("solid", fgColor=GRAY)
def hfont(sz=10, col=WHITE): return Font(name="Meiryo UI", size=sz, bold=True, color=col)
def bfont(sz=9, bold=False): return Font(name="Meiryo UI", size=sz, bold=bold)
def bd():
    s = Side(style="thin")
    return Border(left=s, right=s, top=s, bottom=s)
def sw(ws, col, w): ws.column_dimensions[col].width = w
def sr(ws, r, h): ws.row_dimensions[r].height = h

def hrow(ws, row, cols):
    for letter, text, width in cols:
        c = ws[f"{letter}{row}"]
        c.value = text; c.font = hfont(); c.fill = hf()
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border = bd(); ws.column_dimensions[letter].width = width
    sr(ws, row, 30)

def bc(ws, ref, val="", bold=False, ha="left"):
    c = ws[ref]; c.value = val; c.font = bfont(bold=bold)
    c.alignment = Alignment(horizontal=ha, vertical="center", wrap_text=True); c.border = bd()
    return c

def ec(ws, ref, val=""):
    c = ws[ref]; c.value = val
    c.font = Font(name="Meiryo UI", size=9, color="FF6B7280", italic=True)
    c.fill = gf(); c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    c.border = bd()

def title1(ws, text, cols="A1:H1", h=26):
    row = 1
    ws.merge_cells(cols)
    ws[cols.split(":")[0]].value = text
    ws[cols.split(":")[0]].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws[cols.split(":")[0]].alignment = Alignment(horizontal="center", vertical="center")
    sr(ws, row, h)

def note(ws, ref, text, row_h=20):
    row = int(''.join(filter(str.isdigit, ref)))
    end_col = ref[0]
    ws.merge_cells(f"A{row}:{end_col}{row}")
    ws[f"A{row}"].value = text
    ws[f"A{row}"].font = Font(name="Meiryo UI", size=8, color=RED, bold=True)
    ws[f"A{row}"].alignment = Alignment(horizontal="left", vertical="center")
    sr(ws, row, row_h)

def save(wb, folder, fname):
    p = os.path.join(BASE, folder, fname); wb.save(p)
    print(f"OK: {folder}/{fname}")

def save_doc(doc, folder, fname):
    p = os.path.join(BASE, folder, fname); doc.save(p)
    print(f"OK: {folder}/{fname}")


# ======================================================================
# 01_指定申請様式
# ======================================================================

def make_jizen_soudan():
    """事前相談票（京都府）"""
    wb = Workbook(); ws = wb.active; ws.title = "事前相談票"
    ws.merge_cells("A1:F1")
    ws["A1"].value = "指定障害福祉サービス事業者　事前相談票（京都府 参考様式）"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center"); sr(ws, 1, 28)
    ws.merge_cells("A2:F2")
    ws["A2"].value = "※ 本様式は参考です。実際の申請前に京都府障害者支援課（TEL:075-414-4600）または京都市障害保健福祉推進室（TEL:075-222-4161）へご確認ください。"
    ws["A2"].font = Font(name="Meiryo UI", size=8, color=RED, bold=True)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, 2, 22)

    rows = [
        ("法人名称", ""), ("法人種別", "社会福祉法人 / NPO法人 / 株式会社 / 合同会社 / その他"),
        ("法人所在地", "〒"), ("代表者氏名", ""), ("電話番号", ""), ("FAX", ""),
        ("担当者氏名", ""), ("担当者連絡先（携帯等）", ""),
        ("申請予定のサービス種別", "訪問看護 / 就労継続支援A型 / 就労継続支援B型 / 相談支援 / その他"),
        ("予定する事業所の所在地", "〒"), ("予定開設時期", "　　年　　月　　日"),
        ("利用定員（予定）", "　　名"), ("現在の法人設立状況", "設立済み（　　年　　月設立） / 設立予定"),
        ("他に運営している事業所", "なし / あり（事業所名：　　　　　　）"),
        ("相談したい主な内容", "人員基準 / 設備基準 / 運営基準 / 指定申請書類 / その他（　　　）"),
    ]
    for i, (label, val) in enumerate(rows, start=3):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = label
        ws.merge_cells(f"C{i}:F{i}"); ws[f"C{i}"].value = val
        ws[f"A{i}"].font = bfont(bold=True); ws[f"A{i}"].fill = sf()
        ws[f"A{i}"].alignment = Alignment(horizontal="right", vertical="center"); ws[f"A{i}"].border = bd()
        ws[f"C{i}"].font = bfont(); ws[f"C{i}"].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        ws[f"C{i}"].border = bd(); sr(ws, i, 22)

    row = len(rows) + 3
    ws.merge_cells(f"A{row}:F{row}"); ws[f"A{row}"].value = "【事前相談希望日時】（第1希望〜第3希望）"
    ws[f"A{row}"].font = hfont(sz=9); ws[f"A{row}"].fill = hf()
    ws[f"A{row}"].alignment = Alignment(horizontal="left", vertical="center"); ws[f"A{row}"].border = bd(); sr(ws, row, 20)
    for j in range(3):
        r = row + j + 1
        ws.merge_cells(f"A{r}:F{r}"); ws[f"A{r}"].value = f"第{j+1}希望：　　年　　月　　日（　）　　時〜　　時"
        ws[f"A{r}"].font = bfont(); ws[f"A{r}"].alignment = Alignment(horizontal="left", vertical="center")
        ws[f"A{r}"].border = bd(); sr(ws, r, 22)

    for col, w in zip("ABCDEF", [16, 14, 16, 14, 14, 14]):
        ws.column_dimensions[col].width = w
    save(wb, "01_指定申請様式", "01_事前相談票（京都府）.xlsx")


def make_houjin_gaiyo():
    """申請法人概要"""
    wb = Workbook(); ws = wb.active; ws.title = "申請法人概要"
    ws.merge_cells("A1:G1")
    ws["A1"].value = "申請法人概要書（参考様式）"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center"); sr(ws, 1, 28)
    ws.merge_cells("A2:G2")
    ws["A2"].value = "※ 実際の申請では京都府・京都市の最新様式を使用してください。本様式は内容整理用の参考です。"
    ws["A2"].font = Font(name="Meiryo UI", size=8, color=RED, bold=True)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, 2, 20)

    sections = [
        ("【1. 法人基本情報】", [
            ("法人名称（正式）", ""), ("法人名称（フリガナ）", ""),
            ("法人種別", ""), ("設立年月日", "　　年　　月　　日"),
            ("法人番号（13桁）", ""), ("主たる事務所の所在地", "〒"),
            ("代表者役職・氏名", ""), ("電話番号", ""), ("FAX番号", ""),
            ("法人メールアドレス", ""),
        ]),
        ("【2. 申請する事業所情報】", [
            ("事業所の名称", ""), ("事業所の所在地", "〒"),
            ("申請するサービス種別", ""), ("申請する指定番号（更新の場合）", "新規 / 更新（　　　　　）"),
            ("利用定員", "　　名"), ("開設（予定）年月日", "　　年　　月　　日"),
            ("管理者氏名・資格", ""), ("サービス管理責任者氏名・資格", ""),
        ]),
        ("【3. 関係法令の遵守確認】", [
            ("欠格事由の有無", "なし（欠格事由に該当しないことを確認済み）"),
            ("社会保険・労働保険の加入", "加入済み / 加入予定（　　年　　月）"),
            ("就業規則の整備", "整備済み / 整備中"),
        ]),
    ]

    row = 3
    for sec_title, items in sections:
        ws.merge_cells(f"A{row}:G{row}"); ws[f"A{row}"].value = sec_title
        ws[f"A{row}"].font = hfont(sz=9); ws[f"A{row}"].fill = hf()
        ws[f"A{row}"].alignment = Alignment(horizontal="left", vertical="center"); ws[f"A{row}"].border = bd(); sr(ws, row, 20)
        row += 1
        for label, val in items:
            ws.merge_cells(f"A{row}:C{row}"); ws[f"A{row}"].value = label
            ws.merge_cells(f"D{row}:G{row}"); ws[f"D{row}"].value = val
            ws[f"A{row}"].font = bfont(bold=True); ws[f"A{row}"].fill = sf()
            ws[f"A{row}"].alignment = Alignment(horizontal="right", vertical="center"); ws[f"A{row}"].border = bd()
            ws[f"D{row}"].font = bfont(); ws[f"D{row}"].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[f"D{row}"].border = bd(); sr(ws, row, 22); row += 1

    ws.merge_cells(f"A{row}:G{row}")
    ws[f"A{row}"].value = "申請日：　　年　　月　　日　　代表者署名：　　　　　　　　　　（印）"
    ws[f"A{row}"].font = bfont(); ws[f"A{row}"].alignment = Alignment(horizontal="left", vertical="center")
    ws[f"A{row}"].border = bd(); sr(ws, row, 28)

    for col, w in zip("ABCDEFG", [16, 12, 12, 14, 12, 12, 12]):
        ws.column_dimensions[col].width = w
    save(wb, "01_指定申請様式", "02_申請法人概要（京都府）.xlsx")


def make_juugyousha_ichiran():
    """従業者一覧"""
    wb = Workbook(); ws = wb.active; ws.title = "従業者一覧"
    ws.merge_cells("A1:J1")
    ws["A1"].value = "従業者の勤務体制及び勤務形態一覧表（参考様式）"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center"); sr(ws, 1, 28)
    ws.merge_cells("A2:J2")
    ws["A2"].value = "事業所名：　　　　　　　　　　　　サービス種別：　　　　　　　　　　作成日：　　年　　月　　日"
    ws["A2"].font = bfont(); ws["A2"].alignment = Alignment(horizontal="left", vertical="center"); ws["A2"].border = bd(); sr(ws, 2, 22)

    cols = [
        ("A","氏名",14), ("B","職種\n（役職）",12), ("C","雇用形態\n（常勤/非常勤）",10),
        ("D","所有資格・免許\n（登録番号）",20), ("E","実務経験年数",8),
        ("F","週所定\n労働時間",8), ("G","常勤換算\n（h）",8),
        ("H","業務内容\n（担当サービス）",18), ("I","兼務状況",12), ("J","備考",12)
    ]
    hrow(ws, 3, cols)

    ex_rows = [
        ["山田 花子", "管理者\n看護師", "常勤", "看護師免許（第123456号）", "15年", "40", "40.0", "管理・訪問看護", "管理者兼務", ""],
        ["鈴木 太郎", "看護師", "常勤", "看護師免許（第234567号）", "8年", "40", "40.0", "訪問看護", "なし", ""],
        ["田中 美穂", "准看護師", "非常勤", "准看護師免許（第345678号）", "5年", "20", "20.0", "訪問看護", "なし", "週3日勤務"],
    ]
    for i, row_data in enumerate(ex_rows, start=4):
        for j, val in enumerate(row_data):
            ec(ws, f"{get_column_letter(j+1)}{i}", val)
        sr(ws, i, 32)

    for row in range(7, 22):
        for col in range(1, 11):
            bc(ws, f"{get_column_letter(col)}{row}")
        sr(ws, row, 28)

    ws.merge_cells(f"A22:J22")
    ws["A22"].value = "【常勤換算計算】常勤の所定労働時間（週：　　h）/ 合計常勤換算数：　　人　（訪問看護の基準：看護師等2.5人以上）"
    ws["A22"].font = Font(name="Meiryo UI", size=9, color=RED, bold=True)
    ws["A22"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, 22, 22)
    save(wb, "01_指定申請様式", "03_従業者一覧（京都府）.xlsx")


def make_shikingakeikaku():
    """資金計画"""
    wb = Workbook(); ws = wb.active; ws.title = "資金計画"
    ws.merge_cells("A1:F1")
    ws["A1"].value = "資金計画書（参考様式）"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center"); sr(ws, 1, 28)
    ws.merge_cells("A2:F2")
    ws["A2"].value = "※ 本様式は参考です。実際の申請様式は各行政窓口で確認してください。金額は概算で記入し、添付書類（通帳コピー等）と合わせて提出します。"
    ws["A2"].font = Font(name="Meiryo UI", size=8, color=RED, bold=True)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, 2, 22)

    ws.merge_cells("A3:F3"); ws["A3"].value = "【収入計画（初年度）】"
    ws["A3"].font = hfont(sz=9); ws["A3"].fill = hf(); ws["A3"].alignment = Alignment(horizontal="left", vertical="center"); ws["A3"].border = bd(); sr(ws, 3, 20)

    income_cols = [("A","収入項目",24),("B","月額見込（円）",14),("C","年間見込（円）",14),("D","算出根拠・備考",28)]
    hrow(ws, 4, income_cols)
    income_items = [
        ["障害福祉サービス費（介護給付費）", "", "", "利用者数×単価×利用日数で算出"],
        ["自立支援医療費（訪問看護の場合）", "", "", ""],
        ["利用者負担金（1割）", "", "", ""],
        ["その他収入（生産活動収入等）", "", "", "就労系のみ"],
        ["合計収入", "", "", ""],
    ]
    for i, row in enumerate(income_items, start=5):
        for j, v in enumerate(row):
            bc(ws, f"{get_column_letter(j+1)}{i}", v, bold=(i == 9))
        sr(ws, i, 22)

    ws.merge_cells("A10:F10"); ws["A10"].value = "【支出計画（初年度）】"
    ws["A10"].font = hfont(sz=9); ws["A10"].fill = hf(); ws["A10"].alignment = Alignment(horizontal="left", vertical="center"); ws["A10"].border = bd(); sr(ws, 10, 20)

    exp_cols = [("A","支出項目",24),("B","月額見込（円）",14),("C","年間見込（円）",14),("D","算出根拠・備考",28)]
    hrow(ws, 11, exp_cols)
    exp_items = [
        ["人件費（給与・賞与・社会保険料等）", "", "", ""],
        ["賃料・リース料", "", "", ""],
        ["水道光熱費", "", "", ""],
        ["通信・交通費", "", "", ""],
        ["消耗品・備品費", "", "", ""],
        ["研修・採用費", "", "", ""],
        ["その他経費", "", "", ""],
        ["合計支出", "", "", ""],
    ]
    for i, row in enumerate(exp_items, start=12):
        for j, v in enumerate(row):
            bc(ws, f"{get_column_letter(j+1)}{i}", v, bold=(i == 19))
        sr(ws, i, 22)

    ws.merge_cells("A20:D20"); ws["A20"].value = "【収支差額（初年度見込）】　収入合計 - 支出合計 ＝　　　　　円"
    ws["A20"].font = Font(name="Meiryo UI", size=10, bold=True, color=TEAL)
    ws["A20"].alignment = Alignment(horizontal="left", vertical="center"); ws["A20"].border = bd(); sr(ws, 20, 24)

    ws.merge_cells("A21:D21"); ws["A21"].value = "【手元資金（開設時）】　運転資金として概ね3ヶ月分の支出額相当を確保することが推奨されます。　保有額：　　　　　円"
    ws["A21"].font = bfont(); ws["A21"].alignment = Alignment(horizontal="left", vertical="center"); ws["A21"].border = bd(); sr(ws, 21, 24)

    for col, w in zip("ABCD", [24, 14, 14, 28]): ws.column_dimensions[col].width = w
    save(wb, "01_指定申請様式", "04_資金計画（京都府）.xlsx")


# ======================================================================
# 02_訪問看護記録様式（不足分）
# ======================================================================

def make_houmon_keikaku():
    """訪問看護計画書（別紙様式1準拠）"""
    wb = Workbook(); ws = wb.active; ws.title = "訪問看護計画書"
    ws.merge_cells("A1:H1")
    ws["A1"].value = "訪問看護計画書（別紙様式1準拠）"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center"); sr(ws, 1, 28)
    ws.merge_cells("A2:H2")
    ws["A2"].value = "※ 保険医療機関の主治医の訪問看護指示書に基づき作成。利用者・家族に交付し、同意を得ること（健保・医保共通）。"
    ws["A2"].font = Font(name="Meiryo UI", size=8, color=TEAL, bold=True)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, 2, 20)

    info = [
        ("利用者氏名", "（様）", "生年月日", "　　年　月　日（　歳）"),
        ("住所", "〒", "電話", ""),
        ("主傷病名", "", "副傷病名", ""),
        ("主治医・医療機関名", "", "医師TEL", ""),
        ("計画作成日", "　　年　月　日", "計画期間", "　　年　月〜　　年　月"),
        ("担当看護師", "", "管理者確認", ""),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info, start=3):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value = l2
        ws.merge_cells(f"G{i}:H{i}"); ws[f"G{i}"].value = v2
        for cell in [f"A{i}", f"E{i}"]:
            ws[cell].font = bfont(bold=True); ws[cell].fill = sf()
            ws[cell].alignment = Alignment(horizontal="right", vertical="center"); ws[cell].border = bd()
        for cell in [f"C{i}", f"G{i}"]:
            ws[cell].font = bfont(); ws[cell].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[cell].border = bd()
        sr(ws, i, 22)

    secs = [
        (9, "【訪問頻度・時間】", 30, "例）週　回　/　回あたり　分〜　分　/ 主な訪問曜日："),
        (11, "【問題点・看護の目標】", 80,
         "①（問題）：　　　　　（目標）：\n②（問題）：　　　　　（目標）：\n③（問題）：　　　　　（目標）："),
        (13, "【訪問看護の内容（予定するケアの概要）】", 100,
         "①バイタルサイン測定・全身状態の観察\n②医療処置の補助（処置内容：　　）\n③服薬管理・服薬指導\n④清潔ケア（清拭・入浴介助等）\n⑤家族への指導・支援\n⑥精神的支援・相談"),
        (15, "【留意事項・緊急時対応】", 50,
         "緊急時連絡先：主治医（　　　　）→家族（　　　　）→管理者（　　　　）\n入院先："),
        (17, "【本人・家族の同意】", 30,
         "上記計画について説明を受け、同意します。\n\n署名：　　　　　　　（続柄：　　）　　日付：　　年　月　日"),
    ]
    for rn, title, h, ex in secs:
        ws.merge_cells(f"A{rn}:H{rn}"); ws[f"A{rn}"].value = title
        ws[f"A{rn}"].font = hfont(sz=9); ws[f"A{rn}"].fill = hf()
        ws[f"A{rn}"].alignment = Alignment(horizontal="left", vertical="center"); ws[f"A{rn}"].border = bd(); sr(ws, rn, 20)
        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c = ws[f"A{rn+1}"]
        c.value = ex; c.font = Font(name="Meiryo UI", size=9, color="FF6B7280", italic=True)
        c.fill = gf(); c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True); c.border = bd(); sr(ws, rn+1, h)

    for col, w in zip("ABCDEFGH", [14]*8): ws.column_dimensions[col].width = w
    save(wb, "02_訪問看護記録様式", "01_訪問看護計画書（別紙様式1準拠）.xlsx")


def make_shiji_sho():
    """訪問看護指示書（参考様式）"""
    wb = Workbook(); ws = wb.active; ws.title = "訪問看護指示書"
    ws.merge_cells("A1:H1")
    ws["A1"].value = "訪問看護指示書（参考様式）"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center"); sr(ws, 1, 28)
    ws.merge_cells("A2:H2")
    ws["A2"].value = "※ 正式な訪問看護指示書は主治医（保険医）が記載・交付します。本様式は管理用の参考様式です。実際の指示書は主治医から受領してください。"
    ws["A2"].font = Font(name="Meiryo UI", size=8, color=RED, bold=True)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, 2, 22)

    info = [
        ("患者氏名", "（様）", "生年月日", "　　年　月　日（　歳）"),
        ("主たる傷病名", "", "要介護度", "要介護　/ 要支援　/ なし"),
        ("住所・TEL", "〒", "保険種別", "医療保険 / 介護保険 / 自立支援医療"),
        ("指示期間", "　　年　月　日　〜　　年　月　日", "訪問回数（指定）", "週　回程度"),
        ("訪問看護指示区分", "訪問看護 / 精神科訪問看護 / 特別訪問看護指示", "", ""),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info, start=3):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value = l2
        ws.merge_cells(f"G{i}:H{i}"); ws[f"G{i}"].value = v2
        for cell in [f"A{i}", f"E{i}"]:
            ws[cell].font = bfont(bold=True); ws[cell].fill = sf()
            ws[cell].alignment = Alignment(horizontal="right", vertical="center"); ws[cell].border = bd()
        for cell in [f"C{i}", f"G{i}"]:
            ws[cell].font = bfont(); ws[cell].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[cell].border = bd()
        sr(ws, i, 22)

    secs = [
        (8, "【病状・治療状況】", 60, ""),
        (10, "【訪問看護の指示内容（ケアの内容）】", 80,
         "①観察：バイタル・全身状態\n②処置：（具体的内容）\n③服薬管理：\n④その他："),
        (12, "【緊急時の連絡先・対応】", 40, "緊急連絡先（医療機関）：　　　　　TEL："),
        (14, "【留意事項・注意すべき合併症等】", 50, ""),
    ]
    for rn, title, h, ex in secs:
        ws.merge_cells(f"A{rn}:H{rn}"); ws[f"A{rn}"].value = title
        ws[f"A{rn}"].font = hfont(sz=9); ws[f"A{rn}"].fill = hf()
        ws[f"A{rn}"].alignment = Alignment(horizontal="left", vertical="center"); ws[f"A{rn}"].border = bd(); sr(ws, rn, 20)
        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c = ws[f"A{rn+1}"]
        c.value = ex; c.font = Font(name="Meiryo UI", size=9, color="FF6B7280", italic=True)
        c.fill = gf() if ex else PatternFill()
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True); c.border = bd(); sr(ws, rn+1, h)

    sign = 16
    ws.merge_cells(f"A{sign}:D{sign}"); ws[f"A{sign}"].value = "指示日：　　年　月　日"
    ws.merge_cells(f"E{sign}:H{sign}"); ws[f"E{sign}"].value = "医師氏名（署名・捺印）：　　　　　　　　　　　（印）"
    for cell in [f"A{sign}", f"E{sign}"]:
        ws[cell].font = bfont(); ws[cell].border = bd()
        ws[cell].alignment = Alignment(horizontal="left", vertical="center")
    sr(ws, sign, 30)
    ws.merge_cells(f"A{sign+1}:H{sign+1}"); ws[f"A{sign+1}"].value = "医療機関名・所在地：　　　　　　　　　　　　　　　　TEL："
    ws[f"A{sign+1}"].font = bfont(); ws[f"A{sign+1}"].border = bd()
    ws[f"A{sign+1}"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, sign+1, 28)

    for col, w in zip("ABCDEFGH", [14]*8): ws.column_dimensions[col].width = w
    save(wb, "02_訪問看護記録様式", "02_訪問看護指示書（管理用）.xlsx")


# ======================================================================
# 03_就労支援記録様式（不足分）
# ======================================================================

def make_monitoring_shuro():
    """モニタリング記録（就労継続支援用）"""
    wb = Workbook(); ws = wb.active; ws.title = "モニタリング記録"
    ws.merge_cells("A1:H1")
    ws["A1"].value = "モニタリング記録（就労継続支援A型・B型用）"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center"); sr(ws, 1, 28)
    ws.merge_cells("A2:H2")
    ws["A2"].value = "【モニタリング頻度】就労継続支援は原則6ヶ月毎（状況変化時は随時）。記録を保管し、個別支援計画の見直しに活用すること。"
    ws["A2"].font = Font(name="Meiryo UI", size=8, color=TEAL, bold=True)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, 2, 20)

    info = [
        ("利用者氏名", "（様）", "障害種別", "身体 / 精神 / 知的 / 難病"),
        ("事業所名", "", "サービス種別", "就労継続支援A型 / B型"),
        ("モニタリング実施日", "　　年　月　日", "実施者（サビ管）", ""),
        ("前回モニタリング日", "　　年　月　日", "個別支援計画の期間", "　　年　月〜　　年　月"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info, start=3):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value = l2
        ws.merge_cells(f"G{i}:H{i}"); ws[f"G{i}"].value = v2
        for cell in [f"A{i}", f"E{i}"]:
            ws[cell].font = bfont(bold=True); ws[cell].fill = sf()
            ws[cell].alignment = Alignment(horizontal="right", vertical="center"); ws[cell].border = bd()
        for cell in [f"C{i}", f"G{i}"]:
            ws[cell].font = bfont(); ws[cell].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[cell].border = bd()
        sr(ws, i, 22)

    secs = [
        (7, "【目標の達成状況】（個別支援計画の目標ごとに確認）", 90,
         "目標①「　　　　　　　」\n→達成度：達成 / ほぼ達成 / 未達成　　状況：\n\n目標②「　　　　　　　」\n→達成度：達成 / ほぼ達成 / 未達成　　状況："),
        (9, "【本人の意向・満足度】（本人から聴取した内容）", 60,
         "サービスへの満足度：満足 / ほぼ満足 / 不満　　　理由：\n本人の希望・要望："),
        (11, "【就労・作業状況】", 60,
         "出席日数（当月）：　日/　日（出席率：　　%）\n主な作業内容：\n体調・コンディション："),
        (13, "【健康・生活状況の変化】", 50,
         "医療機関受診状況：\n生活環境の変化：\n服薬状況："),
        (15, "【支援内容の評価・課題】", 60, ""),
        (17, "【今後の支援方針・計画変更の要否】", 50,
         "計画変更：必要（変更内容：　　　） / 不要\n次回モニタリング予定：　　年　月　日"),
        (19, "【本人・家族の確認・同意】", 30,
         "モニタリング結果について説明を受けました。\n署名：　　　　　　　（続柄：　　）　日付：　　年　月　日"),
    ]
    for rn, title, h, ex in secs:
        ws.merge_cells(f"A{rn}:H{rn}"); ws[f"A{rn}"].value = title
        ws[f"A{rn}"].font = hfont(sz=9); ws[f"A{rn}"].fill = hf()
        ws[f"A{rn}"].alignment = Alignment(horizontal="left", vertical="center"); ws[f"A{rn}"].border = bd(); sr(ws, rn, 20)
        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c = ws[f"A{rn+1}"]
        c.value = ex; c.font = Font(name="Meiryo UI", size=9, color="FF6B7280", italic=True)
        c.fill = gf(); c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True); c.border = bd(); sr(ws, rn+1, h)

    for col, w in zip("ABCDEFGH", [14]*8): ws.column_dimensions[col].width = w
    save(wb, "03_就労支援記録様式", "03_モニタリング記録（就労）.xlsx")


# ======================================================================
# 04_相談支援様式（不足分）
# ======================================================================

def make_service_riyou_keikaku():
    """サービス等利用計画書"""
    wb = Workbook(); ws = wb.active; ws.title = "サービス等利用計画書"
    ws.merge_cells("A1:H1")
    ws["A1"].value = "サービス等利用計画書（計画相談支援用）"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center"); sr(ws, 1, 28)
    ws.merge_cells("A2:H2")
    ws["A2"].value = "【作成義務】障害者総合支援法に基づき、指定特定相談支援事業所の相談支援専門員が作成します。利用者本人の同意を得て交付・保管。"
    ws["A2"].font = Font(name="Meiryo UI", size=8, color=TEAL, bold=True)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, 2, 20)

    info = [
        ("利用者氏名", "（様）", "生年月日", "　年　月　日（　歳）"),
        ("住所", "〒", "電話番号", ""),
        ("障害の種別・等級", "身体(　級) / 精神(　級) / 知的(　) / 難病", "障害支援区分", "区分　　　"),
        ("計画作成日", "　　年　月　日", "計画期間", "　　年　月〜　　年　月"),
        ("相談支援専門員氏名", "", "事業所名", ""),
        ("利用者の総合的な生活目標\n（本人の夢・希望）", "", "", ""),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info, start=3):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value = l2
        ws.merge_cells(f"G{i}:H{i}"); ws[f"G{i}"].value = v2
        for cell in [f"A{i}", f"E{i}"]:
            ws[cell].font = bfont(bold=True); ws[cell].fill = sf()
            ws[cell].alignment = Alignment(horizontal="right", vertical="center"); ws[cell].border = bd()
        for cell in [f"C{i}", f"G{i}"]:
            ws[cell].font = bfont(); ws[cell].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[cell].border = bd()
        sr(ws, i, 22)

    row = 9
    ws.merge_cells(f"A{row}:H{row}"); ws[f"A{row}"].value = "【利用するサービス・支援の一覧】"
    ws[f"A{row}"].font = hfont(sz=9); ws[f"A{row}"].fill = hf()
    ws[f"A{row}"].alignment = Alignment(horizontal="left", vertical="center"); ws[f"A{row}"].border = bd(); sr(ws, row, 20)

    svc_cols = [("A","サービス種別",18),("B","事業所名",18),("C","提供内容・頻度",22),
                ("D","担当者",12),("E","利用開始日",10),("F","サービス費用\n（月額概算）",12),("G","備考",12)]
    hrow(ws, 10, svc_cols)
    ex_svcs = [
        ["就労継続支援B型", "○○就労継続支援事業所", "週4日・午前9時〜午後3時", "田中サビ管", "R7.4.1", "37,200円/月", ""],
        ["訪問看護", "○○訪問看護ステーション", "週2回・60分", "山田看護師", "R7.4.1", "24,000円/月", "精神科訪問"],
        ["相談支援（計画相談）", "当事業所", "モニタリング月1回以上", "当担当", "R7.4.1", "0円（給付）", ""],
    ]
    for i, row_data in enumerate(ex_svcs, start=11):
        for j, v in enumerate(row_data):
            ec(ws, f"{get_column_letter(j+1)}{i}", v)
        sr(ws, i, 28)
    for row_i in range(14, 20):
        for col in range(1, 8): bc(ws, f"{get_column_letter(col)}{row_i}")
        sr(ws, row_i, 25)

    secs2 = [
        (20, "【優先して解決すべき課題（ニーズ）】", 60, "①\n②\n③"),
        (22, "【支援の方向性・留意事項】", 50, ""),
        (24, "【緊急時の対応】", 30, "緊急連絡先：　　　　　TEL：　　　　（続柄：　）"),
        (26, "【本人・家族の同意】", 28,
         "上記計画について説明を受け、同意します。\n本人署名：　　　　　　　日付：　　年　月　日"),
    ]
    for rn, title, h, ex in secs2:
        ws.merge_cells(f"A{rn}:H{rn}"); ws[f"A{rn}"].value = title
        ws[f"A{rn}"].font = hfont(sz=9); ws[f"A{rn}"].fill = hf()
        ws[f"A{rn}"].alignment = Alignment(horizontal="left", vertical="center"); ws[f"A{rn}"].border = bd(); sr(ws, rn, 20)
        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c = ws[f"A{rn+1}"]
        c.value = ex; c.font = Font(name="Meiryo UI", size=9, color="FF6B7280", italic=True) if ex else bfont()
        c.fill = gf() if ex else PatternFill()
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True); c.border = bd(); sr(ws, rn+1, h)

    for col, w in zip("ABCDEFGH", [18,18,22,12,10,12,12,12]): ws.column_dimensions[col].width = w
    save(wb, "04_相談支援様式", "01_サービス等利用計画書.xlsx")


def make_monitoring_soudan():
    """モニタリング報告書（相談支援用）"""
    wb = Workbook(); ws = wb.active; ws.title = "モニタリング報告書"
    ws.merge_cells("A1:H1")
    ws["A1"].value = "モニタリング報告書（継続サービス利用支援用）"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center"); sr(ws, 1, 28)
    ws.merge_cells("A2:H2")
    ws["A2"].value = "【モニタリング頻度】原則6ヶ月毎（状態変化時・サービス内容変更時は随時）。実施記録は保管し、計画の見直しに活用すること。"
    ws["A2"].font = Font(name="Meiryo UI", size=8, color=TEAL, bold=True)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, 2, 20)

    info = [
        ("利用者氏名", "（様）", "モニタリング実施日", "　　年　月　日"),
        ("相談支援専門員", "", "前回モニタリング日", "　　年　月　日"),
        ("計画の対象期間", "　　年　月〜　　年　月", "参加者", "本人 / 家族 / 事業所担当者"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info, start=3):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value = l2
        ws.merge_cells(f"G{i}:H{i}"); ws[f"G{i}"].value = v2
        for cell in [f"A{i}", f"E{i}"]:
            ws[cell].font = bfont(bold=True); ws[cell].fill = sf()
            ws[cell].alignment = Alignment(horizontal="right", vertical="center"); ws[cell].border = bd()
        for cell in [f"C{i}", f"G{i}"]:
            ws[cell].font = bfont(); ws[cell].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[cell].border = bd()
        sr(ws, i, 22)

    secs = [
        (6, "【目標の達成状況】（計画上の目標ごとに達成度を確認）", 80,
         "目標①「　　　　　　　　　」　達成度：達成 / ほぼ達成 / 未達成\n→状況：\n\n目標②「　　　　　　　　　」　達成度：達成 / ほぼ達成 / 未達成\n→状況："),
        (8, "【各サービスの提供状況の確認】", 60,
         "就労継続支援：利用回数　日/月　体調・様子：\n訪問看護：利用回数　回/月　主な内容：\nその他："),
        (10, "【本人・家族の意向・満足度】", 60,
         "本人の意向：\n家族の意向：\nサービスへの満足度：満足 / ほぼ満足 / 不満　理由："),
        (12, "【生活状況の変化】", 50,
         "健康状態：\n生活環境：\n社会参加状況："),
        (14, "【支援上の課題・今後の方針】", 60,
         "課題：\n方針：\n計画変更の要否：要（変更内容：　　　） / 不要"),
        (16, "【サービス等利用計画の変更内容（変更ある場合）】", 40, ""),
        (18, "【次回モニタリング予定】", 25, "次回：　　年　月頃（変化があれば随時実施）"),
        (20, "【本人・家族の確認】", 28,
         "モニタリング結果について説明を受け、確認しました。\n署名：　　　　　　　（続柄：　　）　日付：　　年　月　日"),
    ]
    for rn, title, h, ex in secs:
        ws.merge_cells(f"A{rn}:H{rn}"); ws[f"A{rn}"].value = title
        ws[f"A{rn}"].font = hfont(sz=9); ws[f"A{rn}"].fill = hf()
        ws[f"A{rn}"].alignment = Alignment(horizontal="left", vertical="center"); ws[f"A{rn}"].border = bd(); sr(ws, rn, 20)
        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c = ws[f"A{rn+1}"]
        c.value = ex; c.font = Font(name="Meiryo UI", size=9, color="FF6B7280", italic=True) if ex else bfont()
        c.fill = gf() if ex else PatternFill()
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True); c.border = bd(); sr(ws, rn+1, h)

    for col, w in zip("ABCDEFGH", [14]*8): ws.column_dimensions[col].width = w
    save(wb, "04_相談支援様式", "02_モニタリング報告書.xlsx")


def make_assessment_soudan():
    """アセスメントシート（相談支援版）"""
    wb = Workbook(); ws = wb.active; ws.title = "アセスメント票（相談支援）"
    ws.merge_cells("A1:H1")
    ws["A1"].value = "アセスメントシート（相談支援・計画相談用）"
    ws["A1"].font = Font(name="Meiryo UI", size=13, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center"); sr(ws, 1, 28)
    ws.merge_cells("A2:H2")
    ws["A2"].value = "実施日：　　年　月　日　実施者：　　　　　　　参加者：本人 / 家族（　　）/ キーパーソン（　　）/ その他（　　）"
    ws["A2"].font = bfont(); ws["A2"].alignment = Alignment(horizontal="left", vertical="center"); ws["A2"].border = bd(); sr(ws, 2, 22)

    info = [
        ("利用者氏名", "（様）", "生年月日", "　年　月　日（　歳）"),
        ("障害の種別・等級", "", "障害支援区分", "区分　　　 / 認定日："),
        ("住所・連絡先", "〒", "保護者・キーパーソン", ""),
        ("主治医・医療機関", "", "主傷病名", ""),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info, start=3):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value = l2
        ws.merge_cells(f"G{i}:H{i}"); ws[f"G{i}"].value = v2
        for cell in [f"A{i}", f"E{i}"]:
            ws[cell].font = bfont(bold=True); ws[cell].fill = sf()
            ws[cell].alignment = Alignment(horizontal="right", vertical="center"); ws[cell].border = bd()
        for cell in [f"C{i}", f"G{i}"]:
            ws[cell].font = bfont(); ws[cell].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[cell].border = bd()
        sr(ws, i, 22)

    secs = [
        (7, "【本人の希望・将来の夢】", 60,
         "本人から直接聴取した希望（そのままの言葉で記録）：\n将来どんな生活をしたいか："),
        (9, "【現在の生活状況】", 70,
         "起床・就寝時刻、食事・服薬の自己管理、清潔管理、外出・移動の状況：\n居住形態：自宅（独居/家族と）/ グループホーム / その他"),
        (11, "【障害・健康状態】", 60,
         "障害の状態（日常生活への影響、進行性の有無等）：\n服薬状況・通院状況：\n精神状態の波・体調悪化のサイン："),
        (13, "【就労・日中活動の状況】", 60,
         "現在の活動：\n就労能力・希望：\n経済状況（収入・工賃・年金等）："),
        (15, "【家族・社会関係】", 50,
         "同居家族・キーパーソンの状況：\n地域・コミュニティとの関係："),
        (17, "【現在利用中のサービス・支援】", 50,
         "サービス種別 / 事業所名 / 利用頻度：\n①\n②\n③"),
        (19, "【相談に至った経緯・主訴】", 50, ""),
        (21, "【アセスメントから見えた課題・優先ニーズ】", 70,
         "優先課題①：\n優先課題②：\n優先課題③："),
    ]
    for rn, title, h, ex in secs:
        ws.merge_cells(f"A{rn}:H{rn}"); ws[f"A{rn}"].value = title
        ws[f"A{rn}"].font = hfont(sz=9); ws[f"A{rn}"].fill = hf()
        ws[f"A{rn}"].alignment = Alignment(horizontal="left", vertical="center"); ws[f"A{rn}"].border = bd(); sr(ws, rn, 20)
        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c = ws[f"A{rn+1}"]
        c.value = ex; c.font = Font(name="Meiryo UI", size=9, color="FF6B7280", italic=True) if ex else bfont()
        c.fill = gf() if ex else PatternFill()
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True); c.border = bd(); sr(ws, rn+1, h)

    ws.merge_cells(f"A23:H23")
    ws[f"A23"].value = "アセスメント実施者署名：　　　　　　　　　日付：　　年　月　日　確認者（管理者）：　　　　　　　　"
    ws[f"A23"].font = bfont(); ws[f"A23"].alignment = Alignment(horizontal="left", vertical="center"); ws[f"A23"].border = bd(); sr(ws, 23, 26)

    for col, w in zip("ABCDEFGH", [14]*8): ws.column_dimensions[col].width = w
    save(wb, "04_相談支援様式", "04_アセスメントシート（相談支援版）.xlsx")


# ======================================================================
# 05_処遇改善加算様式
# ======================================================================

def make_shoguu_keikaku():
    """処遇改善加算計画書"""
    wb = Workbook(); ws = wb.active; ws.title = "処遇改善計画書"
    ws.merge_cells("A1:H1")
    ws["A1"].value = "障害福祉サービス等処遇改善加算 計画書（R7年度 参考様式）"
    ws["A1"].font = Font(name="Meiryo UI", size=12, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center"); sr(ws, 1, 28)
    ws.merge_cells("A2:H2")
    ws["A2"].value = "【提出期限】毎年4月15日までに都道府県（京都府：障害者支援課 TEL:075-414-4600 / 京都市：障害保健福祉推進室 TEL:075-222-4161）へ提出。"
    ws["A2"].font = Font(name="Meiryo UI", size=8, color=RED, bold=True)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, 2, 22)

    info = [
        ("事業者名称", "", "法人番号", ""),
        ("事業所名称", "", "指定番号", ""),
        ("サービス種別", "訪問看護 / 就労継続A型 / 就労継続B型 / 相談支援", "計画年度", "令和　　年度"),
        ("管理者氏名", "", "作成日", "　　年　月　日"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info, start=3):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value = l2
        ws.merge_cells(f"G{i}:H{i}"); ws[f"G{i}"].value = v2
        for cell in [f"A{i}", f"E{i}"]:
            ws[cell].font = bfont(bold=True); ws[cell].fill = sf()
            ws[cell].alignment = Alignment(horizontal="right", vertical="center"); ws[cell].border = bd()
        for cell in [f"C{i}", f"G{i}"]:
            ws[cell].font = bfont(); ws[cell].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[cell].border = bd()
        sr(ws, i, 22)

    cols7 = [("A","職員種別",14),("B","現在の\n賃金水準（円/月）",14),("C","改善額\n（円/月）",12),
             ("D","改善後\n（円/月）",12),("E","対象職員数",8),("F","改善総額\n（年間・円）",14),("G","財源（加算区分）",16),("H","備考",12)]
    hrow(ws, 7, cols7)
    for row in range(8, 14):
        for col in range(1, 9): bc(ws, f"{get_column_letter(col)}{row}")
        sr(ws, row, 22)

    ws.merge_cells("A14:H14"); ws["A14"].value = "【キャリアパス要件の確認】"
    ws["A14"].font = hfont(sz=9); ws["A14"].fill = hf(); ws["A14"].alignment = Alignment(horizontal="left", vertical="center"); ws["A14"].border = bd(); sr(ws, 14, 20)
    kp_items = [
        ("任用要件・賃金体系の整備", "あり（就業規則等に明記済み） / なし（整備予定：　　年　月）"),
        ("研修の実施（OJT・外部研修等）", "実施済み（年　　回以上） / 計画中"),
        ("職場環境等要件（複数取組）", "①　　　②　　　③　　　（要件の内容を記入）"),
    ]
    for i, (label, val) in enumerate(kp_items, start=15):
        ws.merge_cells(f"A{i}:C{i}"); ws[f"A{i}"].value = label
        ws.merge_cells(f"D{i}:H{i}"); ws[f"D{i}"].value = val
        ws[f"A{i}"].font = bfont(bold=True); ws[f"A{i}"].fill = sf()
        ws[f"A{i}"].alignment = Alignment(horizontal="right", vertical="center"); ws[f"A{i}"].border = bd()
        ws[f"D{i}"].font = bfont(); ws[f"D{i}"].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        ws[f"D{i}"].border = bd(); sr(ws, i, 28)

    ws.merge_cells("A18:H18"); ws["A18"].value = "【処遇改善の方法（具体的な記載が必要）】"
    ws["A18"].font = hfont(sz=9); ws["A18"].fill = hf(); ws["A18"].alignment = Alignment(horizontal="left", vertical="center"); ws["A18"].border = bd(); sr(ws, 18, 20)
    ws.merge_cells("A19:H19"); ws["A19"].value = "例）全職員を対象に基本給を月額○○円引き上げる。加えて処遇改善手当として月額○○円を支給する。特定の職員（勤続3年以上）には特定処遇改善手当として月額○○円を支給する。"
    ws["A19"].font = Font(name="Meiryo UI", size=9, color="FF6B7280", italic=True)
    ws["A19"].fill = gf(); ws["A19"].alignment = Alignment(horizontal="left", vertical="top", wrap_text=True); ws["A19"].border = bd(); sr(ws, 19, 70)

    ws.merge_cells("A20:H20"); ws["A20"].value = "上記の計画書の内容が事実と相違ないことを誓約します。\n代表者署名：　　　　　　　　　　　（印）　　年　月　日"
    ws["A20"].font = bfont(); ws["A20"].alignment = Alignment(horizontal="left", vertical="center"); ws["A20"].border = bd(); sr(ws, 20, 36)

    for col, w in zip("ABCDEFGH", [14,14,12,12,8,14,16,12]): ws.column_dimensions[col].width = w
    save(wb, "05_処遇改善加算様式", "01_処遇改善計画書（R7年度）.xlsx")


def make_shoguu_jisseki():
    """処遇改善実績報告書"""
    wb = Workbook(); ws = wb.active; ws.title = "処遇改善実績報告書"
    ws.merge_cells("A1:H1")
    ws["A1"].value = "障害福祉サービス等処遇改善加算 実績報告書（R7年度 参考様式）"
    ws["A1"].font = Font(name="Meiryo UI", size=12, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center"); sr(ws, 1, 28)
    ws.merge_cells("A2:H2")
    ws["A2"].value = "【提出期限】毎年7月31日までに提出（対象期間：前年度4月〜3月の実績）。計画書の記載内容に沿って実際の賃金改善実績を報告。"
    ws["A2"].font = Font(name="Meiryo UI", size=8, color=RED, bold=True)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, 2, 22)

    info = [
        ("事業者名称", "", "事業所名", ""),
        ("報告対象年度", "令和　　年度（　　年4月〜　　年3月）", "サービス種別", ""),
        ("管理者氏名", "", "作成日", "　　年　月　日"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info, start=3):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value = l2
        ws.merge_cells(f"G{i}:H{i}"); ws[f"G{i}"].value = v2
        for cell in [f"A{i}", f"E{i}"]:
            ws[cell].font = bfont(bold=True); ws[cell].fill = sf()
            ws[cell].alignment = Alignment(horizontal="right", vertical="center"); ws[cell].border = bd()
        for cell in [f"C{i}", f"G{i}"]:
            ws[cell].font = bfont(); ws[cell].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[cell].border = bd()
        sr(ws, i, 22)

    cols6 = [("A","職員種別",14),("B","計画時の\n改善額（円/月）",14),("C","実績の\n改善額（円/月）",14),
             ("D","計画比\n（実績÷計画）",10),("E","対象職員数",8),("F","年間改善総額（円）",14),("G","備考",16)]
    hrow(ws, 6, cols6)
    for row in range(7, 13):
        for col in range(1, 8): bc(ws, f"{get_column_letter(col)}{row}")
        sr(ws, row, 22)

    ws.merge_cells("A13:G13"); ws["A13"].value = "【処遇改善加算 受取額と賃金改善額の比較】"
    ws["A13"].font = hfont(sz=9); ws["A13"].fill = hf(); ws["A13"].alignment = Alignment(horizontal="left", vertical="center"); ws["A13"].border = bd(); sr(ws, 13, 20)

    comp_items = [
        ("処遇改善加算 受取額合計（年間）", "　　　　　円"),
        ("賃金改善額合計（年間）", "　　　　　円"),
        ("差額（受取額 - 改善額）", "　　　　　円　※ 差額が生じた場合の理由："),
    ]
    for i, (label, val) in enumerate(comp_items, start=14):
        ws.merge_cells(f"A{i}:D{i}"); ws[f"A{i}"].value = label
        ws.merge_cells(f"E{i}:G{i}"); ws[f"E{i}"].value = val
        ws[f"A{i}"].font = bfont(bold=True); ws[f"A{i}"].fill = sf()
        ws[f"A{i}"].alignment = Alignment(horizontal="right", vertical="center"); ws[f"A{i}"].border = bd()
        ws[f"E{i}"].font = bfont(); ws[f"E{i}"].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        ws[f"E{i}"].border = bd(); sr(ws, i, 28)

    ws.merge_cells("A17:G17"); ws["A17"].value = "上記の実績報告の内容が事実と相違ないことを誓約します。\n代表者署名：　　　　　　　　　　　（印）　　年　月　日"
    ws["A17"].font = bfont(); ws["A17"].alignment = Alignment(horizontal="left", vertical="center"); ws["A17"].border = bd(); sr(ws, 17, 36)

    for col, w in zip("ABCDEFG", [14,14,14,10,8,14,16]): ws.column_dimensions[col].width = w
    save(wb, "05_処遇改善加算様式", "02_処遇改善実績報告書.xlsx")


def make_tokutei_shoguu():
    """特定処遇改善計画書"""
    wb = Workbook(); ws = wb.active; ws.title = "特定処遇改善計画書"
    ws.merge_cells("A1:H1")
    ws["A1"].value = "特定障害福祉サービス等処遇改善加算 計画書（参考様式）"
    ws["A1"].font = Font(name="Meiryo UI", size=12, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center"); sr(ws, 1, 28)
    ws.merge_cells("A2:H2")
    ws["A2"].value = "【対象】勤続10年以上の経験・技能のある障害福祉人材を中心に賃金を改善するための加算。処遇改善加算（Ⅰ）〜（Ⅲ）を取得している事業所が対象。"
    ws["A2"].font = Font(name="Meiryo UI", size=8, color=TEAL, bold=True)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, 2, 22)

    info = [
        ("事業者名称", "", "事業所名", ""),
        ("サービス種別", "", "計画年度", "令和　　年度"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info, start=3):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value = l2
        ws.merge_cells(f"G{i}:H{i}"); ws[f"G{i}"].value = v2
        for cell in [f"A{i}", f"E{i}"]:
            ws[cell].font = bfont(bold=True); ws[cell].fill = sf()
            ws[cell].alignment = Alignment(horizontal="right", vertical="center"); ws[cell].border = bd()
        for cell in [f"C{i}", f"G{i}"]:
            ws[cell].font = bfont(); ws[cell].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[cell].border = bd()
        sr(ws, i, 22)

    ws.merge_cells("A5:H5"); ws["A5"].value = "【グループ区分と賃金改善の考え方（3グループ設定が必要）】"
    ws["A5"].font = hfont(sz=9); ws["A5"].fill = hf(); ws["A5"].alignment = Alignment(horizontal="left", vertical="center"); ws["A5"].border = bd(); sr(ws, 5, 20)

    grp_cols = [("A","グループ",12),("B","対象者の条件",24),("C","改善方法",18),
                ("D","改善額（円/月）",14),("E","対象者数",8),("F","年間総額（円）",14),("G","備考",16)]
    hrow(ws, 6, grp_cols)
    grps = [
        ["グループ①\n（経験・技能あり）", "経験10年以上・または管理的な役割を担う\n介護福祉士・精神保健福祉士等の有資格者", "基本給引き上げ", "", "", "", "月額8万円以上の改善が原則"],
        ["グループ②\n（その他職員）", "グループ①以外の直接処遇職員", "処遇改善手当", "", "", "", "グループ①の½以上の改善"],
        ["グループ③\n（経営・事務等）", "事務・経営・管理部門等", "必要に応じて", "", "", "", "改善額はグループ①の½未満"],
    ]
    for i, row_data in enumerate(grps, start=7):
        for j, v in enumerate(row_data):
            bc(ws, f"{get_column_letter(j+1)}{i}", v)
        sr(ws, i, 45)

    secs = [
        (10, "【賃上げ効果が継続される仕組みの整備】", 50,
         "就業規則・賃金規程への明記（改定日：　　年　月　日）\n具体的な記載箇所・条文番号："),
        (12, "【配分方法の労働者への周知方法】", 40,
         "周知方法：朝礼・書面配布・掲示　実施日：　　年　月　日"),
    ]
    for rn, title, h, ex in secs:
        ws.merge_cells(f"A{rn}:H{rn}"); ws[f"A{rn}"].value = title
        ws[f"A{rn}"].font = hfont(sz=9); ws[f"A{rn}"].fill = hf()
        ws[f"A{rn}"].alignment = Alignment(horizontal="left", vertical="center"); ws[f"A{rn}"].border = bd(); sr(ws, rn, 20)
        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c = ws[f"A{rn+1}"]
        c.value = ex; c.font = Font(name="Meiryo UI", size=9, color="FF6B7280", italic=True)
        c.fill = gf(); c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True); c.border = bd(); sr(ws, rn+1, h)

    ws.merge_cells("A14:H14"); ws["A14"].value = "代表者署名：　　　　　　　　　　　（印）　　年　月　日"
    ws["A14"].font = bfont(); ws["A14"].alignment = Alignment(horizontal="left", vertical="center"); ws["A14"].border = bd(); sr(ws, 14, 30)

    for col, w in zip("ABCDEFGH", [12,24,18,14,8,14,16,12]): ws.column_dimensions[col].width = w
    save(wb, "05_処遇改善加算様式", "03_特定処遇改善計画書.xlsx")


def make_base_up():
    """ベースアップ等支援加算計画書"""
    wb = Workbook(); ws = wb.active; ws.title = "ベースアップ計画書"
    ws.merge_cells("A1:G1")
    ws["A1"].value = "ベースアップ等支援加算 計画書（参考様式）"
    ws["A1"].font = Font(name="Meiryo UI", size=12, bold=True, color=TEAL)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center"); sr(ws, 1, 28)
    ws.merge_cells("A2:G2")
    ws["A2"].value = "【ポイント】ベースアップ等支援加算は、基本給・決まって毎月支払われる手当の引き上げに2/3以上を充てることが要件（処遇改善手当等の一時金は1/3以下）。"
    ws["A2"].font = Font(name="Meiryo UI", size=8, color=RED, bold=True)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, 2, 22)

    info = [
        ("事業者名称", "", "事業所名", ""),
        ("サービス種別", "", "計画年度", "令和　　年度"),
        ("職員総数（常勤換算）", "　　人", "加算取得区分（予定）", "区分Ⅰ / 区分Ⅱ"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info, start=3):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value = v1 if v1 else ""
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value = l2
        ws[f"G{i}"].value = v2
        for cell in [f"A{i}", f"E{i}"]:
            ws[cell].font = bfont(bold=True); ws[cell].fill = sf()
            ws[cell].alignment = Alignment(horizontal="right", vertical="center"); ws[cell].border = bd()
        for cell in [f"C{i}", f"G{i}"]:
            ws[cell].font = bfont(); ws[cell].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[cell].border = bd()
        sr(ws, i, 22)

    ws.merge_cells("A6:G6"); ws["A6"].value = "【ベースアップの方法（具体的に記載）】"
    ws["A6"].font = hfont(sz=9); ws["A6"].fill = hf(); ws["A6"].alignment = Alignment(horizontal="left", vertical="center"); ws["A6"].border = bd(); sr(ws, 6, 20)

    cols_bu = [("A","対象職員区分",16),("B","改善方法（基本給 or 手当名）",22),
               ("C","改善額（円/月）",14),("D","対象人数",8),("E","年間総額（円）",14),
               ("F","うち基本給\n引き上げ相当",14),("G","備考",12)]
    hrow(ws, 7, cols_bu)
    for row in range(8, 14):
        for col in range(1, 8): bc(ws, f"{get_column_letter(col)}{row}")
        sr(ws, row, 22)

    ws.merge_cells("A14:G14"); ws["A14"].value = "【2/3ルールの確認】基本給・毎月支払い手当への充当額：　　　　円　（全体に占める割合：　　　%） ≧ 2/3（66.7%）であること"
    ws["A14"].font = Font(name="Meiryo UI", size=9, color=RED, bold=True)
    ws["A14"].alignment = Alignment(horizontal="left", vertical="center"); ws["A14"].border = bd(); sr(ws, 14, 26)

    ws.merge_cells("A15:G15"); ws["A15"].value = "【就業規則等への明記確認】基本給表または賃金規程の改定：　済み（改定日：　年　月　日）/ 予定（　年　月）"
    ws["A15"].font = bfont(); ws["A15"].alignment = Alignment(horizontal="left", vertical="center"); ws["A15"].border = bd(); sr(ws, 15, 24)

    ws.merge_cells("A16:G16"); ws["A16"].value = "代表者署名：　　　　　　　　　　　（印）　　年　月　日"
    ws["A16"].font = bfont(); ws["A16"].alignment = Alignment(horizontal="left", vertical="center"); ws["A16"].border = bd(); sr(ws, 16, 30)

    for col, w in zip("ABCDEFG", [16,22,14,8,14,14,12]): ws.column_dimensions[col].width = w
    save(wb, "05_処遇改善加算様式", "04_ベースアップ等支援加算計画書.xlsx")


# ======================================================================
# 07_運営規程ひな形（個人情報保護方針）
# ======================================================================

def make_kojin_joho():
    """個人情報保護方針（Word）"""
    doc = Document()
    doc.styles["Normal"].font.name = "Meiryo UI"
    doc.styles["Normal"].font.size = Pt(10.5)
    sec = doc.sections[0]; sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.5); sec.right_margin = Cm(3)

    h = doc.add_heading("個人情報保護方針（プライバシーポリシー）", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x0E, 0x74, 0x90)
    h.runs[0].font.size = Pt(14)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run("　○○事業所（以下「当事業所」という。）は、障害福祉サービスの提供において取り扱う利用者・家族・職員等の個人情報の保護を重要な社会的責務と捉え、以下の方針（個人情報保護方針）を定め、全職員がこれを遵守します。")

    chapters = [
        ("1. 個人情報の取得",
         "当事業所は、サービスの提供・運営上必要な範囲で、適正・公正な手段により個人情報を取得します。取得に際しては、利用目的を明示し、必要に応じて本人の同意を得ます。"),
        ("2. 個人情報の利用目的",
         "取得した個人情報は、以下の目的のために利用します。\n"
         "①障害福祉サービスの提供・記録・管理\n"
         "②介護報酬・障害福祉サービス費の算定・請求\n"
         "③主治医・医療機関・行政機関・他の事業所との連絡調整\n"
         "④職員の採用・労務管理\n"
         "⑤法令に基づく行政への報告・届出"),
        ("3. 個人情報の第三者提供",
         "当事業所は、以下の場合を除き、本人の同意を得ることなく個人情報を第三者に提供しません。\n"
         "①法令に基づく場合（行政機関への届出・報告等）\n"
         "②人の生命・身体・財産の保護に必要な場合\n"
         "③公衆衛生・児童の健全育成に必要な場合\n"
         "④国の機関等の法令の定める事務への協力が必要な場合\n\n"
         "なお、サービス提供に関連する事業者（医療機関・相談支援事業所・介護事業所等）との情報共有については、利用者・家族の同意書に基づき行います。"),
        ("4. 個人情報の安全管理",
         "①書類・記録は施錠可能な場所に保管し、持ち出しを制限します。\n"
         "②電子データは、パスワード管理・ウイルス対策等のセキュリティ措置を講じます。\n"
         "③個人情報へのアクセスは、業務上必要な職員に限定します。\n"
         "④全職員に対して個人情報保護に関する研修を年1回以上実施します。\n"
         "⑤個人情報の取り扱いに関する責任者（管理者）を定めます。"),
        ("5. 個人情報の保存期間と廃棄",
         "個人情報を含む記録は、法令で定める保存期間（最低5年間）を遵守して保管し、期間満了後は適切な方法（シュレッダー処理・消去等）で廃棄します。"),
        ("6. 個人情報の開示・訂正・削除の請求",
         "本人（またはその法定代理人・成年後見人等）から、自己の個人情報の開示・訂正・削除を請求された場合は、合理的な期間内に対応します。請求方法は管理者に申し出てください。"),
        ("7. 苦情・相談窓口",
         "個人情報の取り扱いに関する苦情・相談は、以下の窓口で受け付けます。\n"
         "担当：管理者　電話：○○　受付時間：平日 9:00〜17:00"),
        ("8. 法令の遵守と見直し",
         "当事業所は、個人情報の保護に関する法律（個人情報保護法）及び関連法令・ガイドラインを遵守します。また、本方針は社会状況・法令改正に応じて随時見直し・改定します。"),
        ("附則",
         "本方針は、令和　　年　　月　　日から施行します。\n"
         "○○事業所　代表者　　　　　　　　　　　（署名・捺印）"),
    ]

    for title, content in chapters:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.paragraph_format.space_before = Pt(10)
        doc.add_paragraph(content)

    save_doc(doc, "07_運営規程ひな形", "06_個人情報保護方針.docx")


if __name__ == "__main__":
    print("=== 追加テンプレートバッチ4 ===")
    print("\n--- 01_指定申請様式 ---")
    make_jizen_soudan()
    make_houjin_gaiyo()
    make_juugyousha_ichiran()
    make_shikingakeikaku()
    print("\n--- 02_訪問看護記録様式 ---")
    make_houmon_keikaku()
    make_shiji_sho()
    print("\n--- 03_就労支援記録様式 ---")
    make_monitoring_shuro()
    print("\n--- 04_相談支援様式 ---")
    make_service_riyou_keikaku()
    make_monitoring_soudan()
    make_assessment_soudan()
    print("\n--- 05_処遇改善加算様式 ---")
    make_shoguu_keikaku()
    make_shoguu_jisseki()
    make_tokutei_shoguu()
    make_base_up()
    print("\n--- 07_運営規程ひな形 ---")
    make_kojin_joho()
    print("\n=== 完了 ===")
