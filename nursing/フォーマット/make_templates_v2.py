"""
障碍者福祉事業所 運営ガイド — テンプレート大幅強化版 (v2)
法令・実地指導の必須記載事項に完全対応した全面改訂版

実行すると既存ファイルを上書きし、新規ファイルも生成する。
"""
import os, sys
sys.stdout.reconfigure(encoding="utf-8")
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))

# ── カラー定義 ──────────────────────────────────────────
TEAL  = "FF0E7490"; WHITE = "FFFFFFFF"; LIGHT = "FFCFFAFE"
GRAY  = "FFF8FAFC"; RED   = "FFDC2626"; AMBER = "FFD97706"
GREEN = "FF166534"; GLIGHT= "FFE0F2E8"; WARN  = "FFFEF3C7"

def hf():  return PatternFill("solid", fgColor=TEAL)
def sf():  return PatternFill("solid", fgColor=LIGHT)
def gf():  return PatternFill("solid", fgColor=GRAY)
def wf():  return PatternFill("solid", fgColor=WARN)
def rf():  return PatternFill("solid", fgColor="FFFEE2E2")

def hfont(sz=10, col=WHITE): return Font(name="Meiryo UI", size=sz, bold=True, color=col)
def bfont(sz=9, bold=False, col=None):
    return Font(name="Meiryo UI", size=sz, bold=bold, color=col) if col else Font(name="Meiryo UI", size=sz, bold=bold)
def efont(sz=9): return Font(name="Meiryo UI", size=sz, italic=True, color="FF6B7280")
def bd():
    s = Side(style="thin")
    return Border(left=s, right=s, top=s, bottom=s)
def sr(ws, r, h): ws.row_dimensions[r].height = h
def sw(ws, c, w): ws.column_dimensions[c].width = w

def hrow(ws, row, cols):
    for letter, text, width in cols:
        c = ws[f"{letter}{row}"]
        c.value = text; c.font = hfont(); c.fill = hf()
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border = bd(); ws.column_dimensions[letter].width = width
    sr(ws, row, 30)

def bc(ws, ref, val="", bold=False, ha="left", fill=None):
    c = ws[ref]; c.value = val; c.font = bfont(bold=bold)
    c.alignment = Alignment(horizontal=ha, vertical="center", wrap_text=True); c.border = bd()
    if fill: c.fill = fill
    return c

def ec(ws, ref, val=""):
    c = ws[ref]; c.value = val; c.font = efont()
    c.fill = gf(); c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True); c.border = bd()

def title_row(ws, text, cols_range, row=1, h=28, color=TEAL):
    ws.merge_cells(cols_range)
    ref = cols_range.split(":")[0]
    ws[ref].value = text
    ws[ref].font = Font(name="Meiryo UI", size=13, bold=True, color=color)
    ws[ref].alignment = Alignment(horizontal="center", vertical="center")
    sr(ws, row, h)

def note_row(ws, text, cols_range, row, h=20, color=RED):
    ws.merge_cells(cols_range)
    ref = cols_range.split(":")[0]
    ws[ref].value = text
    ws[ref].font = Font(name="Meiryo UI", size=8, bold=True, color=color)
    ws[ref].alignment = Alignment(horizontal="left", vertical="center")
    sr(ws, row, h)

def section_hdr(ws, text, cols_range, row, h=20):
    ws.merge_cells(cols_range)
    ref = cols_range.split(":")[0]
    ws[ref].value = text; ws[ref].font = hfont(sz=9)
    ws[ref].fill = hf(); ws[ref].alignment = Alignment(horizontal="left", vertical="center")
    ws[ref].border = bd(); sr(ws, row, h)

def info_row(ws, row, label, val, col_l="A", col_v="C", col_end="H", label_w=14, val_bold=False):
    ws.merge_cells(f"{col_l}{row}:{chr(ord(col_v)-1)}{row}")
    ws[f"{col_l}{row}"].value = label
    ws[f"{col_l}{row}"].font = bfont(bold=True); ws[f"{col_l}{row}"].fill = sf()
    ws[f"{col_l}{row}"].alignment = Alignment(horizontal="right", vertical="center")
    ws[f"{col_l}{row}"].border = bd()
    ws.merge_cells(f"{col_v}{row}:{col_end}{row}")
    ws[f"{col_v}{row}"].value = val; ws[f"{col_v}{row}"].font = bfont(bold=val_bold)
    ws[f"{col_v}{row}"].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    ws[f"{col_v}{row}"].border = bd()
    sr(ws, row, 22)

def text_block(ws, row, title, hint, cols_range_title, cols_range_body, h_body=60):
    t_row_str = cols_range_title.split(":")[0]
    b_row_str = cols_range_body.split(":")[0]
    ws.merge_cells(cols_range_title)
    ws[t_row_str].value = title; ws[t_row_str].font = hfont(sz=9); ws[t_row_str].fill = hf()
    ws[t_row_str].alignment = Alignment(horizontal="left", vertical="center"); ws[t_row_str].border = bd(); sr(ws, row, 20)
    ws.merge_cells(cols_range_body)
    c = ws[b_row_str]; c.value = hint; c.font = efont()
    c.fill = gf(); c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    c.border = bd(); sr(ws, row+1, h_body)

def save(wb, folder, fname):
    p = os.path.join(BASE, folder, fname); wb.save(p)
    print(f"OK: {folder}/{fname}")

def save_doc(doc, folder, fname):
    p = os.path.join(BASE, folder, fname); doc.save(p)
    print(f"OK: {folder}/{fname}")

# ======================================================================
# 02_訪問看護記録様式
# ======================================================================

def make_kiroku1_facesheet():
    """訪問看護記録書Ⅰ（フェイスシート） — 法令必須項目完全版"""
    wb = Workbook(); ws = wb.active; ws.title = "フェイスシート"
    title_row(ws, "訪問看護記録書Ⅰ（基本情報・フェイスシート）", "A1:J1")
    note_row(ws, "【保管義務】完結の日から5年間保管。状態変化・療養環境変化時に随時更新すること。", "A2:J2", 2)

    # ── 基本情報 ─────────────────────────────────────────
    section_hdr(ws, "【基本情報】", "A3:J3", 3)
    rows_basic = [
        ("利用者氏名（ふりがな）", "（　　　　　　　　　　　　　　　）様"),
        ("生年月日・年齢", "　　年　月　日生（　　歳）　男 ・ 女"),
        ("住所・電話番号", "〒　　－　　　　　TEL："),
        ("主傷病名（主たる傷病）", ""),
        ("副傷病名・合併症", ""),
        ("現病歴（発症経緯・経過・現在の状態）", "例）〇年〇月に脳梗塞発症、左片麻痺残存。現在外来リハ継続中。"),
        ("既往歴（手術歴・入院歴・重要な病歴）", "例）H25年 大腸がん手術（人工肛門造設）、糖尿病10年"),
        ("アレルギー（薬剤・食品・その他）", "なし / あり（　　　　　　　　　　）"),
        ("感染症・注意事項", "なし / あり（　　　　　　　　　　）"),
    ]
    for i, (label, val) in enumerate(rows_basic, start=4):
        info_row(ws, i, label, val, "A", "D", "J")

    # ── 保険・認定情報 ──────────────────────────────────
    section_hdr(ws, "【保険・認定情報】", "A13:J13", 13)
    rows_ins = [
        ("保険の種類", "健康保険 / 国民健康保険 / 後期高齢者医療 / 自立支援医療（精神通院） / その他（　　）"),
        ("自立支援医療 受給者証番号", "なし / あり（番号：　　　　　　　上限月額：　　　円）"),
        ("要介護認定区分", "なし / 要支援（　）/ 要介護（　）  認定日：　　年　月　日"),
        ("障害支援区分", "なし / 区分（　）  認定日：　　年　月　日"),
        ("身体障害者手帳", "なし / あり（　　級・種別：　　　　）"),
        ("精神障害者保健福祉手帳", "なし / あり（　　級）"),
        ("訪問看護指示書 有効期限", "　　年　月　日まで　　　　医療機関名："),
    ]
    for i, (label, val) in enumerate(rows_ins, start=14):
        info_row(ws, i, label, val, "A", "D", "J")

    # ── 医療情報 ──────────────────────────────────────
    section_hdr(ws, "【主治医・医療機関情報】", "A21:J21", 21)
    rows_med = [
        ("主治医氏名・診療科", ""),
        ("医療機関名・所在地", ""),
        ("医療機関 電話番号", "TEL：　　　　　　　　　  FAX："),
        ("訪問看護の依頼目的", "例）褥瘡処置・服薬管理・リハビリ・精神科訪問看護・ターミナルケア"),
        ("訪問頻度（指示内容）", "週　　回　時間/回　担当職種：看護師 / 准看護師 / PT / OT / ST"),
    ]
    for i, (label, val) in enumerate(rows_med, start=22):
        info_row(ws, i, label, val, "A", "D", "J")

    # ── 緊急連絡先・関係機関 ──────────────────────────
    section_hdr(ws, "【緊急連絡先・関係機関】", "A27:J27", 27)
    rows_emer = [
        ("緊急連絡先①（続柄・氏名）", "　続柄：　　　　　氏名：　　　　　TEL（日中）：　　　　TEL（夜間）："),
        ("緊急連絡先②（続柄・氏名）", "　続柄：　　　　　氏名：　　　　　TEL（日中）：　　　　TEL（夜間）："),
        ("緊急搬送先病院", "　医療機関名：　　　　　　　　　　TEL：　　　　　　　　（主治医の指定先）"),
        ("担当相談支援専門員", "氏名：　　　　　　　　　事業所：　　　　　　　　TEL："),
        ("ケアマネジャー（介護保険）", "氏名：　　　　　　　　　事業所：　　　　　　　　TEL："),
        ("居宅介護支援事業所", "事業所名：　　　　　　　　　　　　TEL："),
        ("その他関係機関（デイ・訪介等）", ""),
    ]
    for i, (label, val) in enumerate(rows_emer, start=28):
        info_row(ws, i, label, val, "A", "D", "J")

    # ── 生活・介護状況 ────────────────────────────────
    section_hdr(ws, "【生活状況・療養環境】", "A35:J35", 35)
    rows_life = [
        ("居住形態", "一人暮らし / 家族と同居 / グループホーム / 施設（　　　　　　）"),
        ("住環境・バリアフリー", "戸建て / マンション　エレベーター：あり/なし　段差：あり/なし  その他："),
        ("主な介護者", "氏名：　　　　　　続柄：　　　　　介護力：良好 / 不十分 / なし"),
        ("生活歴（仕事歴・趣味・習慣）", "例）元会社員、趣味は読書・囲碁、毎朝散歩の習慣あり"),
        ("本人の療養・生活に対する意向", "例）「できるだけ家にいたい」「家族に迷惑をかけたくない」"),
        ("家族の介護・療養に対する意向", ""),
    ]
    for i, (label, val) in enumerate(rows_life, start=36):
        info_row(ws, i, label, val, "A", "D", "J")

    # ── ADL評価 ───────────────────────────────────────
    section_hdr(ws, "【ADL評価（Barthel Index 準拠）】　　評価日：　　年　月　日", "A42:J42", 42)
    hrow(ws, 43, [("A","項目",18),("B","自立/介助",12),("C","状況の詳細",32),
                  ("D","前回評価",12),("E","変化",8),("F","備考",20)])
    adl_items = [
        ("食事", "自立/一部介助/全介助", ""), ("移乗", "自立/一部介助/全介助", ""),
        ("整容", "自立/介助", ""), ("トイレ動作", "自立/一部介助/全介助", ""),
        ("入浴", "自立/介助", ""), ("歩行", "自立/補装具/介助/車椅子", ""),
        ("階段", "自立/介助/不可", ""), ("更衣", "自立/一部介助/全介助", ""),
        ("排便コントロール", "自立/時々失禁/常に失禁", ""),
        ("排尿コントロール", "自立/時々失禁/常に失禁", ""),
    ]
    for i, (item, val, detail) in enumerate(adl_items, start=44):
        bc(ws, f"A{i}", item, bold=True, fill=sf()); ws[f"A{i}"].border = bd()
        bc(ws, f"B{i}", val); bc(ws, f"C{i}", detail)
        bc(ws, f"D{i}"); bc(ws, f"E{i}"); bc(ws, f"F{i}")
        sr(ws, i, 22)

    # ── 認知・意思疎通 ──────────────────────────────
    section_hdr(ws, "【認知機能・意思疎通・精神状態】", "A54:J54", 54)
    cog_rows = [
        ("認知症の有無・程度", "なし / 軽度 / 中等度 / 重度　日常生活自立度：Ⅰ/Ⅱa/Ⅱb/Ⅲa/Ⅲb/Ⅳ/M"),
        ("意思疎通の状況", "良好 / やや困難 / 困難（コミュニケーション方法：　　　　　　　　　　）"),
        ("精神状態・情動", "安定 / 不安定（内容：　　　　　　　　　　　　　　　　　　　　　　　）"),
        ("睡眠状況", "良好 / 不良（　　時間程度　中途覚醒：あり/なし　睡眠薬：あり/なし）"),
    ]
    for i, (label, val) in enumerate(cog_rows, start=55):
        info_row(ws, i, label, val, "A", "D", "J")

    # ── 医療処置・服薬 ──────────────────────────────
    section_hdr(ws, "【医療処置・医療機器管理】", "A59:J59", 59)
    text_block(ws, 59, "【医療処置・医療機器管理】", "例）気管切開・吸引（1日　回）　人工呼吸器（　　）　経管栄養（NG/胃ろう/腸ろう）　留置カテーテル　輸液（　）　褥瘡処置（　部位）　点滴",
               "A59:J59", "A60:J60", 60)
    text_block(ws, 60, "【服薬管理・処方薬一覧】", "例）降圧薬（　　）朝1錠、血糖降下薬（　　）毎食後、抗精神病薬（　　）就寝前　→ 服薬管理：自己管理/一包化/ヘルパー管理",
               "A61:J61", "A62:J62", 55)
    text_block(ws, 62, "【特記事項・支援上の注意点】", "例）水分摂取が少ない傾向あり（誤嚥リスク）。右足に浮腫あり。家族への指導事項：体位変換2時間毎の励行",
               "A63:J63", "A64:J64", 50)

    # 署名欄
    r = 65
    ws.merge_cells(f"A{r}:J{r}")
    ws[f"A{r}"].value = "作成日：　　年　月　日　担当看護師（署名）：　　　　　　　　　最終更新日：　　年　月　日　更新者：　　　　　　　　"
    ws[f"A{r}"].font = bfont(); ws[f"A{r}"].alignment = Alignment(horizontal="left", vertical="center")
    ws[f"A{r}"].border = bd(); sr(ws, r, 28)

    for col, w in zip("ABCDEFGHIJ", [18,12,14,12,12,12,12,12,12,12]):
        ws.column_dimensions[col].width = w

    save(wb, "02_訪問看護記録様式", "06_訪問看護記録書Ⅰ（フェイスシート）.xlsx")


def make_kiroku2_keika():
    """訪問看護記録書Ⅱ（経過記録） — SOAP・バイタル・職種・疼痛スケール完全版"""
    wb = Workbook(); ws = wb.active; ws.title = "訪問看護記録書Ⅱ"
    ws.page_setup.orientation = "landscape"
    title_row(ws, "訪問看護記録書Ⅱ（経過記録）", "A1:O1")
    note_row(ws, "【修正方法】修正液使用禁止。二重線で消し、訂正者署名・日付を記入。遡及記録・追記禁止。【保管】5年間", "A2:O2", 2)

    # 基本情報行
    ws.merge_cells("A3:D3"); ws["A3"].value = "利用者氏名："
    ws.merge_cells("E3:H3")
    ws.merge_cells("I3:J3"); ws["I3"].value = "生年月日："
    ws.merge_cells("K3:L3")
    ws.merge_cells("M3:O3"); ws["M3"].value = "（　　　　）年度"
    for ref in ["A3","I3","M3"]: ws[ref].font = bfont(bold=True)
    sr(ws, 3, 20)

    # ヘッダー
    cols = [
        ("A","訪問日",7), ("B","訪問\n職種",7), ("C","開始\n時刻",6), ("D","終了\n時刻",6),
        ("E","体温\n(℃)",5), ("F","血圧\n(mmHg)",8), ("G","脈拍\n(/分)",6),
        ("H","SpO2\n(%)",6), ("I","呼吸\n(/分)",6), ("J","疼痛\nNRS",5),
        ("K","意識\n(JCS)",6),
        ("L","訪問看護の内容・利用者の状態\n（S：主観 O：客観 A：評価 P：計画）",52),
        ("M","実施した\n看護・リハ内容",20), ("N","特記事項・\n連絡事項",18), ("O","記録者",8)
    ]
    hrow(ws, 4, cols)

    # 記入例
    ex = ["5/1(木)","看護師","10:00","11:00","36.5","128/76","74","97","18","2","0",
          "【S】昨夜は眠れた。足のむくみが気になると本人。\n【O】顔色良好。両下腿浮腫(+)右>左。褥瘡なし。排泄自立。服薬確認済み。\n【A】バイタル安定。浮腫は前回比不変。誤嚥リスク継続観察要。\n【P】次回も同様のケア継続。排泄状況・浮腫を継続観察。",
          "バイタル測定・清拭介助・服薬確認・褥瘡観察・患者教育（水分摂取促進）",
          "主治医報告不要。家族（長男）より電話で状況確認→問題なし。","山田 看護師"]
    letters = [get_column_letter(i+1) for i in range(15)]
    for i, val in enumerate(ex):
        ec(ws, f"{letters[i]}5", val)
    sr(ws, 5, 80)

    # データ行（22行）
    for row in range(6, 28):
        for col in range(1, 16):
            c = ws.cell(row=row, column=col)
            c.border = bd(); c.font = bfont(sz=8)
            c.alignment = Alignment(horizontal="left" if col >= 12 else "center",
                                   vertical="top" if col >= 12 else "center",
                                   wrap_text=True)
        sr(ws, row, 50)

    # 凡例・注記
    r = 28
    ws.merge_cells(f"A{r}:O{r}")
    ws[f"A{r}"].value = "【訪問職種凡例】看：看護師　准：准看護師　PT：理学療法士　OT：作業療法士　ST：言語聴覚士　【疼痛NRS】0=痛みなし〜10=最大の痛み　【意識JCS】0=清明　1〜3桁=障害あり"
    ws[f"A{r}"].font = Font(name="Meiryo UI", size=8, color="FF6B7280")
    ws[f"A{r}"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, r, 18)

    save(wb, "02_訪問看護記録様式", "03_訪問看護記録書Ⅱ経過記録.xlsx")


def make_kyotaku_service_kiroku():
    """居宅介護サービス提供記録 — 法令必須項目完全版（新規作成）"""
    wb = Workbook(); ws = wb.active; ws.title = "サービス提供記録"
    ws.page_setup.orientation = "landscape"
    title_row(ws, "居宅介護 サービス提供記録（実施記録）", "A1:M1")
    note_row(ws,
        "【法令根拠】指定障害福祉サービスの事業等の人員・設備及び運営に関する基準第17条：サービス提供の都度記録義務あり。【保管】5年間。"
        "【注意】利用者確認署名は毎回必須。時刻は実際の開始・終了を分単位で記録すること。",
        "A2:M2", 2)

    # 利用者基本情報
    ws.merge_cells("A3:D3"); ws["A3"].value = "利用者氏名："
    ws.merge_cells("E3:H3")
    ws.merge_cells("I3:J3"); ws["I3"].value = "事業所名："
    ws.merge_cells("K3:M3")
    for ref in ["A3","I3"]: ws[ref].font = bfont(bold=True); ws[ref].fill = sf()
    sr(ws, 3, 22)
    ws.merge_cells("A4:D4"); ws["A4"].value = "担当サービス提供責任者："
    ws.merge_cells("E4:H4")
    ws.merge_cells("I4:J4"); ws["I4"].value = "個別支援計画 作成日："
    ws.merge_cells("K4:M4")
    for ref in ["A4","I4"]: ws[ref].font = bfont(bold=True); ws[ref].fill = sf()
    sr(ws, 4, 22)

    # ヘッダー
    cols = [
        ("A","提供\n年月日",10), ("B","曜\n日",4), ("C","担当\nヘルパー氏名",14),
        ("D","サービス\n種別",10), ("E","開始\n時刻",7), ("F","終了\n時刻",7),
        ("G","実施した具体的なサービス内容\n（何を・どのように・どの程度実施したか）",38),
        ("H","利用者の心身の状況\n（状態変化・特記事項）",22),
        ("I","サービス計画との\n相違・理由",14),
        ("J","利用者\n確認署名",12), ("K","ヘルパー\n署名",10),
        ("L","管理者\n確認",8), ("M","備考",10),
    ]
    hrow(ws, 5, cols)

    # 種別凡例行
    ws.merge_cells("A6:M6")
    ws["A6"].value = "【サービス種別の記入方法】身：身体介護　家：家事援助　通身：通院等介助（身体介護あり）　通：通院等介助（身体介護なし）　乗：通院等乗降介助"
    ws["A6"].font = Font(name="Meiryo UI", size=8, color="FF6B7280")
    ws["A6"].alignment = Alignment(horizontal="left", vertical="center"); sr(ws, 6, 16)

    # 記入例
    ex = ["5/10(土)","土","山田 花子","身","10:00","11:30",
          "①全身清拭（背部・四肢・陰部）　②シャツ・ズボン更衣介助　③口腔ケア（歯磨き）　④体位変換（左側臥位→仰臥位→右側臥位）　⑤水分補給100ml確認",
          "バイタル：BP 128/76、P 74、体温 36.5℃。顔色良好。右踵部に発赤なし（前回比改善）。本人「体が楽になった」と。",
          "なし", "○○様（拇印）", "山田 花子", "", ""]
    letters = [get_column_letter(i+1) for i in range(13)]
    for i, val in enumerate(ex):
        ec(ws, f"{letters[i]}7", val)
    sr(ws, 7, 60)

    # データ行（20行）
    for row in range(8, 28):
        for col in range(1, 14):
            c = ws.cell(row=row, column=col)
            c.border = bd(); c.font = bfont(sz=8)
            c.alignment = Alignment(
                horizontal="center" if col in [1,2,4,5,6,10,11,12] else "left",
                vertical="top" if col in [7,8,9] else "center",
                wrap_text=True)
        sr(ws, row, 55)

    # 月次集計行
    r = 28
    ws.merge_cells(f"A{r}:D{r}")
    ws[f"A{r}"].value = "当月合計"
    ws[f"A{r}"].font = bfont(bold=True); ws[f"A{r}"].fill = sf()
    ws[f"A{r}"].alignment = Alignment(horizontal="center", vertical="center"); ws[f"A{r}"].border = bd()
    for col_l, label in [("E","身体介護計（h）"),("G","家事援助計（h）"),("I","通院等計（回）")]:
        bc(ws, f"{col_l}{r}", label, bold=True, fill=sf())
    ws.merge_cells(f"F{r}:F{r}"); bc(ws, f"F{r}")
    ws.merge_cells(f"H{r}:H{r}"); bc(ws, f"H{r}")
    for col in ["J","K","L","M"]: bc(ws, f"{col}{r}")
    sr(ws, r, 24)

    # 管理者月次確認欄
    ws.merge_cells(f"A{r+1}:M{r+1}")
    ws[f"A{r+1}"].value = "【管理者月次確認】サービス提供記録と請求データとの照合実施日：　　年　月　日　確認者：　　　　　　（署名）"
    ws[f"A{r+1}"].font = bfont(); ws[f"A{r+1}"].alignment = Alignment(horizontal="left", vertical="center")
    ws[f"A{r+1}"].border = bd(); sr(ws, r+1, 28)

    save(wb, "02_訪問看護記録様式", "12_居宅介護サービス提供記録.xlsx")


def make_houmon_houkoku_v2():
    """訪問看護報告書 — 主治医への必須報告事項完全版"""
    wb = Workbook(); ws = wb.active; ws.title = "訪問看護報告書"
    title_row(ws, "訪問看護報告書", "A1:H1")
    note_row(ws, "【提出義務】毎月、主治医（指示を行った医師）へ提出。精神科訪問看護は別紙様式4を使用。内容は訪問看護計画書に基づき記載。", "A2:H2", 2)

    info_pairs = [
        ("利用者氏名", "（様）", "報告対象月", "　　年　　月分"),
        ("主治医氏名・医療機関名", "", "訪問看護ステーション名", ""),
        ("疾患・傷病名（主診断）", "", "管理者名", ""),
        ("訪問看護指示書 有効期限", "　　年　月　日まで", "担当看護師", ""),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info_pairs, start=3):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value = l2
        ws.merge_cells(f"G{i}:H{i}"); ws[f"G{i}"].value = v2
        for ref in [f"A{i}", f"E{i}"]:
            ws[ref].font = bfont(bold=True); ws[ref].fill = sf()
            ws[ref].alignment = Alignment(horizontal="right", vertical="center"); ws[ref].border = bd()
        for ref in [f"C{i}", f"G{i}"]:
            ws[ref].font = bfont(); ws[ref].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[ref].border = bd()
        sr(ws, i, 22)

    # 訪問実績サマリー
    section_hdr(ws, "【当月訪問実績】", "A7:H7", 7)
    ws.merge_cells("A8:H8")
    ws["A8"].value = "訪問回数：　　回　 うち緊急訪問：　　回　 特別訪問看護指示期間：　なし　・　あり（　　日〜　　日）　リハ訪問：　　回（PT/OT/ST）"
    ws["A8"].font = bfont(); ws["A8"].alignment = Alignment(horizontal="left", vertical="center"); ws["A8"].border = bd(); sr(ws, 8, 24)

    # 主要報告セクション
    sections = [
        (9, "【病状・バイタルサインの経過（月間）】", 70,
         "・体温：　〜　℃台（平均：　　℃）　・血圧：　/　〜　/　mmHg　・脈拍：　〜　回/分　・SpO2：　〜　%\n"
         "・意識レベル：清明 / 傾眠傾向 / その他（　　　　　　）\n"
         "・特記すべき病状変化："),
        (11, "【看護・リハビリテーションの内容（月間）】", 80,
         "・清拭・入浴介助：　有（　回）　無\n・創傷・褥瘡処置：　有（部位・状態：　　　　　　）　無\n"
         "・服薬管理・指導：　有　無\n・吸引：　有（　回）　無\n・経管栄養管理：　有　無\n"
         "・リハビリ（PT/OT/ST）：　有（内容：　　　　　　　）　無\n・その他："),
        (13, "【ADL・日常生活状況の変化】", 60,
         "前月比較：改善（　　） / 不変 / 低下（　　）\n"
         "食事：　　　移動：　　　排泄：　　　入浴：　　　意思疎通：　　　"),
        (15, "【家族への指導・支援内容】", 50,
         "例）体位変換の方法を指導。吸引手技の確認。緊急時対応手順の再確認。"),
        (17, "【医療機器・医療処置の管理状況】", 40,
         "例）人工呼吸器：設定変更なし、作動良好。カテーテル：交換実施（　　日）。"),
        (19, "【精神状態・療養意欲・生活面の状況】", 50,
         "例）「体の具合はどうですか」の問いに「まあまあです」と返答あり。外出意欲あり。不安・抑うつなし。"),
        (21, "【今後の看護・支援の方向性（主治医へ）】", 60,
         "例）引き続き現状のケアを継続。次月は通院同行の予定あり。〇〇について主治医の指示を仰ぎたい。"),
        (23, "【特記事項・連絡事項（主治医への報告事項）】", 50,
         "例）〇日に転倒あり（処置済み）。〇日に家族より体調不良の連絡を受け緊急訪問実施。"),
    ]
    for rn, title, h, ex in sections:
        ws.merge_cells(f"A{rn}:H{rn}"); ws[f"A{rn}"].value = title
        ws[f"A{rn}"].font = hfont(sz=9); ws[f"A{rn}"].fill = hf()
        ws[f"A{rn}"].alignment = Alignment(horizontal="left", vertical="center"); ws[f"A{rn}"].border = bd(); sr(ws, rn, 20)
        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c = ws[f"A{rn+1}"]
        c.value = ex; c.font = efont()
        c.fill = gf(); c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        c.border = bd(); sr(ws, rn+1, h)

    # 署名欄
    r = 25
    ws.merge_cells(f"A{r}:H{r}")
    ws[f"A{r}"].value = "報告日：　　年　月　日　　担当看護師署名：　　　　　　　　　　管理者確認署名：　　　　　　　　　（確認日：　年　月　日）"
    ws[f"A{r}"].font = bfont(); ws[f"A{r}"].alignment = Alignment(horizontal="left", vertical="center")
    ws[f"A{r}"].border = bd(); sr(ws, r, 30)

    for col, w in zip("ABCDEFGH", [14,14,14,14,14,14,14,14]): ws.column_dimensions[col].width = w
    save(wb, "02_訪問看護記録様式", "04_訪問看護報告書.xlsx")


# ======================================================================
# 03_就労支援記録様式
# ======================================================================

def make_kojin_shien_v2():
    """個別支援計画書（就労継続支援A型・B型） — 法令必須項目完全版"""
    wb = Workbook(); ws = wb.active; ws.title = "個別支援計画書"
    title_row(ws, "個別支援計画書（就労継続支援A型・B型）", "A1:I1")
    note_row(ws,
        "【法令根拠】障害者総合支援法基準省令第58条・59条。サービス管理責任者が作成。利用者に説明し同意を得ること。6ヶ月ごとに見直し。",
        "A2:I2", 2)

    # 基本情報
    section_hdr(ws, "【基本情報】", "A3:I3", 3)
    basic = [
        ("利用者氏名", "（様）　　フリガナ：", "生年月日・年齢", "　　年　月　日生（　　歳）"),
        ("障害の種別・等級", "身体(　級) / 精神(　級) / 知的(療育　) / 難病 / 複合", "障害支援区分", "区分　　　　認定日：　　年　月"),
        ("サービス種別", "就労継続支援　A型 / B型", "利用定員（施設）", "　　名"),
        ("計画作成日", "　　年　月　日", "計画期間", "　　年　月　日　〜　　年　月　日"),
        ("計画の種別", "新規 / 更新（前回作成日：　　年　月　日）", "サービス管理責任者", ""),
        ("前回のモニタリング実施日", "　　年　月　日", "モニタリング頻度", "6ヶ月ごと（状態変化時は随時）"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(basic, start=4):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:E{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"F{i}:G{i}"); ws[f"F{i}"].value = l2
        ws.merge_cells(f"H{i}:I{i}"); ws[f"H{i}"].value = v2
        for ref in [f"A{i}", f"F{i}"]:
            ws[ref].font = bfont(bold=True); ws[ref].fill = sf()
            ws[ref].alignment = Alignment(horizontal="right", vertical="center"); ws[ref].border = bd()
        for ref in [f"C{i}", f"H{i}"]:
            ws[ref].font = bfont(); ws[ref].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[ref].border = bd()
        sr(ws, i, 22)

    # アセスメント（意向）
    row = 10
    text_block(ws, row, "【本人の意向・夢・将来への希望】（本人の言葉でそのまま記録する）",
               "例）「毎日通えるようになりたい」「給料をもらって自分で買い物したい」「いつか一般の仕事をしてみたい」",
               "A10:I10", "A11:I11", 55)
    text_block(ws, 11, "【家族・支援者の意向】",
               "例）「無理せず続けてほしい」「体調の波があるので柔軟に対応してほしい」",
               "A12:I12", "A13:I13", 45)

    # 現状評価
    section_hdr(ws, "【現状評価・アセスメント結果】", "A14:I14", 14)
    assess_items = [
        ("就労能力・作業能力", "作業持続時間：　　h　作業の種類：　　　　　　集中力：高/中/低　指示理解：良好/困難"),
        ("健康・体調管理", "通院状況：　　　　服薬管理：自己/支援要　体調の波：大/小　精神状態：安定/不安定"),
        ("対人関係・社会性", "対人関係：良好/課題あり　　コミュニケーション：良好/困難"),
        ("生活スキル・自立度", "交通機関の利用：自立/支援要　金銭管理：自立/支援要　生活リズム：安定/不安定"),
        ("障害の特性・配慮事項", "例）大きな音に過敏。疲れると自傷行為が出る可能性あり。定期的な休憩が必要。"),
    ]
    for i, (label, val) in enumerate(assess_items, start=15):
        ws.merge_cells(f"A{i}:C{i}"); ws[f"A{i}"].value = label
        ws.merge_cells(f"D{i}:I{i}"); ws[f"D{i}"].value = val
        ws[f"A{i}"].font = bfont(bold=True); ws[f"A{i}"].fill = sf()
        ws[f"A{i}"].alignment = Alignment(horizontal="right", vertical="center"); ws[f"A{i}"].border = bd()
        ws[f"D{i}"].font = bfont(); ws[f"D{i}"].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        ws[f"D{i}"].border = bd(); sr(ws, i, 28)

    # 課題・目標
    section_hdr(ws, "【生活全般の解決すべき課題（ニーズ）と優先順位】", "A20:I20", 20)
    hrow(ws, 21, [("A","優先\n順位",6),("B","課題（ニーズ）",28),("C","本人・家族の意向",20),
                  ("D","長期目標（6ヶ月〜1年）",22),("E","短期目標（〜3ヶ月）",22),
                  ("F","達成時期",10),("G","評価指標",16),("H","担当者",10),("I","備考",8)])
    for row in range(22, 26):
        for col in "ABCDEFGHI": bc(ws, f"{col}{row}")
        sr(ws, row, 40)

    # 支援内容テーブル（法令必須）
    section_hdr(ws, "【提供するサービスの内容・量・担当者】（法令必須記載事項）", "A26:I26", 26)
    hrow(ws, 27, [("A","支援内容の種別",18),("B","具体的な支援内容",28),("C","頻度・時間",12),
                  ("D","支援量（月）",10),("E","担当者・職種",14),("F","開始時期",10),("G","終了時期",10),
                  ("H","留意事項",14),("I","特記",8)])
    support_examples = [
        ["就労・作業支援", "封入作業・箱折り・清掃作業の指導・支援", "週4日・6時間/日", "約96時間/月", "職業指導員　田中", "R7.4.1〜", "計画期間末", "疲れたら申告するよう促す", ""],
        ["生活支援", "毎日の健康チェック・服薬確認・昼食の準備補助", "毎日", "月20日", "生活支援員　鈴木", "R7.4.1〜", "計画期間末", "体調変化時は管理者へ報告", ""],
        ["相談支援", "週1回の個別面談・困りごとの相談", "週1回30分", "月4回", "サービス管理責任者", "R7.4.1〜", "計画期間末", "", ""],
        ["医療機関との連絡・連携", "主治医との定期的な情報共有（月1回）", "月1回", "月1回", "サービス管理責任者", "R7.4.1〜", "継続", "", ""],
    ]
    for i, row_data in enumerate(support_examples, start=28):
        for j, val in enumerate(row_data):
            ec(ws, f"{get_column_letter(j+1)}{i}", val)
        sr(ws, i, 35)
    for row in range(32, 36):
        for col in "ABCDEFGHI": bc(ws, f"{col}{row}")
        sr(ws, row, 32)

    # 工賃目標（B型のみ）
    section_hdr(ws, "【工賃目標・経済的自立（B型のみ・A型は賃金目標を記入）】", "A36:I36", 36)
    ws.merge_cells("A37:I37")
    ws["A37"].value = "当月の目標工賃月額：　　　　円（前月実績：　　　　円）　一般就労移行の希望：あり / なし / 将来的に検討"
    ws["A37"].font = bfont(); ws["A37"].alignment = Alignment(horizontal="left", vertical="center"); ws["A37"].border = bd(); sr(ws, 37, 24)

    # 関係機関連携
    section_hdr(ws, "【関係機関との連携内容】", "A38:I38", 38)
    ws.merge_cells("A39:I39")
    ws["A39"].value = "主治医：　　　　　　（連携方法・頻度：　　　）　相談支援専門員：　　　　（連携方法：　　　）　家族：　　　（連携方法：　　　）"
    ws["A39"].font = bfont(); ws["A39"].alignment = Alignment(horizontal="left", vertical="center"); ws["A39"].border = bd(); sr(ws, 39, 24)

    # 本人同意
    section_hdr(ws, "【本人・家族の同意欄】（法令必須）", "A40:I40", 40)
    ws.merge_cells("A41:I41")
    ws["A41"].value = "上記の個別支援計画について、サービス管理責任者より説明を受け、内容を確認し、同意しました。\n\n本人署名：　　　　　　　　　　　（　　年　月　日）　　保護者・代理人署名：　　　　　　　　続柄：　　　　（　　年　月　日）\n\nサービス管理責任者署名：　　　　　　　　　　（説明実施日：　　年　月　日）　　管理者確認：　　　　　　　　（　　年　月　日）"
    ws["A41"].font = bfont(); ws["A41"].alignment = Alignment(horizontal="left", vertical="top", wrap_text=True); ws["A41"].border = bd(); sr(ws, 41, 80)

    for col, w in zip("ABCDEFGHI", [18,28,12,10,14,10,10,14,8]): ws.column_dimensions[col].width = w
    save(wb, "03_就労支援記録様式", "06_個別支援計画書（就労継続支援版）.xlsx")


def make_kyotaku_keikaku():
    """個別支援計画書（居宅介護・重度訪問介護・同行援護・行動援護用）— 新規作成"""
    wb = Workbook(); ws = wb.active; ws.title = "個別支援計画書（介護系）"
    title_row(ws, "個別支援計画書（居宅介護・重度訪問介護・同行援護・行動援護）", "A1:I1")
    note_row(ws,
        "【法令根拠】指定障害福祉サービスの事業等の人員・設備及び運営に関する基準第24条。サービス提供責任者が作成・交付。6ヶ月ごとに見直し必須。",
        "A2:I2", 2)

    section_hdr(ws, "【基本情報】", "A3:I3", 3)
    basic = [
        ("利用者氏名", "（様）", "生年月日・年齢", "　年　月　日生（　歳）"),
        ("住所", "", "電話番号", ""),
        ("障害の種別・等級", "身体(　級) / 精神(　級) / 知的(療育　) / 難病", "障害支援区分", "区分　　　認定日：　年　月"),
        ("利用するサービス種別", "居宅介護 / 重度訪問介護 / 同行援護 / 行動援護", "支給決定期間", "　年　月　日〜　年　月　日"),
        ("計画作成日", "　　年　月　日", "計画期間", "　年　月　日〜　年　月　日"),
        ("担当サービス提供責任者", "", "前回計画作成日", "　　年　月　日"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(basic, start=4):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:E{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"F{i}:G{i}"); ws[f"F{i}"].value = l2
        ws.merge_cells(f"H{i}:I{i}"); ws[f"H{i}"].value = v2
        for ref in [f"A{i}", f"F{i}"]:
            ws[ref].font = bfont(bold=True); ws[ref].fill = sf()
            ws[ref].alignment = Alignment(horizontal="right", vertical="center"); ws[ref].border = bd()
        for ref in [f"C{i}", f"H{i}"]:
            ws[ref].font = bfont(); ws[ref].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws[ref].border = bd()
        sr(ws, i, 22)

    # 意向・目標
    row = 10
    text_block(ws, row, "【本人の意向・生活の希望】（本人の言葉でそのまま記録）",
               "例）「できる限り自分で生活したい」「週に2回お風呂に入りたい」「デイに行く日は自分で準備したい」",
               "A10:I10", "A11:I11", 50)
    text_block(ws, 11, "【家族・支援者の意向】",
               "例）「介護負担が大きいので週3回の支援をお願いしたい」",
               "A12:I12", "A13:I13", 40)
    text_block(ws, 13, "【生活全般の課題（ニーズ）】",
               "例）①入浴が一人では困難→週2回の入浴介助が必要　②服薬管理が難しい→毎日の服薬確認支援が必要",
               "A14:I14", "A15:I15", 50)
    text_block(ws, 15, "【長期目標（6ヶ月〜1年）】",
               "例）自宅でできる限り安全に生活できる状態を維持する",
               "A16:I16", "A17:I17", 40)
    text_block(ws, 17, "【短期目標（〜3ヶ月）】",
               "例）週2回の入浴を継続し、清潔を保つ。服薬を毎日確認し、飲み忘れゼロを目指す",
               "A18:I18", "A19:I19", 40)

    # サービス内容表（法令必須）
    section_hdr(ws, "【提供するサービスの内容・頻度・担当者】（法令必須記載事項）", "A20:I20", 20)
    hrow(ws, 21, [("A","サービス種別\n・内容の区分",16),("B","具体的な支援内容",30),
                  ("C","曜日・時間帯",12),("D","1回あたり\n時間",9),
                  ("E","月間\n提供回数",8),("F","担当ヘルパー",12),
                  ("G","開始日",10),("H","留意事項・特記",16),("I","備考",8)])
    examples = [
        ["身体介護", "全身清拭・シャツ更衣・口腔ケア・体位変換", "火・金 10:00〜", "90分", "8回", "山田 花子", "R7.4.1〜", "右踵部の発赤確認必須", ""],
        ["身体介護", "通院介助（受診同行・院内移動介助）", "受診日（月1回）", "2時間", "1回", "担当ヘルパー", "R7.4.1〜", "受診先：〇〇クリニック", ""],
        ["家事援助", "調理（夕食）・洗濯・掃除（トイレ・居室）", "月・水 15:00〜", "60分", "8回", "担当ヘルパー", "R7.4.1〜", "食事制限（塩分6g/日）", ""],
    ]
    for i, row_data in enumerate(examples, start=22):
        for j, val in enumerate(row_data):
            ec(ws, f"{get_column_letter(j+1)}{i}", val)
        sr(ws, i, 38)
    for row in range(25, 30):
        for col in "ABCDEFGHI": bc(ws, f"{col}{row}")
        sr(ws, row, 35)

    # 緊急時対応・リスク管理
    section_hdr(ws, "【リスク管理・緊急時対応計画】", "A30:I30", 30)
    ws.merge_cells("A31:I31")
    ws["A31"].value = "緊急連絡先：①家族（　　）TEL：　　　　②主治医（　　　クリニック）TEL：　　　　③事業所（　　）TEL：　　　　\n予測されるリスク：転倒（　）　服薬誤薬（　）　誤嚥（　）　その他（　　　　　　　）\n緊急時の対応手順：本人の安全確保→サービス提供責任者へ報告→家族・主治医へ連絡→状況に応じて救急要請"
    ws["A31"].font = bfont(); ws["A31"].alignment = Alignment(horizontal="left", vertical="top", wrap_text=True); ws["A31"].border = bd(); sr(ws, 31, 70)

    # 同意欄
    section_hdr(ws, "【本人・家族の同意欄】（法令必須）", "A32:I32", 32)
    ws.merge_cells("A33:I33")
    ws["A33"].value = "上記の個別支援計画について、サービス提供責任者より説明を受け、内容を確認し、同意しました。\n\n本人署名：　　　　　　　　　　（　　年　月　日）　　保護者・代理人署名：　　　　　　　　続柄：　　　　（　　年　月　日）\n\nサービス提供責任者署名：　　　　　　　　（説明日：　年　月　日）　　管理者確認：　　　　　（　年　月　日）"
    ws["A33"].font = bfont(); ws["A33"].alignment = Alignment(horizontal="left", vertical="top", wrap_text=True); ws["A33"].border = bd(); sr(ws, 33, 80)

    for col, w in zip("ABCDEFGHI", [16,30,12,9,8,12,10,16,8]): ws.column_dimensions[col].width = w
    save(wb, "03_就労支援記録様式", "13_個別支援計画書（居宅介護版）.xlsx")


def make_monitoring_v2():
    """モニタリング記録（就労継続支援）— 法令必須項目完全版"""
    wb = Workbook(); ws = wb.active; ws.title = "モニタリング記録"
    title_row(ws, "モニタリング記録（就労継続支援A型・B型）", "A1:H1")
    note_row(ws, "【実施頻度】原則6ヶ月ごと。状態変化時・サービス内容変更時は随時実施。利用者宅または事業所での面談が望ましい。", "A2:H2", 2)

    section_hdr(ws, "【基本情報】", "A3:H3", 3)
    basic = [
        ("利用者氏名", "（様）", "実施日", "　　年　月　日"),
        ("サービス種別", "就労継続支援　A型 / B型", "実施者（サービス管理責任者）", ""),
        ("個別支援計画の期間", "　年　月〜　年　月", "前回モニタリング日", "　　年　月　日"),
        ("参加者", "本人 ・ 家族（　　）・ 担当職員 ・ 相談支援専門員 ・ その他（　　）", "次回モニタリング予定", "　　年　月　日頃"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(basic, start=4):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:E{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"F{i}:G{i}"); ws[f"F{i}"].value = l2
        ws[f"H{i}"].value = v2
        for ref in [f"A{i}", f"F{i}"]:
            ws[ref].font = bfont(bold=True); ws[ref].fill = sf()
            ws[ref].alignment = Alignment(horizontal="right", vertical="center"); ws[ref].border = bd()
        for ref in [f"C{i}", f"H{i}"]:
            ws[ref].font = bfont(); ws[ref].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True); ws[ref].border = bd()
        sr(ws, i, 22)

    # 目標達成状況テーブル
    section_hdr(ws, "【個別支援計画 目標の達成状況】", "A8:H8", 8)
    hrow(ws, 9, [("A","目標（計画書より転記）",26),("B","短期/長期",7),
                 ("C","達成度評価",12),("D","達成度の根拠・状況詳細",24),
                 ("E","本人の感想・意向",16),("F","今後の方針",14),("G","計画変更要否",8),("H","備考",8)])
    ec(ws, "C10", "達成/ほぼ達成/未達成")
    for row in range(9, 14):
        for col in "ABCDEFGH": bc(ws, f"{col}{row}")
        sr(ws, row, 40)

    # 出席・作業状況
    section_hdr(ws, "【出席・就労状況（モニタリング期間中）】", "A14:H14", 14)
    ws.merge_cells("A15:H15")
    ws["A15"].value = "出席日数：　日/　日（出席率：　　%）　欠席日数：　日　早退：　回　遅刻：　回\n主な作業内容：　　　　　　　作業能力の変化：向上 / 不変 / 低下（理由：　　　）\n工賃実績（B型）：当月　　　　円　年度累計平均　　　　円"
    ws["A15"].font = bfont(); ws["A15"].alignment = Alignment(horizontal="left", vertical="top", wrap_text=True); ws["A15"].border = bd(); sr(ws, 15, 60)

    # 健康・生活状況
    sections = [
        (16, "【健康・生活状況の変化】", 55,
         "医療機関受診状況：　　　　　　服薬状況：適切/課題あり（　　　）\n体調・精神状態：安定/不安定（内容：　　　　　　）　睡眠：良好/不良\n生活環境の変化：なし / あり（　　　　　　　　）"),
        (18, "【本人の意向・満足度・自己評価】", 60,
         "サービスへの満足度：大変満足 / 満足 / やや不満 / 不満\n「どんなことが良かった・困ったか」（本人の言葉で記録）：\n今後やってみたいこと・要望："),
        (20, "【家族・支援者からの意見・要望】", 45, ""),
        (22, "【支援上の課題・改善事項】", 55, ""),
        (24, "【今後の支援方針・計画変更の要否】", 50,
         "計画変更：必要（変更内容：　　　　　　　　） / 不要（現計画を継続）\n次期の主な支援課題・目標："),
    ]
    for rn, title, h, ex in sections:
        ws.merge_cells(f"A{rn}:H{rn}"); ws[f"A{rn}"].value = title
        ws[f"A{rn}"].font = hfont(sz=9); ws[f"A{rn}"].fill = hf()
        ws[f"A{rn}"].alignment = Alignment(horizontal="left", vertical="center"); ws[f"A{rn}"].border = bd(); sr(ws, rn, 20)
        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c = ws[f"A{rn+1}"]
        c.value = ex; c.font = efont() if ex else bfont()
        c.fill = gf() if ex else PatternFill()
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True); c.border = bd(); sr(ws, rn+1, h)

    # 同意・署名
    r = 26
    ws.merge_cells(f"A{r}:H{r}")
    ws[f"A{r}"].value = "モニタリング結果について説明を受け、確認しました。\n\n本人署名：　　　　　　　　（　年　月　日）　保護者・代理人：　　　　　　　続柄：　　（　年　月　日）\n\nサービス管理責任者署名：　　　　　　　　（実施日：　年　月　日）　管理者確認：　　　　　　（　年　月　日）"
    ws[f"A{r}"].font = bfont(); ws[f"A{r}"].alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    ws[f"A{r}"].border = bd(); sr(ws, r, 80)

    for col, w in zip("ABCDEFGH", [26,7,12,24,16,14,8,8]): ws.column_dimensions[col].width = w
    save(wb, "03_就労支援記録様式", "03_モニタリング記録（就労）.xlsx")


def make_assessment_v2():
    """アセスメント票（A型・B型・居宅介護・相談支援共通） — ICF準拠・強化版"""
    wb = Workbook(); ws = wb.active; ws.title = "アセスメント票"
    title_row(ws, "アセスメントシート（就労支援・居宅介護・相談支援共通版）", "A1:H1")
    note_row(ws, "【実施タイミング】サービス開始前・個別支援計画見直し時・状況変化時。本人・家族等の参加のもと実施することが望ましい。ICF（国際生活機能分類）の視点を活用。", "A2:H2", 2)

    section_hdr(ws, "【基本情報・実施状況】", "A3:H3", 3)
    basic = [
        ("利用者氏名", "（様）", "実施日", "　　年　月　日"),
        ("生年月日・年齢", "　　年　月　日生（　歳）", "実施者", ""),
        ("利用するサービス種別", "", "参加者", "本人 / 家族（　）/ 相談支援専門員 / その他"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(basic, start=4):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value = l2
        ws.merge_cells(f"G{i}:H{i}"); ws[f"G{i}"].value = v2
        for ref in [f"A{i}", f"E{i}"]:
            ws[ref].font = bfont(bold=True); ws[ref].fill = sf()
            ws[ref].alignment = Alignment(horizontal="right", vertical="center"); ws[ref].border = bd()
        for ref in [f"C{i}", f"G{i}"]:
            ws[ref].font = bfont(); ws[ref].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True); ws[ref].border = bd()
        sr(ws, i, 22)

    assess_sections = [
        (7, "【本人の意向・夢・将来の希望】（強みベースで記録）",
         "「どんな生活をしたいか」「何がやりたいか」「どんな仕事・活動をしてみたいか」を本人の言葉で記録。\n例）「自分で稼ぎたい」「料理を覚えたい」「一人暮らしがしたい」", 60),
        (9, "【障害の状況・特性】（強みも含めて記録）",
         "障害の種別・程度・診断名。障害の特性（感覚過敏・パニック時の対応・得意な作業）。コミュニケーションの方法（話す/書く/絵/AAC等）。", 65),
        (11, "【健康・医療の状況】",
         "主傷病名・副傷病名・既往歴。通院状況・頻度・主治医名。服薬状況（種類・管理方法）。体調の波・悪化のサイン・てんかん発作の有無。", 65),
        (13, "【日常生活・ADLの状況】",
         "食事（自立/介助要）　排泄（自立/介助要）　入浴（自立/介助要）　着替え（自立/介助要）　移動（自立/車椅子/介助要）\n服薬管理（自立/支援要）　金銭管理（自立/支援要）　調理（自立/支援要）　掃除・洗濯（自立/支援要）", 70),
        (15, "【就労・活動能力の状況】（就労系サービスのみ）",
         "作業の持続時間（　h）　手先の器用さ（良好/普通/課題あり）　指示理解力（口頭/書面/実演で理解）\n集中力・体力・協調性・自己申告（体調悪化を自分で申し出られるか）", 65),
        (17, "【社会参加・日中活動の状況】",
         "現在利用中のサービス（種別・事業所名・利用頻度）。地域・コミュニティとの関係。余暇活動・趣味。外出頻度・交通機関の利用。", 60),
        (19, "【家族・住環境・支援体制の状況】",
         "同居家族・キーパーソン（氏名・続柄・連絡先・介護力）。住居形態（自宅/GH/施設）。バリアフリー状況。家族の介護力・支援意向。", 60),
        (21, "【経済状況・制度利用状況】",
         "収入源（障害年金/就労賃金・工賃/生活保護/その他）。活用中の制度（障害者手帳・自立支援医療・成年後見等）。経済的な困難・課題。", 55),
        (23, "【本人・家族のニーズ（解決すべき課題）まとめ】",
         "上記のアセスメントを踏まえて、支援の優先課題を箇条書きで整理する。\n① 　② 　③ 　④", 60),
        (25, "【強みと活用できるリソース】（ストレングスベース）",
         "例）手先が器用で封入作業が得意。音楽が大好きで気持ちが安定する。家族のサポートが手厚い。主治医との関係が良好。", 50),
        (27, "【アセスメントから見えた支援の方向性】",
         "個別支援計画の目標・支援内容の方向性について整理する。", 55),
    ]
    for rn, title, ex, h in assess_sections:
        section_hdr(ws, title, f"A{rn}:H{rn}", rn)
        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c = ws[f"A{rn+1}"]
        c.value = ex; c.font = efont(); c.fill = gf()
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True); c.border = bd(); sr(ws, rn+1, h)

    # 署名
    r = 29
    ws.merge_cells(f"A{r}:H{r}")
    ws[f"A{r}"].value = "アセスメント実施者署名：　　　　　　　　　　　（日付：　年　月　日）\nサービス管理責任者・確認者：　　　　　　　　（日付：　年　月　日）"
    ws[f"A{r}"].font = bfont(); ws[f"A{r}"].alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    ws[f"A{r}"].border = bd(); sr(ws, r, 45)

    for col, w in zip("ABCDEFGH", [18,14,14,14,14,14,14,14]): ws.column_dimensions[col].width = w
    save(wb, "03_就労支援記録様式", "10_アセスメント票.xlsx")


# ======================================================================
# 06_事故苦情ヒヤリ様式
# ======================================================================

def make_jiko_v2():
    """事故報告書 — 5W1H・M-SHEL分析・行政報告・再発防止完全版"""
    wb = Workbook(); ws = wb.active; ws.title = "事故報告書"
    ws.merge_cells("A1:H1"); ws["A1"].value = "事故報告書"
    ws["A1"].font = Font(name="Meiryo UI", size=14, bold=True, color=WHITE)
    ws["A1"].fill = PatternFill("solid", fgColor="FFDC2626")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center"); sr(ws, 1, 28)
    note_row(ws, "【速報義務】重大事故（死亡・骨折・誤嚥等）は24時間以内に行政へ速報。その後速やかに詳細報告書を提出。【保管】5年間", "A2:H2", 2)

    # 基本情報
    section_hdr(ws, "【事故基本情報】", "A3:H3", 3)
    basic = [
        ("発生日時", "　　年　月　日（　曜日）　　時　　分頃", "事業所名", ""),
        ("発生場所", "利用者宅 / 事業所内 / 移動中 / 通院先 / その他（　　）", "担当者氏名・職種", ""),
        ("当事者氏名（利用者）", "", "年齢・性別", "　　歳　男 / 女"),
        ("障害種別・支援区分", "", "当事者の状態（平時）", ""),
        ("事故の種類", "転倒 / 転落 / 誤薬 / 誤嚥・窒息 / 行方不明 / 暴力 / 熱傷 / 施設内感染 / 車両事故 / その他（　　）", "重症度", "軽傷 / 中等症 / 重症 / 死亡"),
        ("目撃者の有無・氏名", "なし / あり（　　　　　　職種：　　　）", "救急搬送", "なし / あり（搬送先：　　　　　　　）"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(basic, start=4):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value = l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value = v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value = l2
        ws.merge_cells(f"G{i}:H{i}"); ws[f"G{i}"].value = v2
        for ref in [f"A{i}", f"E{i}"]:
            ws[ref].font = bfont(bold=True); ws[ref].fill = sf()
            ws[ref].alignment = Alignment(horizontal="right", vertical="center"); ws[ref].border = bd()
        for ref in [f"C{i}", f"G{i}"]:
            ws[ref].font = bfont(); ws[ref].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True); ws[ref].border = bd()
        sr(ws, i, 22)

    # 事故の経緯・対応（時系列）
    text_sections = [
        (10, "【事故の経緯・状況（5W1H：いつ・どこで・誰が・何を・なぜ・どのように）】", 100,
         "例）〇時〇分、利用者宅浴室にて入浴介助中、担当ヘルパー山田が目を離した際に転倒した。右膝に擦り傷（約3cm）。本人は「足が滑った」と話している。浴室マットが端に寄っていた。"),
        (12, "【緊急対応の内容（時系列で記録）】", 90,
         "例）〇時〇分：転倒確認・安全確認・出血部位の止血処置\n〇時〇分：管理者・サービス提供責任者へ電話報告\n〇時〇分：主治医（〇〇クリニック）へ連絡→自宅様子見の指示\n〇時〇分：家族（長男）へ電話報告・説明\n翌日：経過確認訪問実施"),
        (14, "【受診状況・診断結果】", 45,
         "受診：なし / あり（受診日：　年　月　日　医療機関：　　　　　診断：　　　　　　治療内容：　　　　　）"),
        (16, "【原因分析（なぜ起きたか：人・設備・環境・手順・ルール・体制の6観点で）】", 90,
         "①人的要因：担当者の注意不足 / 技術不足 / 疲労・健康問題\n②設備要件：浴室マットの劣化・ズレ / 手すりなし\n③環境要因：床が滑りやすい / 照明不十分\n④手順要因：見守り基準が曖昧\n⑤ルール要因：ヒヤリハット共有不足\n⑥体制要因：人手不足による急ぎの対応"),
        (18, "【再発防止策（具体的な対策・実施担当者・実施期限を明記）】", 90,
         "①滑り止めマットを新品に交換（担当：〇〇　期限：〇月〇日）\n②浴室内では利用者から目を離さないよう手順書を改訂（担当：管理者　期限：〇月〇日）\n③スタッフ全員への朝礼での周知・注意喚起（担当：管理者　期限：〇月〇日）\n④1ヶ月後に再発防止策の効果確認を実施"),
        (20, "【行政への報告状況】", 45,
         "速報：要 / 不要（報告先：　　　　　　　速報日：　年　月　日）\n詳細報告：提出予定日：　年　月　日　受理確認："),
        (22, "【保険請求・補償対応】", 35, "賠償責任保険：申請要 / 不要（保険会社：　　　　　　　請求番号：　　　　）"),
        (23, "【フォローアップ実施記録】", 45,
         "①利用者・家族への経過報告日：　年　月　日（内容：　　　）\n②再発防止策の実施確認日：　年　月　日（内容：　　　）\n③全職員への周知・研修実施日：　年　月　日"),
    ]
    for rn, title, h, ex in text_sections:
        ws.merge_cells(f"A{rn}:H{rn}"); ws[f"A{rn}"].value = title
        ws[f"A{rn}"].font = hfont(sz=9); ws[f"A{rn}"].fill = hf()
        ws[f"A{rn}"].alignment = Alignment(horizontal="left", vertical="center"); ws[f"A{rn}"].border = bd(); sr(ws, rn, 20)
        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c = ws[f"A{rn+1}"]
        c.value = ex; c.font = efont() if ex else bfont()
        c.fill = gf() if ex else PatternFill()
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True); c.border = bd(); sr(ws, rn+1, h)

    # 署名
    sign_row = 25
    ws.merge_cells(f"A{sign_row}:D{sign_row}")
    ws[f"A{sign_row}"].value = "報告者氏名・署名：　　　　　　　　　（報告日：　年　月　日）"
    ws.merge_cells(f"E{sign_row}:H{sign_row}")
    ws[f"E{sign_row}"].value = "管理者確認・署名：　　　　　　　　　（確認日：　年　月　日）"
    for ref in [f"A{sign_row}", f"E{sign_row}"]:
        ws[ref].font = bfont(); ws[ref].border = bd()
        ws[ref].alignment = Alignment(horizontal="left", vertical="center")
    sr(ws, sign_row, 30)

    for col, w in zip("ABCDEFGH", [14]*8): ws.column_dimensions[col].width = w
    save(wb, "06_事故苦情ヒヤリ様式", "01_事故報告書.xlsx")


# ======================================================================
# 07_運営規程ひな形（Word版強化）
# ======================================================================

def make_uneikitei_houmon_v2():
    """運営規程ひな形（訪問看護）— 法令必須条項完全版 2024年度対応"""
    doc = Document()
    doc.styles["Normal"].font.name = "Meiryo UI"
    doc.styles["Normal"].font.size = Pt(10)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3); sec.right_margin = Cm(2.5)

    h = doc.add_heading("訪問看護ステーション 運営規程（ひな形）", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x0E, 0x74, 0x90)
    h.runs[0].font.size = Pt(14)

    doc.add_paragraph("※ このひな形は参考例です。実際の申請では事業所の実態に合わせて修正し、京都府・京都市の最新様式・審査を経て使用してください。")
    doc.add_paragraph("【法令根拠】健康保険法第88条・障害者総合支援法・指定訪問看護の事業の人員及び運営に関する基準（厚生省令第80号）")

    chapters = [
        ("第1条（目的）",
         "この運営規程は、○○訪問看護ステーション（以下「当ステーション」という。）が実施する訪問看護事業の適正な運営を確保するために必要な人員及び管理運営に関する事項を定め、利用者が心身の機能の維持・回復を図り、その居宅において自立した日常生活を営むことができるよう、適切な訪問看護を提供することを目的とする。"),
        ("第2条（運営の方針）",
         "当ステーションが実施する訪問看護事業の運営は、以下の方針により行う。\n"
         "①利用者の意思及び人格を尊重し、常にその者の立場に立ったサービスの提供に努める。\n"
         "②地域の医療機関・保健・福祉機関との連携を密にして、総合的なサービス提供に努める。\n"
         "③利用者の人権の擁護、虐待・ハラスメントの防止のため、研修の実施等必要な措置を講じる。\n"
         "④個人情報の保護に努め、業務上知り得た情報を適切に管理する。\n"
         "⑤医師の指示に基づく訪問看護の提供を通じて、在宅での療養生活の質の向上を図る。\n"
         "⑥事業継続計画（BCP）を策定し、災害・感染症等の緊急時においても利用者への支援を継続する。"),
        ("第3条（事業所の名称及び所在地）",
         "事業所の名称：○○訪問看護ステーション\n"
         "事業所の所在地：京都府○○市○○町○○番地\n"
         "電話番号：○○○-○○○○-○○○○　FAX：○○○-○○○○-○○○○\n"
         "指定番号：○○○○○○○○○○（指定日：令和　年　月　日）"),
        ("第4条（従業者の職種、員数及び職務内容）",
         "当ステーションに勤務する職員の職種、員数及び職務内容は次のとおりとする。\n"
         "（１）管理者：1名（常勤・専任）\n"
         "　　職務：訪問看護ステーション全体の管理・業務実施状況の把握・業務改善・行政への届出管理\n"
         "（２）看護師・准看護師：常勤換算2.5名以上\n"
         "　　職務：主治医の指示に基づく訪問看護の実施・訪問看護計画書及び報告書の作成・利用者家族への指導\n"
         "（３）理学療法士・作業療法士・言語聴覚士：必要に応じて配置\n"
         "　　職務：主治医の指示に基づくリハビリテーションの実施\n"
         "（４）事務員：必要に応じて配置\n"
         "　　職務：診療報酬請求事務・書類管理・電話対応"),
        ("第5条（営業日及び営業時間）",
         "営業日：月曜日から金曜日（祝日・年末年始〔12月29日から1月3日〕を除く）\n"
         "営業時間：午前8時30分から午後5時30分まで\n"
         "訪問看護の提供時間：上記営業時間中。\n"
         "24時間対応体制：当ステーションは24時間対応体制加算を（算定している / 算定していない）。\n"
         "　　　　　　　　算定している場合：担当看護師の連絡先を利用者に提供し、緊急時の連絡・相談・訪問に対応する。"),
        ("第6条（訪問看護の内容及び利用料等）",
         "（１）提供するサービスの内容\n"
         "・病状の観察及び判断（バイタルサイン測定・全身状態の観察）\n"
         "・医師の指示に基づく医療処置の補助（点滴管理・褥瘡処置・吸引・経管栄養等）\n"
         "・リハビリテーション（理学療法士・作業療法士・言語聴覚士による場合）\n"
         "・服薬管理・服薬指導\n"
         "・清潔ケア（清拭・入浴介助等）\n"
         "・在宅酸素・人工呼吸器等の医療機器管理\n"
         "・認知症のケア\n"
         "・精神科訪問看護（精神科を標榜する保険医の精神科訪問看護指示書に基づく場合）\n"
         "・ターミナルケア（在宅看取りへの対応）\n"
         "・家族等への介護指導・相談支援\n\n"
         "（２）利用料\n"
         "医療保険（健康保険・国民健康保険・後期高齢者医療）の給付対象となる訪問看護は、訪問看護療養費として保険請求する。利用者負担割合は保険の種類・年齢・所得等により異なる（原則1〜3割）。\n"
         "自立支援医療（精神通院）の対象者は、原則1割（月額上限あり）。\n"
         "上記以外に係る費用（交通費等の実費）については、別途実費を徴収することがある。"),
        ("第7条（苦情処理）",
         "苦情受付窓口：管理者（電話：○○○-○○○○-○○○○　受付時間：営業日の9時〜17時）\n"
         "受付方法：面接・電話・書面・メール\n"
         "苦情処理の手順：受付→記録→原因調査→対応・改善→報告→記録の保管（5年間）\n"
         "第三者委員：○○様（連絡先：○○）\n"
         "外部の苦情処理機関：京都府 障害者支援課（075-414-4600）/ 国民健康保険団体連合会"),
        ("第8条（緊急時等における対応方法）",
         "訪問中に利用者の病状が急変した場合は、速やかに主治医（指示を行った医師）に連絡をとり、指示を受ける。必要に応じて救急搬送の手配を行うとともに、管理者に報告し、利用者家族への連絡も速やかに行う。\n"
         "緊急時の連絡体制：担当看護師→管理者→主治医→家族→（必要に応じて）救急\n"
         "重大事故発生時は、24時間以内に行政（京都府障害者支援課または京都市障害保健福祉推進室）へ速報し、その後詳細報告書を提出する。"),
        ("第9条（虐待の防止のための措置）",
         "当ステーションは利用者への虐待防止のため、以下の措置を講ずる。\n"
         "①虐待防止責任者の設置（管理者が担当）\n"
         "②虐待防止に関する研修の年1回以上の実施・記録の保管\n"
         "③苦情・通報窓口の整備・周知\n"
         "④虐待発見時の速やかな通報（市区町村への通報義務）\n"
         "⑤身体拘束の禁止（緊急やむを得ない場合を除き一切禁止）\n"
         "⑥ハラスメント防止方針の策定・周知"),
        ("第10条（衛生管理・感染症対策）",
         "当ステーションは感染症の予防及びまん延の防止のため、以下の措置を講じる。\n"
         "①感染症対策指針の策定・見直し（年1回以上）\n"
         "②感染症対策委員会の設置・定期開催（責任者：管理者）\n"
         "③感染症対策に関する研修の年2回以上の実施\n"
         "④感染症対策訓練の年2回以上の実施\n"
         "⑤標準予防策（手指消毒・PPE使用等）の徹底\n"
         "⑥訪問前後の手指衛生・器材の洗浄・消毒の徹底"),
        ("第11条（事業継続計画（BCP）の策定）",
         "当ステーションは、感染症・自然災害等の非常時においても、利用者への訪問看護の提供を継続するため、事業継続計画（BCP）を策定し、年1回以上見直しを行う。\n"
         "BCPには、①発生時の初動対応・指揮命令系統、②業務継続の優先順位、③スタッフの安否確認方法、④代替手段の確保（他事業所との連携等）を記載する。\n"
         "BCPに基づく訓練を年2回以上実施し、記録を保管する。"),
        ("第12条（秘密保持等）",
         "従業者は業務上知り得た利用者及びその家族の個人情報の保持に努め、正当な理由なくして第三者に漏洩しない。また、業務を離れた後もこの義務は継続する。\n"
         "採用時に全職員から秘密保持に関する誓約書を徴取する。\n"
         "情報の第三者提供は、利用者の同意書に基づく場合または法令に基づく場合に限る。"),
        ("第13条（記録の整備・保管）",
         "当ステーションは、訪問看護計画書・報告書・記録書等の書類を作成し、利用者ごとに整理・保管する。\n"
         "記録の保管期間：完結の日から5年間（都道府県条例で定める期間以上）\n"
         "電子保存を行う場合は、改ざん防止・バックアップ等の適切な措置を講じる。"),
        ("第14条（研修の実施）",
         "当ステーションは、全職員に対して以下の研修を計画的に実施し、記録を保管する。\n"
         "・虐待防止研修：年1回以上\n"
         "・感染症対策研修：年2回以上\n"
         "・個人情報保護研修：年1回以上\n"
         "・身体拘束禁止研修：年1回以上\n"
         "・BCP（事業継続計画）訓練：年2回以上\n"
         "・その他、資質向上のための研修・外部研修への参加"),
        ("附則",
         "この運営規程は、令和○年○月○日から施行する。\n"
         "改訂履歴：令和○年○月○日改訂（改訂理由：　　　　　　　　　　　）"),
    ]

    for title, content in chapters:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.paragraph_format.space_before = Pt(10)
        doc.add_paragraph(content)

    save_doc(doc, "07_運営規程ひな形", "01_運営規程ひな形_訪問看護.docx")


# ======================================================================
# 実行
# ======================================================================
if __name__ == "__main__":
    print("=== テンプレート大幅強化版 (v2) ===\n")

    print("--- 02_訪問看護記録様式 ---")
    make_kiroku1_facesheet()
    make_kiroku2_keika()
    make_houmon_houkoku_v2()
    make_kyotaku_service_kiroku()   # 新規作成

    print("\n--- 03_就労支援記録様式 ---")
    make_kojin_shien_v2()
    make_kyotaku_keikaku()          # 新規作成（居宅介護版）
    make_monitoring_v2()
    make_assessment_v2()

    print("\n--- 06_事故苦情ヒヤリ様式 ---")
    make_jiko_v2()

    print("\n--- 07_運営規程ひな形 ---")
    make_uneikitei_houmon_v2()

    print("\n=== 完了 ===")
