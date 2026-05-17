"""追加テンプレートバッチ3"""
import os, sys
sys.stdout.reconfigure(encoding='utf-8')
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
TEAL="FF0E7490"; WHITE="FFFFFFFF"; LIGHT="FFCFFAFE"; GRAY="FFF8FAFC"
RED="FFDC2626"; AMBER="FFD97706"

def hf(): return PatternFill("solid", fgColor=TEAL)
def sf(): return PatternFill("solid", fgColor=LIGHT)
def gf(): return PatternFill("solid", fgColor=GRAY)
def rf(): return PatternFill("solid", fgColor=RED)
def hfont(sz=10,col=WHITE): return Font(name="Meiryo UI",size=sz,bold=True,color=col)
def bfont(sz=9,bold=False,col=None):
    if col: return Font(name="Meiryo UI",size=sz,bold=bold,color=col)
    return Font(name="Meiryo UI",size=sz,bold=bold)
def bd():
    s=Side(style="thin")
    return Border(left=s,right=s,top=s,bottom=s)
def sr(ws,r,h): ws.row_dimensions[r].height=h

def hrow(ws,row,cols):
    for letter,text,width in cols:
        c=ws[f"{letter}{row}"]
        c.value=text; c.font=hfont(); c.fill=hf()
        c.alignment=Alignment(horizontal="center",vertical="center",wrap_text=True)
        c.border=bd(); ws.column_dimensions[letter].width=width
    sr(ws,row,30)

def bc(ws,ref,val="",bold=False,ha="left"):
    c=ws[ref]; c.value=val; c.font=bfont(bold=bold)
    c.alignment=Alignment(horizontal=ha,vertical="center",wrap_text=True); c.border=bd()
    return c

def ec(ws,ref,val=""):
    c=ws[ref]; c.value=val
    c.font=Font(name="Meiryo UI",size=9,color="FF6B7280",italic=True)
    c.fill=gf(); c.alignment=Alignment(horizontal="left",vertical="center",wrap_text=True); c.border=bd()

def title_cell(ws,ref,text,row_height=26):
    row = int(ref[1:])
    ws.merge_cells(f"A{row}:H{row}")
    ws[ref].value=text
    ws[ref].font=Font(name="Meiryo UI",size=13,bold=True,color=TEAL)
    ws[ref].alignment=Alignment(horizontal="center",vertical="center")
    sr(ws,row,row_height)

def save(wb,folder,fname):
    p=os.path.join(BASE,folder,fname); wb.save(p)
    print(f"OK: {folder}/{fname}")

def save_doc(doc,folder,fname):
    p=os.path.join(BASE,folder,fname); doc.save(p)
    print(f"OK: {folder}/{fname}")

# ===== 賃金台帳（A型専用） =====
def make_chingin():
    wb=Workbook(); ws=wb.active; ws.title="賃金台帳A型"

    ws.merge_cells("A1:J1")
    ws["A1"].value="賃金台帳（就労継続支援A型用）　　　年　　月分"
    ws["A1"].font=Font(name="Meiryo UI",size=13,bold=True,color=TEAL)
    ws["A1"].alignment=Alignment(horizontal="center",vertical="center"); sr(ws,1,26)

    ws.merge_cells("A2:J2")
    ws["A2"].value="【法的根拠】労働基準法第108条：賃金台帳の作成義務。賃金支払い毎に記入し3年間保存。最低賃金（毎年10月改定）以上の賃金支払いが必須。"
    ws["A2"].font=Font(name="Meiryo UI",size=8,color=RED,bold=True)
    ws["A2"].alignment=Alignment(horizontal="left",vertical="center"); sr(ws,2,20)

    # 最低賃金欄
    ws.merge_cells("A3:E3"); ws["A3"].value="京都府最低賃金（2024年10月〜）：時間額 1,058円　【毎年10月確認必須】"
    ws["A3"].font=Font(name="Meiryo UI",size=9,bold=True,color=RED)
    ws["A3"].fill=PatternFill("solid",fgColor="FFFEE2E2"); ws["A3"].alignment=Alignment(horizontal="left",vertical="center"); ws["A3"].border=bd()
    ws.merge_cells("F3:J3"); ws["F3"].value="事業所名：　　　　　　　　　　　"
    ws["F3"].font=bfont(); ws["F3"].alignment=Alignment(horizontal="left",vertical="center"); ws["F3"].border=bd(); sr(ws,3,22)

    cols=[("A","利用者氏名",14),("B","雇用形態",10),("C","基本時給\n（円）",8),
          ("D","労働日数",7),("E","総労働\n時間（h）",8),("F","基本給\n（円）",10),
          ("G","控除額\n（円）",9),("H","差引支給額\n（円）",11),("I","支払日",10),("J","受領確認\n（署名）",12)]
    hrow(ws,4,cols)

    ex=[["山田 太郎","雇用（週20h）",1100,18,72,79200,3500,75700,"2026/5/25","山田"],
        ["佐藤 花子","雇用（週15h）",1060,15,56,59360,0,59360,"2026/5/25","佐藤"]]
    for i,row in enumerate(ex,start=5):
        for j,v in enumerate(row): ec(ws,f"{get_column_letter(j+1)}{i}",str(v))
        sr(ws,i,22)

    for row in range(7,27):
        for col in range(1,11): bc(ws,f"{get_column_letter(col)}{row}",ha="right" if col in [3,4,5,6,7,8] else "left")
        sr(ws,row,22)

    total=27
    ws.merge_cells(f"A{total}:E{total}"); ws[f"A{total}"].value="合計"
    ws[f"A{total}"].font=bfont(bold=True); ws[f"A{total}"].fill=sf()
    ws[f"A{total}"].alignment=Alignment(horizontal="right",vertical="center"); ws[f"A{total}"].border=bd()
    for col in "FGHIJ": bc(ws,f"{col}{total}",ha="right"); sr(ws,total,22)

    ws.merge_cells(f"A{total+1}:J{total+1}")
    ws[f"A{total+1}"].value="【最低賃金チェック】全利用者の基本時給が最低賃金（1,058円）以上であること。最低賃金改定時（毎年10月）は即時に契約内容を更新すること。"
    ws[f"A{total+1}"].font=Font(name="Meiryo UI",size=9,color=RED,bold=True)
    ws[f"A{total+1}"].alignment=Alignment(horizontal="left",vertical="center"); sr(ws,total+1,20)

    save(wb,"03_就労支援記録様式","11_賃金台帳（A型）.xlsx")


# ===== 業務日誌（管理者用）=====
def make_gyomu():
    wb=Workbook(); ws=wb.active; ws.title="業務日誌"

    ws.merge_cells("A1:G1")
    ws["A1"].value="管理者業務日誌　　　　年　　月分"
    ws["A1"].font=Font(name="Meiryo UI",size=13,bold=True,color=TEAL)
    ws["A1"].alignment=Alignment(horizontal="center",vertical="center"); sr(ws,1,26)

    ws.merge_cells("A2:G2")
    ws["A2"].value="管理者氏名：　　　　　　　　　　事業所名：　　　　　　　　　　サービス種別：　　　　　　　"
    ws["A2"].font=bfont(); ws["A2"].alignment=Alignment(horizontal="left",vertical="center"); ws["A2"].border=bd(); sr(ws,2,20)

    cols=[("A","日付",8),("B","出勤者・欠勤者",16),("C","利用者の状況\n（人数・特記）",22),
          ("D","外部連絡・会議・訪問",20),("E","事務処理・請求作業",18),("F","特記事項・課題",20),("G","確認",8)]
    hrow(ws,3,cols)

    ex_rows=[
        ["5/1(木)","出勤：山田・鈴木・田中\n欠勤：佐藤（有給）","利用者18名出席\n山本さん体調不良で早退",
         "主治医（田中クリニック）へ経過報告\n相談支援専門員・鈴木様と電話",
         "4月分請求データ最終確認\n国保連に電子請求送信","——","✓"],
        ["5/2(金)","出勤：全員5名","利用者20名出席","——",
         "処遇改善加算計画書の最終確認・押印","来週の研修準備（感染症対策）","✓"],
    ]
    for i,row in enumerate(ex_rows,start=4):
        for j,v in enumerate(row): ec(ws,f"{get_column_letter(j+1)}{i}",v)
        sr(ws,i,55)

    for row in range(6,36):
        for col in range(1,8): bc(ws,f"{get_column_letter(col)}{row}")
        sr(ws,row,55)

    for col,w in zip("ABCDEFG",[8,16,22,20,18,20,8]): ws.column_dimensions[col].width=w
    save(wb,"06_事故苦情ヒヤリ様式","10_管理者業務日誌.xlsx")


# ===== 精神科訪問看護計画書（参考様式）=====
def make_seishin_keikaku():
    wb=Workbook(); ws=wb.active; ws.title="精神科訪看計画書"

    ws.merge_cells("A1:H1")
    ws["A1"].value="精神科訪問看護計画書（参考様式）　別紙様式3準拠"
    ws["A1"].font=Font(name="Meiryo UI",size=13,bold=True,color=TEAL)
    ws["A1"].alignment=Alignment(horizontal="center",vertical="center"); sr(ws,1,26)

    ws.merge_cells("A2:H2")
    ws["A2"].value="※ 精神疾患を有する利用者への訪問看護は、一般の訪問看護計画書（別紙様式1）の代わりにこの精神科用様式（別紙様式3）を使用します。"
    ws["A2"].font=Font(name="Meiryo UI",size=8,color=TEAL,bold=True)
    ws["A2"].alignment=Alignment(horizontal="left",vertical="center"); sr(ws,2,20)

    info=[
        ("利用者氏名","（様）","生年月日","　　年　月　日（　歳）"),
        ("訪問看護ステーション名","","計画作成日","　　年　月　日"),
        ("主治医氏名・医療機関","","計画期間","　　年　月〜　　年　月"),
        ("精神疾患名（主診断）","","自立支援医療受給","あり（受給者証番号：　　　　）/ なし"),
    ]
    for i,(l1,v1,l2,v2) in enumerate(info,start=3):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value=l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value=v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value=l2
        ws.merge_cells(f"G{i}:H{i}"); ws[f"G{i}"].value=v2
        for cell in [f"A{i}",f"E{i}"]:
            ws[cell].font=bfont(bold=True); ws[cell].fill=sf()
            ws[cell].alignment=Alignment(horizontal="right",vertical="center"); ws[cell].border=bd()
        for cell in [f"C{i}",f"G{i}"]:
            ws[cell].font=bfont(); ws[cell].alignment=Alignment(horizontal="left",vertical="center",wrap_text=True); ws[cell].border=bd()
        sr(ws,i,22)

    sections=[
        (7,"【現在の病状・精神状態（GAF/GAS等の評価含む）】",60,
         "例）統合失調症。陽性症状は落ち着いているが、陰性症状（意欲低下・社会的引きこもり）が中心。GAF：50。"),
        (9,"【看護の目標（精神症状・生活・社会参加）】",50,
         "短期：服薬を自己管理し、週3回の就労継続支援への通所を維持できる。\n長期：地域で安定した生活を送りながら社会参加の場を広げる。"),
        (11,"【問題点・訪問看護の内容】",120,
         "①服薬管理：抗精神病薬の自己管理支援・副作用観察（錐体外路症状等）\n②精神状態の観察：幻覚・妄想・自傷他害リスク・希死念慮の確認\n③生活支援：生活リズム・清潔管理・食事管理の指導\n④コミュニケーション支援：傾聴・感情表出の支援\n⑤危機介入計画：悪化時の対応フロー（主治医・家族・行政への連絡順序）\n⑥服薬拒否時の対応：主治医への報告・家族との連携"),
        (14,"【緊急時対応計画（クライシスプラン）】",60,
         "①本人サイン（早期警告サイン）：　　　　\n②対応者：本人→訪問看護師→主治医→家族→救急\n③入院先：○○病院（TEL：○○）"),
        (16,"【家族・関係機関との連携】",50,
         "家族：　　（続柄）TEL：○○\n相談支援専門員：○○ TEL：○○\n主治医：○○クリニック TEL：○○"),
        (18,"【本人・家族の同意】",30,
         "上記計画について説明を受け、同意します。\n\n本人署名：　　　　　　日付：　　年　月　日\n家族署名：　　　　　　続柄："),
    ]
    for rn,title,h,example in sections:
        ws.merge_cells(f"A{rn}:H{rn}"); ws[f"A{rn}"].value=title
        ws[f"A{rn}"].font=hfont(sz=9); ws[f"A{rn}"].fill=hf()
        ws[f"A{rn}"].alignment=Alignment(horizontal="left",vertical="center"); ws[f"A{rn}"].border=bd(); sr(ws,rn,20)
        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c=ws[f"A{rn+1}"]
        c.value=example; c.font=Font(name="Meiryo UI",size=9,color="FF6B7280",italic=True)
        c.fill=gf(); c.alignment=Alignment(horizontal="left",vertical="top",wrap_text=True); c.border=bd(); sr(ws,rn+1,h)

    for col,w in zip("ABCDEFGH",[14,14,14,14,14,14,14,14]): ws.column_dimensions[col].width=w
    save(wb,"02_訪問看護記録様式","09_精神科訪問看護計画書（別紙様式3準拠）.xlsx")


# ===== 精神科訪問看護報告書 =====
def make_seishin_houkoku():
    wb=Workbook(); ws=wb.active; ws.title="精神科訪看報告書"

    ws.merge_cells("A1:H1")
    ws["A1"].value="精神科訪問看護報告書（参考様式）　別紙様式4準拠"
    ws["A1"].font=Font(name="Meiryo UI",size=13,bold=True,color=TEAL)
    ws["A1"].alignment=Alignment(horizontal="center",vertical="center"); sr(ws,1,26)

    info=[
        ("利用者氏名","（様）","報告対象月","　　年　　月分"),
        ("主治医氏名・医療機関","","訪問看護ステーション",""),
        ("精神疾患名","","担当看護師",""),
        ("訪問回数（当月）","　　回（うち緊急：　　回）","指示書有効期限","　　年　月　日まで"),
    ]
    for i,(l1,v1,l2,v2) in enumerate(info,start=2):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value=l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value=v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value=l2
        ws.merge_cells(f"G{i}:H{i}"); ws[f"G{i}"].value=v2
        for cell in [f"A{i}",f"E{i}"]:
            ws[cell].font=bfont(bold=True); ws[cell].fill=sf()
            ws[cell].alignment=Alignment(horizontal="right",vertical="center"); ws[cell].border=bd()
        for cell in [f"C{i}",f"G{i}"]:
            ws[cell].font=bfont(); ws[cell].alignment=Alignment(horizontal="left",vertical="center",wrap_text=True); ws[cell].border=bd()
        sr(ws,i,22)

    sections=[
        (6,"【精神症状の経過（当月）】",70,
         "陽性症状：幻聴・妄想なし安定。\n陰性症状：意欲低下継続。外出は週3〜4回（就労継続支援への通所）。\nGAF：55（前月比±0）"),
        (8,"【服薬状況・副作用】",50,
         "内服：処方通り自己管理できている。特別な副作用なし。\n残薬確認：問題なし。"),
        (10,"【生活状況（睡眠・食事・清潔・活動）】",60,
         "睡眠：23時就寝・8時起床。概ね安定。\n食事：自炊週3回。\n清潔：週1〜2回の入浴を本人の確認で確認。"),
        (12,"【社会参加・就労・対人関係】",50,
         "就労継続支援B型に週4日通所。同僚との会話も増えてきた。"),
        (14,"【家族・関係機関との連携】",40,
         "家族（母）：週1回電話にて状況共有。安定しているとのこと。\n相談支援専門員：今月モニタリング実施予定。"),
        (16,"【今後の支援方針（主治医へ）】",50,
         "現状維持で良好。来月は外出機会の増加を目標に支援継続予定。"),
    ]
    for rn,title,h,example in sections:
        ws.merge_cells(f"A{rn}:H{rn}"); ws[f"A{rn}"].value=title
        ws[f"A{rn}"].font=hfont(sz=9); ws[f"A{rn}"].fill=hf()
        ws[f"A{rn}"].alignment=Alignment(horizontal="left",vertical="center"); ws[f"A{rn}"].border=bd(); sr(ws,rn,20)
        ws.merge_cells(f"A{rn+1}:H{rn+1}"); c=ws[f"A{rn+1}"]
        c.value=example; c.font=Font(name="Meiryo UI",size=9,color="FF6B7280",italic=True)
        c.fill=gf(); c.alignment=Alignment(horizontal="left",vertical="top",wrap_text=True); c.border=bd(); sr(ws,rn+1,h)

    ws.merge_cells("A18:H18")
    ws["A18"].value="報告日：　　年　月　日　担当看護師署名：　　　　　　　　管理者確認：　　　　　　"
    ws["A18"].font=bfont(); ws["A18"].alignment=Alignment(horizontal="left",vertical="center"); ws["A18"].border=bd(); sr(ws,18,28)

    for col,w in zip("ABCDEFGH",[14,14,14,14,14,14,14,14]): ws.column_dimensions[col].width=w
    save(wb,"02_訪問看護記録様式","10_精神科訪問看護報告書（別紙様式4準拠）.xlsx")


# ===== 退院前カンファレンス記録 =====
def make_taiin_conf():
    wb=Workbook(); ws=wb.active; ws.title="退院前カンファレンス"

    ws.merge_cells("A1:G1")
    ws["A1"].value="退院前カンファレンス記録"
    ws["A1"].font=Font(name="Meiryo UI",size=13,bold=True,color=TEAL)
    ws["A1"].alignment=Alignment(horizontal="center",vertical="center"); sr(ws,1,26)

    ws.merge_cells("A2:G2")
    ws["A2"].value="【加算根拠】退院・退所加算（相談支援）/ 訪問看護情報提供療養費（Ⅰ・Ⅱ）/ 入院時情報連携加算 ← 参加記録が必須"
    ws["A2"].font=Font(name="Meiryo UI",size=9,bold=True,color=TEAL)
    ws["A2"].alignment=Alignment(horizontal="left",vertical="center"); sr(ws,2,20)

    info=[
        ("開催日時","　　年　月　日（　）　　時〜　　時","開催場所",""),
        ("利用者氏名","（様）","入院医療機関",""),
        ("退院予定日","　　年　月　日","担当訪問看護師",""),
    ]
    for i,(l1,v1,l2,v2) in enumerate(info,start=3):
        ws.merge_cells(f"A{i}:B{i}"); ws[f"A{i}"].value=l1
        ws.merge_cells(f"C{i}:D{i}"); ws[f"C{i}"].value=v1
        ws.merge_cells(f"E{i}:F{i}"); ws[f"E{i}"].value=l2
        ws[f"G{i}"].value=v2
        for cell in [f"A{i}",f"E{i}"]:
            ws[cell].font=bfont(bold=True); ws[cell].fill=sf()
            ws[cell].alignment=Alignment(horizontal="right",vertical="center"); ws[cell].border=bd()
        for cell in [f"C{i}",f"G{i}"]:
            ws[cell].font=bfont(); ws[cell].alignment=Alignment(horizontal="left",vertical="center",wrap_text=True); ws[cell].border=bd()
        sr(ws,i,22)

    part_cols=[("A","氏名",16),("B","所属機関",20),("C","役割・職種",14),("D","参加形態",12)]
    hrow(ws,6,part_cols)
    ws.merge_cells("E6:G6"); ws["E6"].value="備考"
    ws["E6"].font=hfont(sz=9); ws["E6"].fill=hf(); ws["E6"].alignment=Alignment(horizontal="center",vertical="center"); ws["E6"].border=bd(); sr(ws,6,28)

    for row in range(7,13):
        for col in "ABCD": bc(ws,f"{col}{row}")
        ws.merge_cells(f"E{row}:G{row}"); bc(ws,f"E{row}")
        sr(ws,row,20)

    secs=[
        (13,"【退院時の状態・引き継ぎ情報】",60,"・診断名・障害名：\n・現在の服薬：\n・医療処置：\n・ADL・機能レベル：\n・精神状態："),
        (15,"【退院後の訪問看護計画】",60,"・訪問頻度：週　回（当初2週間：週　回）\n・主な看護内容：\n・緊急時連絡体制："),
        (17,"【関係機関との役割分担・連携事項】",50,"・訪問看護：\n・相談支援専門員：\n・居宅介護：\n・主治医："),
        (19,"【退院後に必要な手続き・準備事項】",50,"・介護保険：申請済み / 未申請\n・自立支援医療：更新期限確認\n・障害福祉サービス：サービス等利用計画の更新"),
    ]
    for rn,title,h,example in secs:
        ws.merge_cells(f"A{rn}:G{rn}"); ws[f"A{rn}"].value=title
        ws[f"A{rn}"].font=hfont(sz=9); ws[f"A{rn}"].fill=hf()
        ws[f"A{rn}"].alignment=Alignment(horizontal="left",vertical="center"); ws[f"A{rn}"].border=bd(); sr(ws,rn,20)
        ws.merge_cells(f"A{rn+1}:G{rn+1}"); c=ws[f"A{rn+1}"]
        c.value=example; c.font=Font(name="Meiryo UI",size=9,color="FF6B7280",italic=True)
        c.fill=gf(); c.alignment=Alignment(horizontal="left",vertical="top",wrap_text=True); c.border=bd(); sr(ws,rn+1,h)

    ws.merge_cells("A21:G21")
    ws["A21"].value="記録者署名：　　　　　　　　　　　日付：　　年　月　日　　（本記録は5年間保管）"
    ws["A21"].font=bfont(); ws["A21"].alignment=Alignment(horizontal="left",vertical="center"); ws["A21"].border=bd(); sr(ws,21,26)

    for col,w in zip("ABCDEFG",[16,20,14,12,12,12,14]): ws.column_dimensions[col].width=w
    save(wb,"02_訪問看護記録様式","11_退院前カンファレンス記録.xlsx")


# ===== 守秘義務誓約書（Word）=====
def make_himitsu():
    doc=Document()
    doc.styles["Normal"].font.name="Meiryo UI"
    doc.styles["Normal"].font.size=Pt(10.5)
    sec=doc.sections[0]; sec.top_margin=Cm(3); sec.bottom_margin=Cm(3)
    sec.left_margin=Cm(3.5); sec.right_margin=Cm(3)

    h=doc.add_heading("秘密保持に関する誓約書", level=1)
    h.runs[0].font.color.rgb=RGBColor(0x0E,0x74,0x90)
    h.runs[0].font.size=Pt(16)
    h.alignment=WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p=doc.add_paragraph()
    p.add_run("　私は、○○事業所（以下「事業所」という。）に在職（又は就業）するにあたり、以下の事項を遵守することを誓約いたします。")

    items=[
        "業務を通じて知り得た利用者及びその家族に関する情報（氏名・病状・家族状況・生活状況・医療情報・財産状況等）を、業務遂行の目的以外に使用しないこと。",
        "利用者及びその家族に関する情報を、正当な理由なく第三者（家族・友人・SNS等を含む）に漏洩しないこと。",
        "業務上知り得た事業所の経営情報・職員情報・サービス内容・内部情報を外部に漏洩しないこと。",
        "個人情報が記録された書類・データ（紙・電子媒体）を許可なく持ち出し・複製・送信しないこと。",
        "退職・契約終了後も、在職中に知り得た情報の秘密保持義務を継続して遵守すること。",
        "本誓約書に違反した場合は、事業所が被った損害を賠償する責任を負うことを了解すること。",
        "個人情報の保護に関する法律（個人情報保護法）及び関連法令を遵守すること。",
    ]
    for i,item in enumerate(items,start=1):
        p=doc.add_paragraph(style="List Number")
        p.add_run(item)

    doc.add_paragraph()
    doc.add_paragraph("　上記の各事項について確認し、誓約します。")
    doc.add_paragraph()
    doc.add_paragraph("　　　　　　　年　　月　　日")
    doc.add_paragraph()
    doc.add_paragraph("氏名（自署）：　　　　　　　　　　　　　　　　　　　　　　　　　（印）")
    doc.add_paragraph()
    doc.add_paragraph("住所：　　　　　　　　　　　　　　　　　　　　　　　　　　　　　")
    doc.add_paragraph()
    doc.add_paragraph("【事業所保管用】　事業所名：○○事業所　　受領者（管理者）確認：　　　　　　　　")

    save_doc(doc,"07_運営規程ひな形","04_秘密保持誓約書.docx")


# ===== 身体拘束禁止方針・チェックシート（Word）=====
def make_taibatsu():
    doc=Document()
    doc.styles["Normal"].font.name="Meiryo UI"
    doc.styles["Normal"].font.size=Pt(10.5)
    sec=doc.sections[0]; sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
    sec.left_margin=Cm(3.5); sec.right_margin=Cm(3)

    h=doc.add_heading("身体拘束禁止・虐待防止方針", level=1)
    h.runs[0].font.color.rgb=RGBColor(0x0E,0x74,0x90)
    h.runs[0].font.size=Pt(14)
    h.alignment=WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p=doc.add_paragraph()
    p.add_run("　○○事業所は、利用者の尊厳と権利を守るため、以下の方針を定め、全職員がこれを遵守します。")

    secs=[
        ("1. 身体拘束の禁止",
         "当事業所では、緊急やむを得ない場合を除き、いかなる形態の身体拘束も行いません。\n\n「緊急やむを得ない」の3要件（すべて満たす必要あり）：\n①切迫性：利用者本人または他の利用者等の生命・身体が危険にさらされている\n②非代替性：拘束以外の方法では危険を回避できない\n③一時性：一時的なものであり、できる限り早急に拘束を解除する\n\n上記3要件を満たす場合でも、管理者・家族への説明と同意、記録の作成が必要です。"),
        ("2. 禁止される行為（例示）",
         "・車椅子・椅子・ベッドへのベルト・抑制帯による固定\n・ミトン型手袋の装着\n・四肢をひも等で縛ること\n・向精神薬の過剰投与・不適切投与による行動抑制\n・ドアの施錠による隔離\n・立ち上がれないよう椅子の高さを調節すること"),
        ("3. 虐待の防止",
         "虐待の種類：身体的虐待・精神的虐待・性的虐待・経済的虐待・放棄（ネグレクト）\n\n職員の通報義務：職員が虐待を発見した場合は、速やかに管理者に報告し、市区町村（障害者虐待防止センター）に通報します。\n\n虐待防止委員会：管理者が委員長となり、年1回以上の開催と研修を実施します。"),
        ("4. 研修・周知",
         "全職員に対して年1回以上、身体拘束禁止・虐待防止に関する研修を実施し、記録を保管します。"),
        ("附則",
         "本方針は令和　　年　　月　　日から施行します。\n事業所名：○○事業所　管理者：　　　　　　　（署名）"),
    ]
    for title,content in secs:
        p=doc.add_paragraph()
        p.add_run(title).bold=True
        p.paragraph_format.space_before=Pt(12)
        doc.add_paragraph(content)

    save_doc(doc,"07_運営規程ひな形","05_身体拘束禁止・虐待防止方針.docx")


# ===== 相談支援用 運営規程ひな形（Word）=====
def make_soudan_uneikitei():
    doc=Document()
    doc.styles["Normal"].font.name="Meiryo UI"
    doc.styles["Normal"].font.size=Pt(10)
    sec=doc.sections[0]; sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
    sec.left_margin=Cm(3); sec.right_margin=Cm(2.5)

    h=doc.add_heading("特定相談支援事業所 運営規程（ひな形）", level=1)
    h.runs[0].font.color.rgb=RGBColor(0x0E,0x74,0x90)
    h.runs[0].font.size=Pt(14)

    doc.add_paragraph("※ このひな形は参考例です。一般相談支援（地域移行・地域定着）も行う場合は第6条のサービス内容に追記してください。")

    chapters=[
        ("第1条（目的）","この運営規程は、○○相談支援事業所（以下「当事業所」という。）が実施する特定相談支援事業の運営の適正を確保するために必要な事項を定め、障碍者が自立した日常生活または社会生活を営めるよう支援することを目的とする。"),
        ("第2条（運営の方針）","①利用者の意思及び人格を尊重し、公正中立な立場でサービス等利用計画を作成する。\n②常に利用者の立場に立ち、その意向を尊重した総合的な相談支援を行う。\n③関係機関・事業所と連携し、地域のネットワークを活かした支援を行う。\n④個人情報を適切に管理する。\n⑤利用者の人権擁護・虐待防止のために必要な措置を講じる。"),
        ("第3条（事業所の名称及び所在地）","事業所の名称：○○相談支援事業所\n事業所の所在地：京都府○○市○○町○○番地"),
        ("第4条（従業者の職種、員数及び職務内容）","（１）管理者：1名（常勤・専従）事業所全体の管理\n（２）相談支援専門員：1名以上（常勤換算）\n　資格要件：相談支援従事者初任者研修修了・実務経験（3〜5年以上）\n　職務：サービス等利用計画の作成・モニタリング・相談支援"),
        ("第5条（営業日及び営業時間）","営業日：月曜日から金曜日（祝日・年末年始を除く）\n営業時間：午前9時00分から午後5時30分まで\n相談受付時間：上記営業時間内"),
        ("第6条（提供するサービスの内容）","（１）基本相談支援：福祉に関する様々な問題に関する相談に応じ、必要な情報提供・助言等を行う。\n（２）計画相談支援\n　ア）サービス利用支援：利用者のサービス等利用計画を作成し、サービス事業所等との連絡調整を行う。\n　イ）継続サービス利用支援（モニタリング）：定期的に計画の達成状況を評価・見直しを行う。"),
        ("第7条（利用料）","指定相談支援（計画相談支援・地域相談支援）は、全額給付費で賄われるため、利用者の負担はありません。ただし、計画作成に伴う実費（交通費等）は利用者の負担となります。"),
        ("第8条（苦情処理）","苦情受付窓口：管理者（TEL：○○）\n第三者委員：○○様\n受付方法：面接・電話・書面\n処理手順：受付→記録→調査→対応→報告→記録保管"),
        ("第9条（虐待の防止）","虐待防止責任者を管理者と定め、年1回以上の研修を実施する。発見した場合は市区町村に通報する。"),
        ("第10条（秘密保持）","相談を通じて知り得た個人情報を正当な理由なく第三者に漏洩しない。退職後も同様。"),
        ("附則","この運営規程は、令和○年○月○日から施行する。"),
    ]
    for title,content in chapters:
        p=doc.add_paragraph(); p.add_run(title).bold=True
        p.paragraph_format.space_before=Pt(12)
        doc.add_paragraph(content)

    save_doc(doc,"07_運営規程ひな形","06_運営規程ひな形_特定相談支援事業所.docx")


if __name__=="__main__":
    print("=== 追加テンプレートバッチ3 ===")
    print("\n--- 就労支援 ---")
    make_chingin()
    make_gyomu()
    print("\n--- 訪問看護 ---")
    make_seishin_keikaku()
    make_seishin_houkoku()
    make_taiin_conf()
    print("\n--- Word書類 ---")
    make_himitsu()
    make_taibatsu()
    make_soudan_uneikitei()
    print("\n=== 完了 ===")
